"""
Cortex AI Agent IDE — Main Window
Full 3-panel layout: Sidebar | Editor Tabs | AI Chat + Terminal
"""

import json
import os
import sys
import platform
from pathlib import Path
from typing import Optional
from PyQt6.QtWidgets import (
    QMainWindow, QWidget, QHBoxLayout, QVBoxLayout, QSplitter,
    QTabWidget, QLabel, QPushButton, QStatusBar, QFileDialog,
    QToolBar, QMenuBar, QMenu, QMessageBox, QInputDialog, QTabBar,
    QFrame, QSizePolicy, QApplication, QListWidget, QListWidgetItem, QComboBox, QDialog,
    QStackedWidget, QScrollArea, QTreeView, QLineEdit
)
from PyQt6.QtCore import Qt, QSize, pyqtSignal, pyqtSlot, QTimer, QRect, QProcessEnvironment, QSignalBlocker, QEventLoop, QDir, QModelIndex, QThread
from PyQt6.QtGui import QFileSystemModel, QAction, QKeySequence, QIcon, QFont, QPainter, QColor, QMouseEvent, QCloseEvent, QPixmap
from PyQt6.QtWebEngineCore import QWebEnginePage

from src.config.settings import get_settings
from src.config.theme_manager import get_theme_manager
from src.config.points_manager import get_points_manager, InsufficientPointsError
from src.core.project_manager import ProjectManager
from src.core.file_manager import FileManager
from src.core.session_manager import SessionManager
from src.core.codebase_index import get_codebase_index
# Agent bridge - connects Cortex to agent module
try:
    from src.ai.agent_bridge import get_agent_bridge as AIAgent
    HAS_AGENT_BRIDGE = True
except ImportError:
    from src.ai.stub_agent import get_stub_agent as AIAgent  # Fallback stub
    HAS_AGENT_BRIDGE = False

# File tree delegate for chevron arrows
from src.ui.components.sidebar import FileTreeDelegate
import subprocess as _subprocess
import logging as _logging

_log = _logging.getLogger(__name__)


class _GitStatusWorker(QThread):
    """Background worker that collects git status without blocking the UI."""
    status_ready = pyqtSignal(dict)

    def __init__(self, repo_path: str, git_manager, gh_cached: object = None):
        super().__init__()
        self._repo_path = repo_path
        self._git_manager = git_manager
        self._gh_cached = gh_cached  # None = not checked yet

    # ---- helpers (run off-thread) ------------------------------------
    @staticmethod
    def _parse_numstat(output: str):
        for line in output.strip().splitlines():
            parts = line.split('\t')
            if len(parts) >= 2:
                try:
                    a = int(parts[0]) if parts[0] != '-' else 0
                    d = int(parts[1]) if parts[1] != '-' else 0
                    return a, d
                except ValueError:
                    pass
        return 0, 0

    def _get_file_diff_stats(self, file_path: str):
        total_add, total_del = 0, 0
        try:
            # FIX: Prevent console window popup
            kwargs = dict(cwd=self._repo_path, capture_output=True, text=True, timeout=5)
            if sys.platform == 'win32':
                kwargs['creationflags'] = subprocess.CREATE_NO_WINDOW
            
            r = _subprocess.run(
                ["git", "diff", "--numstat", file_path],
                **kwargs
            )
            if r.returncode == 0 and r.stdout.strip():
                a, d = self._parse_numstat(r.stdout)
                total_add += a; total_del += d
        except Exception:
            pass
        try:
            # FIX: Prevent console window popup
            kwargs = dict(cwd=self._repo_path, capture_output=True, text=True, timeout=5)
            if sys.platform == 'win32':
                kwargs['creationflags'] = subprocess.CREATE_NO_WINDOW
            
            r = _subprocess.run(
                ["git", "diff", "--cached", "--numstat", file_path],
                **kwargs
            )
            if r.returncode == 0 and r.stdout.strip():
                a, d = self._parse_numstat(r.stdout)
                total_add += a; total_del += d
        except Exception:
            pass
        return total_add, total_del

    # ---- main work (runs in QThread) --------------------------------
    def run(self):
        data = {}
        # 1) branch name
        try:
            # FIX: Prevent console window popup
            kwargs = dict(cwd=self._repo_path, capture_output=True, text=True, timeout=5)
            if sys.platform == 'win32':
                kwargs['creationflags'] = subprocess.CREATE_NO_WINDOW
            
            r = _subprocess.run(
                ["git", "rev-parse", "--abbrev-ref", "HEAD"],
                **kwargs
            )
            data['branch'] = r.stdout.strip() if r.returncode == 0 else None
        except Exception:
            data['branch'] = None

        # 2) git file status
        try:
            git_files = self._git_manager.get_status()
        except Exception:
            git_files = []
        data['git_files'] = git_files
        data['unstaged'] = sum(1 for f in git_files if not f.staged)
        data['untracked'] = sum(1 for f in git_files if f.status.name == 'UNTRACKED')
        data['staged'] = sum(1 for f in git_files if f.staged)
        data['total'] = len(git_files)

        # 3) gh cli (cache-aware)
        if self._gh_cached is not None:
            data['gh'] = self._gh_cached
        else:
            try:
                # FIX: Prevent console window popup
                kwargs = dict(capture_output=True, text=True, timeout=5)
                if sys.platform == 'win32':
                    kwargs['creationflags'] = subprocess.CREATE_NO_WINDOW
                
                r = _subprocess.run(
                    ["gh", "--version"],
                    **kwargs
                )
                if r.returncode == 0:
                    data['gh'] = r.stdout.strip().split('\n')[0]
                else:
                    data['gh'] = ''
            except FileNotFoundError:
                data['gh'] = ''
            except Exception:
                data['gh'] = ''

        # 4) per-file diff stats (batched, max 100)
        file_stats = []
        seen = set()
        for gf in git_files[:100]:
            if gf.path in seen:
                continue
            seen.add(gf.path)
            a, d = self._get_file_diff_stats(gf.path)
            file_stats.append((gf, a, d))
        data['file_stats'] = file_stats

        self.status_ready.emit(data)


class CodeAnalyzer:
    """Simple prompt builder for AI code actions."""

    def build_explain_prompt(self, code: str, language: str) -> str:
        return f"Explain this {language} code in detail. Break down what each part does and why:\n\n```{language}\n{code}\n```"

    def build_refactor_prompt(self, code: str, language: str) -> str:
        return f"Refactor this {language} code to be cleaner, more efficient, and follow best practices. Explain your changes:\n\n```{language}\n{code}\n```"

    def build_test_prompt(self, code: str, language: str) -> str:
        return f"Write comprehensive unit tests for this {language} code. Include edge cases and error handling:\n\n```{language}\n{code}\n```"

    def build_debug_prompt(self, code: str, error: str, language: str) -> str:
        return f"Help me debug this {language} code. Error: {error}\n\n```{language}\n{code}\n```\n\nWhat's causing this error and how do I fix it?"
# from src.ai.file_edit_tracker import FileEditTracker
from src.core.git_manager import GitManager
from src.ui.components.sidebar import SidebarWidget
# CommandPalette removed - not implemented in AI-first mode
# from src.ui.components.command_palette import CommandPalette
from src.ui.components.editor import CodeEditor
from src.ui.components.ai_chat import AIChatWidget
from src.ui.components.xterm_terminal import XTermWidget
from src.ui.components.find_replace import FindReplaceDialog

from src.ui.dialogs.diff_viewer import DiffWindow
from src.utils.icons import make_icon
# Live Server removed in AI-first mode - AI handles code execution
# from src.core.live_server import LiveServer
from src.utils.helpers import detect_language, shorten_path
from src.utils.logger import get_logger
from src.utils.notifications import show_task_complete_notification

log = get_logger("main_window")

try:
    from src.ui.syntax_highlighting_config import (
        UniversalCodeColorizer,
        MarkdownColorizer, 
        DRACULA_COLORS,
        FONTS
    )
    HAS_SYNTAX_HIGHLIGHTING = True
except ImportError:
    HAS_SYNTAX_HIGHLIGHTING = False
    log.warning("Syntax highlighting module not available")


class CleanTabBar(QTabBar):
    """Tab bar that draws a clean × close button instead of Qt's default box."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setMouseTracking(True)
        self._hovered_tab = -1
        # Inherit theme from parent window if possible
        from src.config.theme_manager import get_theme_manager
        self._is_dark = get_theme_manager().is_dark


    def set_dark(self, is_dark: bool):
        self._is_dark = is_dark
        self.update()

    def mouseMoveEvent(self, event):
        idx = self.tabAt(event.pos())
        if idx != self._hovered_tab:
            self._hovered_tab = idx
            self.update()
        super().mouseMoveEvent(event)

    def leaveEvent(self, event):
        self._hovered_tab = -1
        self.update()
        super().leaveEvent(event)

    def tabSizeHint(self, index):
        s = super().tabSizeHint(index)
        s.setHeight(34)
        return s

    def paintEvent(self, event):
        """Draw tabs with a clean × button — theme-aware."""
        from PyQt6.QtGui import QPainter, QColor
        from PyQt6.QtCore import QRect, Qt
        painter = QPainter(self)
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)

        d = self._is_dark
        # Cursor IDE Anysphere Dark Theme - Tab Colors
        col_sel_bg    = QColor("#181818") if d else QColor("#fafafa")  # editor.background
        col_hover_bg  = QColor("#1f1f1f") if d else QColor("#f0f0f0")  # tab.hoverBackground
        col_normal_bg = QColor("#141414") if d else QColor("#f3f3f3")  # tab.inactiveBackground
        col_accent    = QColor("#228df2") if d else QColor("#0078d4")  # cursor.accent
        col_divider   = QColor("#2a2a2a") if d else QColor("#e5e5e5")  # sideBar.border
        col_sel_fg    = QColor("#ffffff") if d else QColor("#2c2c2c")  # tab.activeForeground
        col_hover_fg  = QColor("#d6d6dd") if d else QColor("#333333")  # editor.foreground
        col_normal_fg = QColor("#6d6d6d") if d else QColor("#616161")  # tab.inactiveForeground
        col_close     = QColor("#f14c4c") if d else QColor("#d73a49")  # RED close button
        col_close_hover = QColor("#ff6b6b") if d else QColor("#ff4444")  # Brighter red on hover

        for i in range(self.count()):
            rect = self.tabRect(i)
            is_selected = (i == self.currentIndex())
            is_hovered  = (i == self._hovered_tab)

            # Background
            if is_selected:
                painter.fillRect(rect, col_sel_bg)
            elif is_hovered:
                painter.fillRect(rect, col_hover_bg)
            else:
                painter.fillRect(rect, col_normal_bg)

            # Accent top border on active tab
            if is_selected:
                painter.fillRect(QRect(rect.x(), rect.y(), rect.width(), 2), col_accent)

            # Right divider
            painter.fillRect(QRect(rect.right(), rect.y() + 4, 1, rect.height() - 8), col_divider)

            # Tab label
            text = self.tabText(i)
            
            # ── Premium Tab Icon ───────────────────────────────────────────────
            icon_x = rect.x() + 10
            icon_size = 14
            
            # Determine icon based on file extension
            filepath = None
            if hasattr(self.parent(), '_files'):
                filepath = self.parent()._files.get(i)
            
            from src.ui.components.sidebar import _get_icon_name
            icon_name = _get_icon_name(filepath) if filepath else "files"
            if text == "Welcome": icon_name = "ai"
            if "Terminal" in text: icon_name = "terminal"
            
            # Cursor IDE Syntax Colors for File Icons
            colors = {
                "python":   "#83d6c5",  # teal - keyword
                "html":     "#87c3ff",  # light blue - class/tag
                "css":      "#87c3ff",  # light blue - class
                "javascript": "#e394dc", # pink - string
                "typescript": "#87c3ff", # light blue - class
                "markdown": "#d6d6dd", # primary text
                "json":     "#efb080",  # orange - number
                "java":     "#83d6c5",  # teal
                "rust":     "#efb080",  # orange
                "go":       "#87c3ff",  # light blue
                "sql":      "#83d6c5",  # teal
                "ai":       "#228df2",  # accent blue
                "terminal": "#6d6d6d"   # muted
            }
            icon_color = colors.get(icon_name, "#abb2bf")
            
            # Using a fallback if terminal icon not found, or use the shell icon
            icon_pixmap = make_icon(icon_name, icon_color, icon_size).pixmap(icon_size, icon_size)
            painter.drawPixmap(icon_x, rect.y() + (rect.height() - icon_size)//2, icon_pixmap)
            
            # Reserve space for close button (14px + 2px right padding)
            btn_reserved = 16  # 14 + 2px right margin
            
            # Label - leave space for close button
            label_x = icon_x + icon_size + 6
            label_width = rect.width() - (label_x - rect.x()) - btn_reserved - 2
            label_rect = QRect(label_x, rect.y(), max(0, label_width), rect.height())
            
            # Draw label with eliding to prevent overflow
            painter.save()
            fg = col_sel_fg if is_selected else (col_hover_fg if is_hovered else col_normal_fg)
            painter.setPen(fg)
            f = painter.font()
            f.setPointSize(9)
            painter.setFont(f)
            # Use elided text to prevent overflow into button area
            from PyQt6.QtGui import QFontMetrics
            fm = QFontMetrics(f)
            elided_text = fm.elidedText(text, Qt.TextElideMode.ElideRight, label_width)
            painter.drawText(label_rect,
                             Qt.AlignmentFlag.AlignVCenter | Qt.AlignmentFlag.AlignLeft,
                             elided_text)
            painter.restore()

            # × close button — ONLY show on hovered or selected tabs
            if is_hovered or is_selected:
                # Fixed position from right edge - 2px padding
                btn_size = 14
                btn_x = rect.right() - btn_size - 2
                btn_y = rect.y() + (rect.height() - btn_size) // 2
                btn_rect = QRect(btn_x, btn_y, btn_size, btn_size)
                
                # Check if cursor is over the × button
                cursor_pos = self.mapFromGlobal(self.cursor().pos())
                is_close_hovered = btn_rect.contains(cursor_pos)

                # Determine colors
                if is_close_hovered:
                    bg_color = QColor("#f14c4c") if d else QColor("#d73a49")  # Red
                    x_color = QColor("#ffffff")  # White X
                else:
                    bg_color = QColor(255, 255, 255, 40) if d else QColor(0, 0, 0, 30)
                    x_color = QColor("#f14c4c") if d else QColor("#d73a49")  # Red X
                
                # CRITICAL: Fill button area with tab background color first
                bg_fill = col_sel_bg if is_selected else col_hover_bg
                painter.fillRect(btn_rect, bg_fill)
                
                # Draw button background
                painter.setBrush(bg_color)
                painter.setPen(Qt.PenStyle.NoPen)
                painter.drawRoundedRect(btn_rect, 2, 2)
                
                # Draw X centered
                painter.setPen(x_color)
                font = painter.font()
                font.setPointSize(10)
                font.setBold(True)
                painter.setFont(font)
                painter.drawText(btn_rect, Qt.AlignmentFlag.AlignCenter, "×")

        painter.end()

    def mousePressEvent(self, event):
        """Handle × button click to close the tab."""
        from PyQt6.QtCore import Qt, QRect
        from PyQt6.QtGui import QMouseEvent
        i = self.tabAt(event.pos())
        if i >= 0:
            # Only check if tab is hovered or selected (matching paint logic)
            is_selected = (i == self.currentIndex())
            is_hovered = (i == self._hovered_tab)
            if is_hovered or is_selected:
                rect = self.tabRect(i)
                btn_size = 14
                btn_x = rect.right() - btn_size - 2  # Match paintEvent
                btn_y = rect.y() + (rect.height() - btn_size) // 2
                btn_rect = QRect(btn_x, btn_y, btn_size, btn_size)
                if btn_rect.contains(event.pos()):
                    self.tabCloseRequested.emit(i)
                    return
        super().mousePressEvent(event)


class ClickableLabel(QLabel):
    """A label that behaves like a button — supports RichText and hover cursor."""
    clicked = pyqtSignal()
    
    def __init__(self, text, parent=None):
        super().__init__(text, parent)
        self.setCursor(Qt.CursorShape.PointingHandCursor)
        self.setTextFormat(Qt.TextFormat.RichText)
        
    def mousePressEvent(self, event: QMouseEvent):
        if event.button() == Qt.MouseButton.LeftButton:
            self.clicked.emit()
        super().mousePressEvent(event)



class EditorTabWidget(QTabWidget):
    """Central editor area with tabs."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setTabBar(CleanTabBar(self))
        self.setTabsClosable(True)
        self.setMovable(True)
        self.setDocumentMode(True)
        self.tabCloseRequested.connect(self._close_tab)
        self._files: dict[int, str] = {}   # tab_index -> filepath
        self._modified: set[str] = set()

    def open_diff_tab(self, file_path: str, original: str, modified: str, is_dark: bool = True):
        """Open a read-only diff tab next to the file tab. Tab name: '⟷ filename'."""
        import difflib
        from PyQt6.QtWidgets import QTextBrowser
        from PyQt6.QtGui import QFont
        from pathlib import Path as _P

        file_name = _P(file_path).name
        tab_label = f'⟷ {file_name}'

        # If a diff tab for this file already exists, switch to it
        for idx in range(self.count()):
            if self.tabText(idx) == tab_label:
                self.setCurrentIndex(idx)
                return

        # Build unified diff HTML
        diff_lines = list(difflib.unified_diff(
            original.splitlines(keepends=True),
            modified.splitlines(keepends=True),
            fromfile='Original',
            tofile='Modified',
            n=3
        ))

        bg   = '#1e1e1e' if is_dark else '#ffffff'
        fg   = '#cccccc' if is_dark else '#333333'
        add_bg = 'rgba(46,160,67,0.2)'  if is_dark else 'rgba(46,160,67,0.15)'
        add_fg = '#56d364'              if is_dark else '#1a7f37'
        rem_bg = 'rgba(248,81,73,0.2)' if is_dark else 'rgba(255,129,130,0.15)'
        rem_fg = '#f85149'              if is_dark else '#cf222e'
        info_fg = '#8b949e'             if is_dark else '#6e7781'

        def esc(t): return t.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')

        parts = [f"<div style='background:{bg};color:{fg};white-space:pre;"
                 f"font-family:\"Cascadia Code\",Consolas,monospace;font-size:13px;line-height:1.5;padding:10px;'>"]

        if not diff_lines:
            parts.append(f"<div style='color:{info_fg};padding:20px;'>No changes detected.</div>")
        else:
            for line in diff_lines:
                line = line.rstrip('\n')
                s = esc(line)
                if line.startswith('+++') or line.startswith('---'):
                    parts.append(f"<div style='color:{info_fg};font-weight:bold;background:rgba(128,128,128,0.1);padding-left:6px;'>{s}</div>")
                elif line.startswith('@@'):
                    parts.append(f"<div style='color:{info_fg};padding-left:6px;margin-top:6px;'>{s}</div>")
                elif line.startswith('+'):
                    parts.append(f"<div style='color:{add_fg};background:{add_bg};padding-left:6px;'>{s}</div>")
                elif line.startswith('-'):
                    parts.append(f"<div style='color:{rem_fg};background:{rem_bg};padding-left:6px;'>{s}</div>")
                else:
                    parts.append(f"<div style='color:{fg};padding-left:6px;'>{s}</div>")
        parts.append('</div>')

        browser = QTextBrowser()
        browser.setReadOnly(True)
        browser.setOpenExternalLinks(False)
        font = QFont('Cascadia Code', 10)
        font.setStyleHint(QFont.StyleHint.Monospace)
        browser.setFont(font)
        browser.setStyleSheet(f'background:{bg};border:none;')
        browser.setHtml(''.join(parts))

        idx = self.addTab(browser, tab_label)
        self.setCurrentIndex(idx)

    def open_file(self, filepath: str, content: str, language: str, is_dark: bool = True) -> int:
        """Open a file in a new tab (or switch to existing)."""
        # Check if already open
        for idx, fp in self._files.items():
            if fp == filepath:
                # File already open - check if content matches, if not update it
                editor = self.widget(idx)
                if isinstance(editor, CodeEditor):
                    current_content = editor.toPlainText()
                    if current_content != content:
                        # Content changed, reload it
                        with QSignalBlocker(editor.document()):
                            editor.set_content(content, language, filepath)
                        print(f"[EditorTabs] Updated content for already-open file: {filepath}")
                    else:
                        print(f"[EditorTabs] File already open with same content: {filepath}")
                self.setCurrentIndex(idx)
                return idx
        
        # Create editor
        editor = CodeEditor(language=language)
        
        # Apply current theme to the new editor
        editor.set_theme(is_dark)
        
        # Disconnect the internal document→editor connection temporarily
        # Use blockSignals instead of disconnect to avoid errors
        with QSignalBlocker(editor.document()):
            editor.set_content(content, language, filepath)
        
        # NOW connect OUR handler - anything after this is a user edit
        editor.content_modified.connect(lambda: self._mark_modified(filepath))

        name = Path(filepath).name
        idx = self.addTab(editor, name)
        self._files[idx] = filepath
        self.setCurrentIndex(idx)
        self.setTabToolTip(idx, filepath)
        
        # Set file type icon on tab
        self._set_tab_icon(idx, filepath)
        
        return idx
    
    def _set_tab_icon(self, idx: int, filepath: str):
        """Set the tab icon based on file extension."""
        from PyQt6.QtGui import QIcon, QPixmap
        from pathlib import Path
        
        ext = Path(filepath).suffix.lower()
        icon_map = {
            '.py': 'python.png',
            '.js': 'javascript.png',
            '.ts': 'typescript.png',
            '.jsx': 'javascript.png',
            '.tsx': 'typescript.png',
            '.html': 'html.png',
            '.htm': 'html.png',
            '.css': 'css.png',
            '.json': 'json.png',
            '.md': 'markdown.png',
            '.java': 'java.png',
            '.rs': 'rust.png',
            '.csv': 'csv.png',
            '.env': 'env.png',
        }
        
        icon_file = icon_map.get(ext)
        if icon_file:
            icon_path = Path(__file__).parent / 'assets' / 'icons' / icon_file
            if icon_path.exists():
                pixmap = QPixmap(str(icon_path))
                if not pixmap.isNull():
                    # Scale to appropriate size for tab (16x16)
                    scaled = pixmap.scaled(16, 16, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
                    self.setTabIcon(idx, QIcon(scaled))

    def _mark_modified(self, filepath: str):
        """Mark file as modified and update tab with white dot."""
        self._modified.add(filepath)
        for idx, fp in self._files.items():
            if fp == filepath:
                name = Path(fp).name
                # Set tab text with white dot (●) prefix
                self.setTabText(idx, f"● {name}")
                # Set tab tooltip to show modified status
                self.setTabToolTip(idx, f"{filepath} (Modified)")
                break
    
    def _mark_saved(self, filepath: str):
        """Mark file as saved and remove dot from tab."""
        self._modified.discard(filepath)
        for idx, fp in self._files.items():
            if fp == filepath:
                name = Path(fp).name
                # Remove dot and restore normal tab text
                self.setTabText(idx, name)
                # Restore normal tooltip
                self.setTabToolTip(idx, filepath)
                break

    def _close_tab(self, index: int):
        """Close tab with save confirmation if modified."""
        widget = self.widget(index)
        filepath = self._files.get(index)
        
        # Check if file has unsaved changes
        if filepath and filepath in self._modified:
            from PyQt6.QtWidgets import QMessageBox
            
            reply = QMessageBox.question(
                self,
                "Save Changes?",
                f"Do you want to save the changes to '{Path(filepath).name}'?",
                QMessageBox.StandardButton.Save | 
                QMessageBox.StandardButton.Discard | 
                QMessageBox.StandardButton.Cancel
            )
            
            if reply == QMessageBox.StandardButton.Cancel:
                return  # User cancelled - don't close tab
            elif reply == QMessageBox.StandardButton.Save:
                # Save the file before closing
                editor = self.current_editor() if index == self.currentIndex() else None
                if editor:
                    content = editor.get_all_text()
                    try:
                        Path(filepath).write_text(content, encoding='utf-8')
                        self._mark_saved(filepath)
                    except Exception as e:
                        from PyQt6.QtWidgets import QMessageBox
                        QMessageBox.critical(self, "Save Error", f"Failed to save file: {e}")
                        return  # Don't close tab on save error
        
        # Proceed with closing tab
        self.removeTab(index)
        
        # Cleanup files mapping
        new_files = {}
        for old_idx, fp in self._files.items():
            if old_idx == index:
                continue
            new_idx = old_idx if old_idx < index else old_idx - 1
            new_files[new_idx] = fp
        self._files = new_files
        
        if filepath:
            self._modified.discard(filepath)
            
        # If no tabs left, we could show welcome again or just keep it empty
        if self.count() == 0:
            # Maybe tell parent to show welcome? 
            # For now just let it be empty to match VS Code behavior
            pass

    def current_editor(self) -> CodeEditor | None:
        w = self.currentWidget()
        return w if isinstance(w, CodeEditor) else None

    def current_filepath(self) -> str | None:
        return self._files.get(self.currentIndex())

    def save_current(self, file_manager: FileManager) -> bool:
        """Save current file and remove modified indicator."""
        editor = self.current_editor()
        fp = self.current_filepath()
        if not editor or not fp:
            return False
        content = editor.get_all_text()
        ok = file_manager.write(fp, content)
        if ok:
            # Use _mark_saved to properly update tab text and tooltip
            self._mark_saved(fp)
        return ok
    
    def save_file(self, filepath: str, content: str) -> bool:
        """Save a specific file and update its modified state."""
        from pathlib import Path as _Path
        try:
            _Path(filepath).write_text(content, encoding='utf-8')
            self._mark_saved(filepath)
            return True
        except Exception as e:
            print(f"Save error: {e}")
            return False

    def get_open_files(self) -> list[str]:
        return list(self._files.values())

    def close_current_tab(self):
        """Close the currently active tab."""
        current_idx = self.currentIndex()
        if current_idx >= 0:
            self._close_tab(current_idx)

    def close_all_tabs(self):
        """Close all open tabs."""
        while self.count() > 0:
            self._close_tab(0)

    def update_theme(self, is_dark: bool):
        # Update tab bar colours
        if isinstance(self.tabBar(), CleanTabBar):
            self.tabBar().set_dark(is_dark)

        # Update individual editor widgets
        for i in range(self.count()):
            w = self.widget(i)
            if isinstance(w, CodeEditor):
                w.set_theme(is_dark)


class CortexMainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        log.info("MainWindow: __init__ START")
        log.info("MainWindow: Initializing managers...")
        self._settings = get_settings()
        self._theme_manager = get_theme_manager()
        self._project_manager = ProjectManager()
        self._file_manager = FileManager()
        self._session_manager = SessionManager()
        self._live_server: Optional[LiveServer] = None  # built-in HTML Live Server
        # Git manager for source control integration
        self._git_manager = GitManager()
        self._git_manager.status_changed.connect(self._update_git_summary)
        log.info("[GIT] GitManager initialized")
        
        # Agent bridge connects Cortex to agent module
        self._ai_agent = AIAgent(file_manager=self._file_manager)
        if HAS_AGENT_BRIDGE:
            log.info("[AGENT] Agent bridge initialized - full integration active")
        else:
            log.info("[AGENT] Stub agent initialized - bridge not available")
        
        # Legacy AI components removed - agent_bridge.py is the active runtime
        self._file_tracker = None
        self._diff_window = DiffWindow(self)
        self._codebase_index = None
        self._inline_edit_context = None
        # Live Server removed in AI-first mode
        self._live_server = None
        
        # Initialize UI components to None to prevent theme application crashes if build fails
        self._toolbar = None
        self._toolbar_sep = None
        self._toolbar_logo = None
        self._toolbar_btns = []
        self._memory_btn = None
        self._settings_btn = None

        try:
            log.info("MainWindow: Building UI...")
            self._build_ui()
            log.info("MainWindow: Building Menu...")
            # Build menu bar for all modes (Codex-style has menu bar)
            self._build_menu()
            log.info("MainWindow: Building Status Bar...")
            self._build_status_bar()
        except Exception as e:
            log.error(f"UI Build Error: {e}", exc_info=True)
            print(f"UI Build Error: {e}")
            raise  # Re-raise to stop execution

        log.info("MainWindow: Connecting signals...")
        self._connect_signals()
        log.info("MainWindow: Applying initial theme...")
        self._apply_initial_theme()
        log.info("MainWindow: Restoring session...")
        self._restore_session()
        log.info("MainWindow: Initialization complete.")

        # Enable drag and drop of folders/files onto the main window
        self.setAcceptDrops(True)

        # Set Window Icon (Title Bar + Taskbar) - BEFORE show() to prevent flash
        # Uses pre-generated taskbar_rounded.png (run generate_icons.py once)
        if getattr(sys, 'frozen', False):
            base = sys._MEIPASS
        else:
            base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        logo_dir = os.path.join(base, "src", "assets", "logo")
        if not os.path.isdir(logo_dir):
            # Fallback: try exe directory
            exe_dir = os.path.dirname(sys.executable) if getattr(sys, 'frozen', False) else os.getcwd()
            logo_dir = os.path.join(exe_dir, "src", "assets", "logo")
        if not os.path.isdir(logo_dir):
            # Fallback: try _internal directory (PyInstaller onedir)
            exe_dir = os.path.dirname(sys.executable) if getattr(sys, 'frozen', False) else os.getcwd()
            logo_dir = os.path.join(exe_dir, "_internal", "src", "assets", "logo")
        if not os.path.isdir(logo_dir):
            logo_dir = os.path.join(os.getcwd(), "src", "assets", "logo")

        icon_candidates = [
            os.path.join(logo_dir, "taskbar_rounded.png"),
            os.path.join(logo_dir, "taskbar.png"),
            os.path.join(logo_dir, "taskbar.ico"),
        ]

        icon = QIcon()
        found_icon = False
        for candidate in icon_candidates:
            if os.path.exists(candidate):
                from PyQt6.QtGui import QPixmap
                from PyQt6.QtCore import Qt
                pm = QPixmap(candidate)
                if not pm.isNull():
                    for sz in [16, 32, 48, 64, 128, 256]:
                        icon.addPixmap(pm.scaled(sz, sz, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation))
                    found_icon = True
                    log.info(f"[ICON] Successfully loaded icon: {candidate}")
                    break

        if not found_icon:
            log.error(f"[ICON] No valid icon found in candidates: {icon_candidates}")
            log.error(f"[ICON] Checked paths: {[os.path.abspath(candidate) for candidate in icon_candidates]}")

        if not icon.isNull():
            self.setWindowIcon(icon)
            # Also set app-level icon for taskbar grouping
            from PyQt6.QtWidgets import QApplication
            QApplication.instance().setWindowIcon(icon)
        else:
            log.error("[ICON] Failed to set window icon: QIcon is null")

        # Window geometry
        w = self._settings.get("window", "width") or 1400
        h = self._settings.get("window", "height") or 900
        self.resize(w, h)
        self.setGeometry(100, 100, w, h)
        if self._settings.get("window", "maximized"):
            self.showMaximized()
        else:
            self.show()

    def _on_new_chat(self):
        """Handle new chat request from sidebar or navigation"""
        # Create new chat via JavaScript
        if hasattr(self, '_ai_chat') and self._ai_chat:
            self._ai_chat.run_javascript("window.newChat()")
            log.info("New chat created via sidebar")
        else:
            log.warning("AI chat not ready, cannot create new chat")
    
    def _on_settings_requested(self):
        """Handle settings request from top navigation bar"""
        self._open_settings()
    
    def _open_settings(self):
        """Open settings dialog (Ctrl+,)."""
        dialog = QDialog(self)
        dialog.setWindowTitle("Settings")
        dialog.setMinimumWidth(500)
        dialog.setMinimumHeight(400)
        
        layout = QVBoxLayout(dialog)
        
        # Theme setting
        theme_layout = QHBoxLayout()
        theme_label = QLabel("Theme:")
        theme_combo = QComboBox()
        theme_combo.addItems(["Dark", "Light"])
        theme_combo.setCurrentText("Dark" if self._settings.theme == "dark" else "Light")
        theme_combo.currentTextChanged.connect(lambda t: self._set_theme(t.lower()))
        theme_layout.addWidget(theme_label)
        theme_layout.addWidget(theme_combo)
        theme_layout.addStretch()
        layout.addLayout(theme_layout)
        
        layout.addStretch()
        
        # Close button
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.accept)
        layout.addWidget(close_btn)
        
        dialog.exec()
    
    def _on_chat_rename_requested(self, item):
        """Handle double-click or pencil icon to rename a chat."""
        chat_id = item.data(Qt.ItemDataRole.UserRole)
        # Get current title from widget
        widget = self._sidebar_chat_list.itemWidget(item)
        current_title = ''
        if widget:
            label = widget.findChild(QLabel, 'chat_title_label')
            if label:
                current_title = label.text()
        if not current_title:
            current_title = item.text() or 'Untitled'
        
        from PyQt6.QtWidgets import QInputDialog
        new_title, ok = QInputDialog.getText(
            self, 'Rename Chat', 'New name:', text=current_title
        )
        if ok and new_title.strip():
            new_title = new_title.strip()
            # Update label in widget
            if widget:
                label = widget.findChild(QLabel, 'chat_title_label')
                if label:
                    label.setText(new_title)
            # Update in JS
            safe_title = new_title.replace("'", "\\'").replace('"', '\\"')
            self._ai_chat.run_javascript(f"""
                if (window.updateChatTitle) {{
                    window.updateChatTitle('{chat_id}', '{safe_title}');
                }}
            """)
            log.info(f"Chat renamed to: {new_title}")
    
    def _on_chat_context_menu(self, position):
        """Show context menu for chat items (rename, delete)."""
        item = self._sidebar_chat_list.itemAt(position)
        if not item:
            return
        
        menu = QMenu(self)
        rename_action = menu.addAction('✏️ Rename')
        rename_action.triggered.connect(lambda: self._on_chat_rename_requested(item))
        menu.addSeparator()
        delete_action = menu.addAction('🗑️ Delete Chat')
        delete_action.setStyleSheet('color: #ef4444;')
        delete_action.triggered.connect(lambda: self._delete_chat_by_id(item))
        menu.exec(self._sidebar_chat_list.viewport().mapToGlobal(position))
    
    def _delete_chat_by_id(self, item):
        """Delete a chat: remove from sidebar and from JS storage."""
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QHBoxLayout as _HBox, QLabel as _Label, QPushButton as _Btn
        from PyQt6.QtCore import Qt as _Qt
        chat_id = item.data(Qt.ItemDataRole.UserRole)
        widget = self._sidebar_chat_list.itemWidget(item)
        chat_title = 'this chat'
        if widget:
            label = widget.findChild(QLabel, 'chat_title_label')
            if label:
                chat_title = label.toolTip() or label.text()
        
        # --- Compact confirmation dialog ---
        dialog = QDialog(self)
        dialog.setWindowTitle('Delete Chat')
        dialog.setFixedSize(280, 110)
        dialog.setWindowFlags(Qt.WindowType.Dialog | Qt.WindowType.FramelessWindowHint | Qt.WindowType.WindowStaysOnTopHint)
        dialog.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        
        # Outer layout (no margins) — dialog itself is transparent
        outer_layout = QVBoxLayout(dialog)
        outer_layout.setContentsMargins(0, 0, 0, 0)
        outer_layout.setSpacing(0)
        
        # Inner container with dark background, border, and rounded corners
        container = QWidget()
        container.setStyleSheet("""
            QWidget {
                background-color: #151515;
                border: 1px solid #333;
                border-radius: 5px;
            }
        """)
        outer_layout.addWidget(container)
        
        # Content layout inside the styled container
        dlg_layout = QVBoxLayout(container)
        dlg_layout.setContentsMargins(20, 16, 20, 14)
        dlg_layout.setSpacing(16)
        
        # Message
        msg_lbl = _Label(f"Delete '{chat_title}'?")
        msg_lbl.setStyleSheet('color: #d0d0d0; font-size: 13px; font-weight: 500; background: transparent;')
        msg_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        dlg_layout.addWidget(msg_lbl)
        
        dlg_layout.addStretch()
        
        # Buttons row - centered equally
        btn_layout = _HBox()
        btn_layout.setContentsMargins(0, 0, 0, 0)
        btn_layout.setSpacing(12)
        btn_layout.addStretch(1)
        
        cancel_btn = _Btn('Cancel')
        cancel_btn.setFixedSize(80, 30)
        cancel_btn.setStyleSheet("""
            QPushButton {
                background-color: #2d2d2d;
                color: #c0c0c0;
                border: 1px solid #555;
                border-radius: 5px;
                font-size: 12px;
            }
            QPushButton:hover { background-color: #383838; }
        """)
        cancel_btn.clicked.connect(dialog.reject)
        btn_layout.addWidget(cancel_btn)
        
        delete_btn = _Btn('Delete')
        delete_btn.setFixedSize(80, 30)
        delete_btn.setStyleSheet("""
            QPushButton {
                background-color: #c0392b;
                color: #ffffff;
                border: none;
                border-radius: 5px;
                font-size: 12px;
            }
            QPushButton:hover { background-color: #e74c3c; }
        """)
        delete_btn.clicked.connect(dialog.accept)
        btn_layout.addWidget(delete_btn)
        
        btn_layout.addStretch(1)
        dlg_layout.addLayout(btn_layout)
        
        if dialog.exec() == QDialog.DialogCode.Accepted:
            row = self._sidebar_chat_list.row(item)
            self._sidebar_chat_list.takeItem(row)
            # Delete from JS (bypass JS confirm dialog by directly removing)
            self._ai_chat.run_javascript(f"""
                (function() {{
                    chats = chats.filter(function(c) {{ return c.id != '{chat_id}'; }});
                    if (typeof bridge !== 'undefined' && bridge.delete_chat_from_sqlite) {{
                        bridge.delete_chat_from_sqlite('{chat_id}');
                    }}
                    var key = typeof getStorageKey === 'function' ? getStorageKey() : '';
                    if (key) localStorage.removeItem(key);
                    if (typeof saveChats === 'function') saveChats();
                    if (typeof currentChatId !== 'undefined' && currentChatId == '{chat_id}') {{
                        if (chats.length > 0) loadChat(chats[0].id);
                        else startNewChat();
                    }}
                }})();
            """)
            log.info(f"Chat deleted: {chat_title} (ID: {chat_id})")
    
    def _on_chat_item_clicked(self, item):
        """Handle click on a chat item to load that chat."""
        chat_id = item.data(Qt.ItemDataRole.UserRole)
        # Get title from widget label
        widget = self._sidebar_chat_list.itemWidget(item)
        chat_title = ''
        if widget:
            label = widget.findChild(QLabel, 'chat_title_label')
            if label:
                chat_title = label.toolTip() or label.text()
        
        if chat_id and self._ai_chat:
            self._ai_chat.run_javascript(f"""
                if (window.loadChat) {{
                    window.loadChat('{chat_id}');
                    console.log('[CHAT] Loaded chat from sidebar: {chat_title}');
                }}
            """)
            log.info(f"Chat loaded from sidebar: {chat_title} (ID: {chat_id})")
    
    def _refresh_sidebar_chat_list(self, chat_list_json: str = None):
        """Refresh the sidebar chat list from chat data passed directly from JS."""
        log.info("[ChatList] _refresh_sidebar_chat_list called")
        try:
            if not chat_list_json:
                log.warning("[ChatList] No chat_list_json provided")
                return
            
            import json
            chat_list = json.loads(chat_list_json)
            log.info(f"[ChatList] Parsed {len(chat_list)} chats from JS data")
            
            # Clear existing items
            self._sidebar_chat_list.clear()
            
            # Add chats to sidebar with custom widgets
            for chat in chat_list:
                title = chat.get('title', 'New Chat')
                chat_id = chat.get('id', '')
                item, widget = self._create_chat_list_item(title, chat_id)
                self._sidebar_chat_list.addItem(item)
                self._sidebar_chat_list.setItemWidget(item, widget)
            
            log.info(f"[ChatList] Sidebar refreshed with {len(chat_list)} chats")
        except Exception as e:
            log.warning(f"[ChatList] Failed to refresh sidebar chat list: {e}")
            import traceback
            log.warning(traceback.format_exc())
    
    def _create_chat_list_item(self, title: str, chat_id: str):
        """Create a QListWidgetItem + custom widget with SVG pencil and trash icons."""
        from PyQt6.QtCore import QSize
        from src.utils.icons import make_icon
        
        item = QListWidgetItem()
        item.setData(Qt.ItemDataRole.UserRole, chat_id)
        item.setSizeHint(QSize(0, 34))
        
        # Container widget
        widget = QWidget()
        widget.setObjectName('chat_item_widget')
        widget.setStyleSheet("""
            QWidget#chat_item_widget { background: transparent; }
            QLabel { color: #a0a0a0; font-size: 13px; background: transparent; }
            QPushButton#pencil_btn { background: transparent; border: none; padding: 0px; }
            QPushButton#pencil_btn:hover { background: rgba(255, 255, 255, 0.08); border-radius: 3px; }
            QPushButton#trash_btn { background: transparent; border: none; padding: 0px; }
            QPushButton#trash_btn:hover { background: rgba(255, 80, 80, 0.15); border-radius: 3px; }
        """)
        
        row_layout = QHBoxLayout(widget)
        row_layout.setContentsMargins(8, 2, 4, 2)
        row_layout.setSpacing(4)
        
        # Title label (truncated)
        title_label = QLabel(title)
        title_label.setObjectName('chat_title_label')
        title_label.setMaximumWidth(160)
        metrics_text = title if len(title) <= 22 else title[:22] + '…'
        title_label.setText(metrics_text)
        title_label.setToolTip(title)
        row_layout.addWidget(title_label, 1)
        
        # Pencil (rename) button — clean SVG icon in gold
        pencil_btn = QPushButton()
        pencil_btn.setObjectName('pencil_btn')
        pencil_btn.setFixedSize(24, 24)
        pencil_btn.setIcon(make_icon('chat-pencil', '#e6a817', 16))
        pencil_btn.setIconSize(QSize(16, 16))
        pencil_btn.setToolTip('Rename')
        pencil_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        row_layout.addWidget(pencil_btn)
        
        # Trash (delete) button — clean SVG icon in red
        trash_btn = QPushButton()
        trash_btn.setObjectName('trash_btn')
        trash_btn.setFixedSize(24, 24)
        trash_btn.setIcon(make_icon('chat-trash', '#cc3333', 16))
        trash_btn.setIconSize(QSize(16, 16))
        trash_btn.setToolTip('Delete')
        trash_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        row_layout.addWidget(trash_btn)
        
        # Connect buttons — capture item reference
        def on_pencil(checked=False, _item=item):
            self._on_chat_rename_requested(_item)
        def on_trash(checked=False, _item=item):
            self._delete_chat_by_id(_item)
        
        pencil_btn.clicked.connect(on_pencil)
        trash_btn.clicked.connect(on_trash)
        
        return item, widget
    
    def _set_theme(self, theme: str):
        """Set the application theme."""
        self._settings.theme = theme
        self._apply_initial_theme()
    
    def _toggle_theme(self):
        """Toggle between dark and light theme"""
        current_theme = self._settings.theme
        new_theme = "light" if current_theme == "dark" else "dark"
        self._settings.theme = new_theme
        self._apply_initial_theme()
        log.info(f"Theme toggled to: {new_theme}")
    def _build_ui(self):
        """Build AI-First UI Layout - Codex-style with 2-panel and 4-panel states."""
        self.setWindowTitle("Cortex AI Agent")
        central = QWidget()
        self._central = central
        self.setCentralWidget(central)

        # === STATE MANAGEMENT ===
        # Always show 4-panel layout with chat ready
        self._is_welcome_state = False
        self._chat_started = True

        # === CODEX-STYLE LAYOUT WITH SPLITTERS ===
        main_splitter = QSplitter(Qt.Orientation.Horizontal)
        main_splitter.setChildrenCollapsible(True)
        main_splitter.setHandleWidth(4)
        main_splitter.setStyleSheet("""
            QSplitter::handle {
                background-color: #2a2a2a;
            }
            QSplitter::handle:hover {
                background-color: #4d78cc;
            }
        """)

        # Create all panels
        # Panel 1: Left Sidebar (220px) - Collapsible
        self._left_sidebar = self._create_left_sidebar()
        self._left_sidebar.setMinimumWidth(180)
        self._left_sidebar.setMaximumWidth(400)
        main_splitter.addWidget(self._left_sidebar)

        # Panel 2: Chat Panel - Main AI conversation (ALWAYS VISIBLE, flexible)
        self._chat_panel = self._create_chat_panel()
        self._chat_panel.setMinimumWidth(300)
        main_splitter.addWidget(self._chat_panel)

        # Panel 3: Review Panel (380px) - Summary/Review tabs - Collapsible
        self._review_panel = self._create_review_panel()
        self._review_panel.setMinimumWidth(250)
        self._review_panel.setMaximumWidth(600)
        main_splitter.addWidget(self._review_panel)

        # Panel 4: File Tree Panel (280px) - Changed files - Collapsible
        self._file_tree_panel = self._create_file_tree_panel()
        self._file_tree_panel.setMinimumWidth(200)
        self._file_tree_panel.setMaximumWidth(500)
        main_splitter.addWidget(self._file_tree_panel)

        # Set initial sizes (proportions)
        main_splitter.setSizes([220, 600, 380, 280])

        # Add splitter to main layout
        root_layout = QVBoxLayout(central)
        root_layout.setContentsMargins(0, 0, 0, 0)
        root_layout.setSpacing(0)
        root_layout.addWidget(main_splitter, 1)

        # Store reference for later access
        self._main_splitter = main_splitter

        # Panel toggle state tracking
        self._left_sidebar_hidden = False
        self._file_tree_hidden = False
        self._chat_panel_hidden = False
        self._review_panel_hidden = False
        self._summary_panel_hidden = False
        self._git_panel_hidden = False

        # Store minimum widths for panel toggle restore
        self._left_sidebar_min_width = 220
        self._chat_panel_min_width = 300
        self._review_panel_min_width = 380
        self._file_tree_min_width = 280

        # Keep old components for backward compatibility
        self._ai_splitter = None  # Replaced by 4-panel layout
        self._editor_tabs = EditorTabWidget()
        self._editor_tabs.setMinimumSize(200, 150)
        self._editor_tabs.hide()  # Hidden in Codex mode, shown in editor mode

        # Terminal tabs
        self._terminal_tabs = QTabWidget()
        self._terminal_tabs.setTabBar(CleanTabBar(self._terminal_tabs))
        self._terminal_tabs.setTabsClosable(True)
        self._terminal_tabs.setDocumentMode(True)
        self._terminal_tabs.setMovable(True)
        self._terminal_tabs.setVisible(False)
        self._terminal_tabs.setMinimumHeight(120)
        self._terminal_tabs.tabCloseRequested.connect(self._close_terminal_tab)

        # Find/Replace Dialog
        self._find_replace_dialog = FindReplaceDialog(self)
        self._find_replace_dialog.find_requested.connect(self._on_find_requested)
        self._find_replace_dialog.replace_requested.connect(self._on_replace_requested)
        self._find_replace_dialog.replace_all_requested.connect(self._on_replace_all_requested)

        # Command Palette - REMOVED (not implemented in AI-first mode)
        # self._command_palette = CommandPalette(self)
        # self._command_palette.command_selected.connect(self._on_command_selected)

        # === KEEP SIDEBAR & OLD COMPONENTS (Hidden but available) ===
        # These are kept for backward compatibility and can be shown via commands
        self._sidebar = SidebarWidget(self._file_manager, git_manager=self._git_manager)
        self._sidebar.setVisible(False)  # Hidden by default in AI-first mode
        self._sidebar.setMinimumWidth(44)
        self._sidebar.setMaximumWidth(700)

        # Initialize project info on welcome screen
        self._update_welcome_project_info()

        # Initialize Git summary after UI is built
        QTimer.singleShot(500, self._update_git_summary)
        
        # Auto-refresh git status every 30 seconds to detect push/commit from terminal
        self._git_refresh_timer = QTimer(self)
        self._git_refresh_timer.setInterval(30000)  # 30 seconds
        self._git_refresh_timer.timeout.connect(self._update_git_summary)
        self._git_refresh_timer.start()
        self._git_worker = None          # guard: only one worker at a time
        self._gh_version_cache = None    # cache gh --version across session

    def _show_welcome(self):
        """Show a VS Code-like welcome screen in the editor tabs."""
        from PyQt6.QtWidgets import QScrollArea
        
        # Remove existing welcome tab if present
        for i in range(self._editor_tabs.count()):
            tab_text = self._editor_tabs.tabText(i)
            if tab_text == "Welcome":
                widget = self._editor_tabs.widget(i)
                self._editor_tabs.removeTab(i)
                if widget:
                    widget.deleteLater()
                break
        
        self._welcome_scroll = QScrollArea()
        self._welcome_scroll.setWidgetResizable(True)
        self._welcome_scroll.setFrameShape(QFrame.Shape.NoFrame)
        self._welcome_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        
        self._welcome_widget = QWidget()
        welcome = self._welcome_widget
        wlay = QVBoxLayout(welcome)
        wlay.setAlignment(Qt.AlignmentFlag.AlignCenter)
        wlay.setSpacing(20)
        wlay.setContentsMargins(40, 40, 40, 40)

        # Logo and Title
        self._welcome_title = QLabel("⚡ Cortex AI Agent")
        self._welcome_title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self._welcome_title.setObjectName("welcome_title")
        wlay.addWidget(self._welcome_title)

        self._welcome_subtitle = QLabel("Your AI-powered development environment")
        self._welcome_subtitle.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self._welcome_subtitle.setObjectName("welcome_subtitle")
        wlay.addWidget(self._welcome_subtitle)

        # Dynamic Project Info
        self._welcome_project_info = QLabel()
        self._welcome_project_info.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self._welcome_project_info.setObjectName("welcome_project_info")
        wlay.addWidget(self._welcome_project_info)
        self._update_welcome_project_info()

        # Initialize hints list for theme styling
        self._welcome_hints = []


        # AI Chat Card (prominent, GitHub Copilot style) - SIMPLIFIED
        self._welcome_ai_card = QWidget()
        self._welcome_ai_card.setObjectName("welcome_ai_card")
        ai_card_layout = QVBoxLayout(self._welcome_ai_card)
        ai_card_layout.setContentsMargins(16, 16, 16, 16)
        ai_card_layout.setSpacing(12)
        
        ai_title = QLabel("Chat with Cortex")
        ai_title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        ai_title.setObjectName("ai_card_title")
        ai_card_layout.addWidget(ai_title)
        
        ai_subtitle = QLabel("Describe what you want to build or explore")
        ai_subtitle.setAlignment(Qt.AlignmentFlag.AlignCenter)
        ai_subtitle.setObjectName("ai_card_subtitle")
        ai_card_layout.addWidget(ai_subtitle)
        
        ai_prompt = ClickableLabel("<b>💬 Ask a question...</b>")
        ai_prompt.setAlignment(Qt.AlignmentFlag.AlignCenter)
        ai_prompt.setObjectName("ai_prompt")
        ai_prompt.clicked.connect(self._focus_ai_chat)
        ai_card_layout.addWidget(ai_prompt)
        self._welcome_hints.append(ai_prompt)
        
        wlay.addWidget(self._welcome_ai_card)

        # Add stretch to center content vertically
        wlay.addStretch()

        self._welcome_scroll.setWidget(self._welcome_widget)
        idx = self._editor_tabs.addTab(self._welcome_scroll, "Welcome")
        self._editor_tabs.setCurrentIndex(idx)
        # Apply background immediately so no white flash at startup
        self._apply_welcome_theme(self._theme_manager.is_dark)

    def _apply_welcome_theme(self, is_dark: bool):
        """Style the welcome widget to match the current theme."""
        if not hasattr(self, '_welcome_widget') or self._welcome_widget is None:
            return
        bg = "#1e1e1e" if is_dark else "#ffffff"
        fg = "#d4d4d4" if is_dark else "#1a1a1a"
        hint_fg = "#cccccc" if is_dark else "#333333"
        shortcut_color = "#858585" if is_dark else "#6c757d"
        sep_color = "#3e3e42" if is_dark else "#dee2e6"
        subtitle_color = "#858585" if is_dark else "#6c757d"
        
        # Extract conditional colors for AI card
        ai_card_bg = "#2d2d30" if is_dark else "#f3f3f3"
        ai_card_border = "#3e3e42" if is_dark else "#ddd"
        ai_prompt_bg = "#3e3e42" if is_dark else "#e8e8e8"
        ai_prompt_hover_bg = "#4e4e52" if is_dark else "#d4d4d4"
        hint_hover_bg = "#3e3e42" if is_dark else "#e9ecef"
        project_info_color = "#c678dd" if is_dark else "#9b30ff"

        # Base stylesheet with responsive sizing
        sw = self.width()
        # Scale fonts based on window width (min 800, max 1920)
        scale = min(max(sw / 1920.0, 0.6), 1.0)
        title_size = max(int(32 * scale), 18)
        subtitle_size = max(int(16 * scale), 12)
        hint_size = max(int(14 * scale), 11)
        project_size = max(int(14 * scale), 11)

        self._welcome_widget.setStyleSheet(
            f"background-color:{bg}; color:{fg};"
        )
        
        # Use object names for centralized styling
        styles = f"""
            QLabel#welcome_title {{
                font-size: {title_size}px;
                font-weight: bold;
                color: #007acc;
                margin-bottom: 8px;
            }}
            QLabel#welcome_subtitle {{
                font-size: {subtitle_size}px;
                color: {subtitle_color};
                margin-bottom: 12px;
            }}
            QFrame#welcome_sep {{
                color: {sep_color};
                margin: 10px 20%;
            }}
            QLabel#welcome_project_info {{
                font-size: {project_size}px;
                color: {project_info_color};
                margin: 8px 0 16px 0;
            }}
            QLabel#quick_actions_label,
            QLabel#recent_label,
            QLabel#help_label {{
                font-size: {hint_size + 1}px;
                font-weight: bold;
                color: {fg};
                margin-top: 16px;
                margin-bottom: 8px;
            }}
            ClickableLabel#welcome_hint {{
                font-size: {hint_size}px;
                color: {hint_fg};
                padding: 6px 12px;
                border-radius: 4px;
            }}
            ClickableLabel#welcome_hint:hover {{
                background-color: {hint_hover_bg};
            }}
            QLabel#recent_projects {{
                font-size: {hint_size - 1}px;
                color: {subtitle_color};
                padding: 4px 12px;
                font-style: italic;
            }}
            /* AI Chat Card - GitHub Copilot style */
            QWidget#welcome_ai_card {{
                background-color: {ai_card_bg};
                border: 1px solid {ai_card_border};
                border-radius: 8px;
                padding: 12px;
                margin: 8px 0;
            }}
            QLabel#ai_card_title {{
                font-size: 15px;
                font-weight: bold;
                color: {fg};
            }}
            QLabel#ai_card_subtitle {{
                font-size: 13px;
                color: {hint_fg};
            }}
            ClickableLabel#ai_prompt {{
                font-size: 14px;
                color: #569cd6;
                padding: 10px 14px;
                background-color: {ai_prompt_bg};
                border-radius: 6px;
            }}
            ClickableLabel#ai_prompt:hover {{
                background-color: {ai_prompt_hover_bg};
            }}
        """
        
        if hasattr(self, '_welcome_scroll') and self._welcome_scroll:
            self._welcome_scroll.setStyleSheet(f"background-color:{bg};")
        
        if hasattr(self, '_welcome_title'):
            self._welcome_title.setStyleSheet(f"font-size:{title_size}px; font-weight:bold; color:#007acc;")
        if hasattr(self, '_welcome_subtitle'):
            self._welcome_subtitle.setStyleSheet(f"font-size:{subtitle_size}px; color:{subtitle_color};")
        if hasattr(self, '_welcome_sep'):
            self._welcome_sep.setStyleSheet(f"color:{sep_color}; margin:10px 20%;")
        if hasattr(self, '_welcome_hints'):
            for row in self._welcome_hints:
                row.setStyleSheet(f"font-size:{hint_size}px; color:{hint_fg}; padding:6px 12px; border-radius:4px;")
        if hasattr(self, '_welcome_project_info') and self._welcome_project_info:
            self._welcome_project_info.setStyleSheet(f"font-size:{project_size}px; color:{'#c678dd' if is_dark else '#9b30ff'}; margin: 8px 0;")

    # ------------------------------------------------------------------
    # Codex-Style 4-Panel Layout Methods
    # ------------------------------------------------------------------

    def _create_left_sidebar(self) -> QWidget:
        """Create Left Sidebar (220px) - Navigation + Chat History."""
        sidebar = QWidget()
        sidebar.setObjectName("leftSidebar")
        layout = QVBoxLayout(sidebar)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # Codex-style colors
        bg_color = "#1a1a1a"
        text_color = "#cccccc"
        muted_color = "#888888"
        section_label_color = "#555555"

        sidebar.setStyleSheet(f"""
            QWidget#leftSidebar {{
                background-color: {bg_color};
            }}
            QLabel {{
                color: {text_color};
            }}
        """)

        # REMOVED: App Controls section (arrows, icon, title) and separator
        # Sidebar now starts directly with New chat button

        # Top Actions: New chat + Search
        actions_widget = QWidget()
        actions_layout = QVBoxLayout(actions_widget)
        actions_layout.setContentsMargins(12, 8, 12, 12)  # Reduced top margin from 12 to 8
        actions_layout.setSpacing(8)

        new_chat_btn = QPushButton("+ New chat")
        new_chat_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: #2d2d2d;
                color: {text_color};
                border: 1px solid #2a2a2a;
                border-radius: 6px;
                padding: 8px 12px;
                font-size: 13px;
                text-align: left;
            }}
            QPushButton:hover {{
                background-color: #3d3d3d;
            }}
        """)
        new_chat_btn.clicked.connect(self._on_new_chat)
        actions_layout.addWidget(new_chat_btn)

        # REMOVED: Search button (non-functional)

        layout.addWidget(actions_widget)

        # Section label style
        section_label_style = f"""
            color: {section_label_color};
            font-size: 11px;
            font-weight: 600;
            text-transform: uppercase;
            padding: 16px 12px 8px 12px;
        """

        # REMOVED: PLUGINS and AUTOMATIONS sections

        # Projects section
        projects_label = QLabel("Projects")
        projects_label.setStyleSheet(section_label_style)
        layout.addWidget(projects_label)

        # Current project name (dynamic)
        from pathlib import Path as _Path
        project_root = getattr(self._project_manager, 'root', None) if hasattr(self, '_project_manager') else None
        display_name = _Path(project_root).name if project_root else "No project"
        self._sidebar_project_item = QLabel(f"  📁 {display_name}")
        self._sidebar_project_item.setStyleSheet(f"color: {text_color}; font-size: 13px; padding: 6px 12px;")
        layout.addWidget(self._sidebar_project_item)

        # Chat List Section - Below Projects
        chats_header = QWidget()
        chats_layout = QHBoxLayout(chats_header)
        chats_layout.setContentsMargins(12, 12, 12, 8)
        chats_layout.setSpacing(0)

        chats_label = QLabel("Chats")
        chats_label.setStyleSheet(section_label_style.replace("padding: 16px 12px 8px 12px;", ""))
        chats_layout.addWidget(chats_label)
        chats_layout.addStretch()

        # REMOVED: New chat button from CHATS header (redundant with top "+ New chat" button)

        layout.addWidget(chats_header)

        # Chat list with editable titles and delete functionality
        self._sidebar_chat_list = QListWidget()
        self._sidebar_chat_list.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self._sidebar_chat_list.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self._sidebar_chat_list.setUniformItemSizes(True)
        self._sidebar_chat_list.setSizePolicy(
            QSizePolicy.Policy.Expanding,
            QSizePolicy.Policy.Expanding
        )
        self._sidebar_chat_list.setStyleSheet(f"""
            QListWidget {{
                background-color: transparent;
                border: none;
                color: {text_color};
                font-size: 13px;
                outline: none;
            }}
            QListWidget::item {{
                padding: 0px;
                border-radius: 4px;
                margin: 2px 8px;
            }}
            QListWidget::item:selected {{
                background-color: #2d2d2d;
            }}
            QListWidget::item:hover {{
                background-color: #252525;
            }}
        """)
        # Style scrollbar directly on the scrollbar widget — most reliable approach
        self._sidebar_chat_list.verticalScrollBar().setStyleSheet("""
            QScrollBar:vertical {
                background: #252525;
                width: 6px;
                border: none;
                margin: 0px;
            }
            QScrollBar::handle:vertical {
                background: #5a5a5a;
                min-height: 20px;
                border-radius: 3px;
            }
            QScrollBar::handle:vertical:hover {
                background: #6a6a6a;
            }
            QScrollBar::add-line:vertical,
            QScrollBar::sub-line:vertical {
                height: 0px;
                background: none;
            }
            QScrollBar::add-page:vertical,
            QScrollBar::sub-page:vertical {
                background: none;
            }
        """)
        
        # Double-click to rename, right-click for context menu, click to load chat
        self._sidebar_chat_list.itemDoubleClicked.connect(self._on_chat_rename_requested)
        self._sidebar_chat_list.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self._sidebar_chat_list.customContextMenuRequested.connect(self._on_chat_context_menu)
        self._sidebar_chat_list.itemClicked.connect(self._on_chat_item_clicked)
        
        layout.addWidget(self._sidebar_chat_list, 1)  # stretch=1 so it fills remaining space

        # Bottom: Settings + Upgrade
        bottom_widget = QWidget()
        bottom_layout = QHBoxLayout(bottom_widget)
        bottom_layout.setContentsMargins(12, 12, 12, 12)
        bottom_layout.setSpacing(8)

        settings_btn = QPushButton("⚙ Settings")
        settings_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: transparent;
                color: {muted_color};
                border: none;
                padding: 6px;
                font-size: 13px;
            }}
            QPushButton:hover {{
                color: {text_color};
            }}
        """)
        settings_btn.clicked.connect(self._on_settings_requested)
        bottom_layout.addWidget(settings_btn)
        bottom_layout.addStretch()

        upgrade_btn = QPushButton("Upgrade")
        upgrade_btn.setStyleSheet("""
            QPushButton {
                background-color: #f0a500;
                color: #1a1a1a;
                border: none;
                border-radius: 6px;
                padding: 6px 12px;
                font-size: 12px;
                font-weight: 600;
            }
            QPushButton:hover {
                background-color: #ffb020;
            }
        """)
        bottom_layout.addWidget(upgrade_btn)

        layout.addWidget(bottom_widget)

        return sidebar

    def _create_chat_panel(self) -> QWidget:
        """Create Chat Panel (flexible) - Main AI conversation."""
        panel = QWidget()
        panel.setObjectName("chatPanel")
        layout = QVBoxLayout(panel)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        bg_color = "#1e1e1e"
        border_color = "#2a2a2a"

        panel.setStyleSheet(f"""
            QWidget#chatPanel {{
                background-color: {bg_color};
            }}
        """)

        # Top bar - REMOVED hardcoded header (hi Cortex + Commit button)
        # Clean minimal design without hardcoded elements
        
        # AI Chat widget (existing)
        self._ai_chat = AIChatWidget()
        self._ai_chat.run_command.connect(self._on_ai_run_command)
        self._ai_chat.stop_requested.connect(self._on_ai_stop_requested)
        self._ai_chat.open_file_requested.connect(self._open_file)
        self._ai_chat.accept_file_edit_requested.connect(self._on_accept_file_edit)
        self._ai_chat.reject_file_edit_requested.connect(self._on_reject_file_edit)
        self._ai_chat.open_terminal_requested.connect(self._show_terminal_panel)
        self._ai_chat.run_in_terminal_requested.connect(self._show_terminal_and_run)
        self._ai_chat.set_code_context_callback(self._get_code_context)
        self._ai_chat.toggle_autogen_requested.connect(self._on_toggle_autogen)
        
        # Connect chat list update signal to refresh sidebar
        self._ai_chat.chat_list_updated_with_data.connect(self._refresh_sidebar_chat_list)
        log.info("[ChatList] Connected chat_list_updated_with_data signal to _refresh_sidebar_chat_list")
        layout.addWidget(self._ai_chat, 1)

        return panel

    def _create_review_panel(self) -> QWidget:
        """Create Review Panel (380px) - Clickable Summary/Review tabs."""
        panel = QWidget()
        panel.setObjectName("reviewPanel")
        layout = QVBoxLayout(panel)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        bg_color = "#1e1e1e"

        panel.setStyleSheet(f"""
            QWidget#reviewPanel {{
                background-color: {bg_color};
            }}
        """)

        # Tab bar with clickable buttons
        tab_bar = QWidget()
        tab_bar.setFixedHeight(48)
        tab_layout = QHBoxLayout(tab_bar)
        tab_layout.setContentsMargins(12, 0, 12, 0)
        tab_layout.setSpacing(8)

        # Summary tab button
        self._summary_tab_btn = QPushButton("≡ Summary")
        self._summary_tab_btn.setFlat(True)
        self._summary_tab_btn.setStyleSheet("""
            QPushButton {
                color: #ffffff;
                font-size: 13px;
                font-weight: 500;
                background-color: transparent;
                border: none;
                padding: 8px 12px;
                border-radius: 6px;
            }
            QPushButton:hover {
                background-color: #2d2d2d;
            }
            QPushButton:checked {
                background-color: #2d2d2d;
            }
        """)
        self._summary_tab_btn.setCheckable(True)
        self._summary_tab_btn.setChecked(True)
        self._summary_tab_btn.clicked.connect(self._on_summary_tab_clicked)
        tab_layout.addWidget(self._summary_tab_btn)

        # Review tab button
        self._review_tab_btn = QPushButton("📄 Review")
        self._review_tab_btn.setFlat(True)
        self._review_tab_btn.setStyleSheet("""
            QPushButton {
                color: #888888;
                font-size: 13px;
                background-color: transparent;
                border: none;
                padding: 8px 12px;
                border-radius: 6px;
            }
            QPushButton:hover {
                background-color: #2d2d2d;
                color: #cccccc;
            }
            QPushButton:checked {
                background-color: #2d2d2d;
                color: #ffffff;
            }
        """)
        self._review_tab_btn.setCheckable(True)
        self._review_tab_btn.clicked.connect(self._on_review_tab_clicked)
        tab_layout.addWidget(self._review_tab_btn)

        tab_layout.addStretch()

        add_tab = QLabel("+")
        add_tab.setStyleSheet("color: #888888; font-size: 16px;")
        tab_layout.addWidget(add_tab)

        layout.addWidget(tab_bar)

        # Separator
        sep = QFrame()
        sep.setFrameShape(QFrame.Shape.HLine)
        sep.setStyleSheet("background-color: #2a2a2a; max-height: 1px;")
        layout.addWidget(sep)

        # Stacked widget for tab content
        self._review_stack = QStackedWidget()

        # === SUMMARY TAB CONTENT ===
        summary_content = QWidget()
        summary_layout = QVBoxLayout(summary_content)
        summary_layout.setContentsMargins(16, 16, 16, 16)
        summary_layout.setSpacing(12)

        # Progress section
        progress_label = QLabel("Progress")
        progress_label.setStyleSheet("color: #aaaaaa; font-size: 12px; font-weight: bold;")
        summary_layout.addWidget(progress_label)

        progress_text = QLabel("Progress displayed for longer responses")
        progress_text.setStyleSheet("color: #666666; font-size: 13px; padding: 8px 0;")
        summary_layout.addWidget(progress_text)

        # Branch details
        branch_label = QLabel("Branch details")
        branch_label.setStyleSheet("color: #aaaaaa; font-size: 12px; font-weight: bold; margin-top: 8px;")
        summary_layout.addWidget(branch_label)

        # Current branch name (will be updated dynamically)
        self._branch_name_label = QLabel("🌿 No repository")
        self._branch_name_label.setStyleSheet("color: #888888; font-size: 13px; padding: 4px 0;")
        summary_layout.addWidget(self._branch_name_label)

        # GitHub CLI status (will be updated dynamically)
        self._github_status_label = QLabel("🐙 GitHub CLI not installed")
        self._github_status_label.setStyleSheet("color: #888888; font-size: 13px; padding: 4px 0;")
        summary_layout.addWidget(self._github_status_label)

        # Changes count (will be updated dynamically)
        self._changes_label = QLabel("✏️ 0 unstaged, 0 untracked")
        self._changes_label.setStyleSheet("color: #cccccc; font-size: 13px; padding: 4px 0;")
        summary_layout.addWidget(self._changes_label)

        # Artifacts
        artifacts_label = QLabel("Artifacts")
        artifacts_label.setStyleSheet("color: #aaaaaa; font-size: 12px; font-weight: bold; margin-top: 16px;")
        summary_layout.addWidget(artifacts_label)

        artifacts_text = QLabel("View and open referenced files")
        artifacts_text.setStyleSheet("color: #666666; font-size: 13px; padding: 8px 0;")
        summary_layout.addWidget(artifacts_text)

        # Sources
        sources_label = QLabel("Sources")
        sources_label.setStyleSheet("color: #aaaaaa; font-size: 12px; font-weight: bold; margin-top: 16px;")
        summary_layout.addWidget(sources_label)

        sources_text = QLabel("Track sources used")
        sources_text.setStyleSheet("color: #666666; font-size: 13px; padding: 8px 0;")
        summary_layout.addWidget(sources_text)

        summary_layout.addStretch()
        self._review_stack.addWidget(summary_content)

        # === REVIEW TAB CONTENT ===
        review_content = QWidget()
        review_layout = QVBoxLayout(review_content)
        review_layout.setContentsMargins(16, 16, 16, 16)
        review_layout.setSpacing(12)

        # Unstaged files header
        unstaged_header = QWidget()
        unstaged_layout = QHBoxLayout(unstaged_header)
        unstaged_layout.setContentsMargins(0, 0, 0, 0)

        unstaged_label = QLabel("Unstaged")
        unstaged_label.setStyleSheet("color: #ffffff; font-size: 13px; font-weight: 500;")
        unstaged_layout.addWidget(unstaged_label)

        # Unstaged count (will be updated dynamically)
        self._unstaged_count_label = QLabel("0 ▾")
        self._unstaged_count_label.setStyleSheet("color: #888888; font-size: 13px;")
        unstaged_layout.addWidget(self._unstaged_count_label)
        unstaged_layout.addStretch()

        more_options = QLabel("⋯")
        more_options.setStyleSheet("color: #888888; font-size: 16px;")
        unstaged_layout.addWidget(more_options)

        review_layout.addWidget(unstaged_header)

        # Large diff notice
        self._diff_notice = QLabel("Large diff detected — showing one file at a time.")
        self._diff_notice.setStyleSheet("color: #888888; font-size: 12px; padding: 8px 0;")
        self._diff_notice.hide()  # Hidden by default
        review_layout.addWidget(self._diff_notice)

        # File list container with scrollbar
        self._file_list_scroll = QScrollArea()
        self._file_list_scroll.setWidgetResizable(True)
        self._file_list_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self._file_list_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self._file_list_scroll.setStyleSheet("""
            QScrollArea {
                border: none;
                background-color: transparent;
            }
            QScrollBar:vertical {
                border: none;
                background: #252525;
                width: 5px;
                margin: 0px;
                border-radius: 6px;
            }
            QScrollBar::handle:vertical {
                background: #5a5a5a;
                min-height: 30px;
                border-radius: 6px;
            }
            QScrollBar::handle:vertical:hover {
                background: #6a6a6a;
            }
            QScrollBar::add-line:vertical,
            QScrollBar::sub-line:vertical {
                border: none;
                background: none;
                height: 0px;
            }
            QScrollBar::add-page:vertical,
            QScrollBar::sub-page:vertical {
                background: none;
            }
        """)
        
        self._file_list_container = QWidget()
        self._file_list_layout = QVBoxLayout(self._file_list_container)
        self._file_list_layout.setContentsMargins(0, 0, 0, 0)
        self._file_list_layout.setSpacing(0)  # No spacing between files
        
        self._file_list_scroll.setWidget(self._file_list_container)
        review_layout.addWidget(self._file_list_scroll, 1)
        # NOTE: No addStretch() here — scroll area must fill all remaining space

        self._review_stack.addWidget(review_content)
        self._review_stack.setCurrentIndex(0)  # Show Summary by default

        layout.addWidget(self._review_stack, 1)

        return panel

    def _on_summary_tab_clicked(self):
        """Handle Summary tab click."""
        self._summary_tab_btn.setChecked(True)
        self._review_tab_btn.setChecked(False)
        self._review_stack.setCurrentIndex(0)

    def _on_review_tab_clicked(self):
        """Handle Review tab click."""
        self._review_tab_btn.setChecked(True)
        self._summary_tab_btn.setChecked(False)
        self._review_stack.setCurrentIndex(1)

    def _update_git_summary(self):
        """Kick off a background thread to collect git status (non-blocking)."""
        if not hasattr(self, '_git_manager'):
            log.warning("[GIT] GitManager not available")
            self._set_no_git_status()
            return

        if not self._git_manager.is_repo():
            log.info("[GIT] No repository set")
            self._set_no_git_status()
            return

        # Guard: skip if a worker is already running
        if self._git_worker is not None and self._git_worker.isRunning():
            return

        log.info("[GIT] Updating git summary (background)...")
        self._git_worker = _GitStatusWorker(
            self._git_manager._repo_path,
            self._git_manager,
            gh_cached=self._gh_version_cache,
        )
        self._git_worker.status_ready.connect(self._on_git_status_ready)
        self._git_worker.start()

    @pyqtSlot(dict)
    def _on_git_status_ready(self, data: dict):
        """Receive git data from background worker and update UI labels."""
        # Branch
        branch = data.get('branch')
        if branch:
            self._branch_name_label.setText(f"🌿 {branch}")
            self._branch_name_label.setStyleSheet("color: #4ec94e; font-size: 13px; padding: 4px 0;")
        else:
            self._branch_name_label.setText("🌿 Unknown branch")
            self._branch_name_label.setStyleSheet("color: #888888; font-size: 13px; padding: 4px 0;")

        # GitHub CLI (cache the result for the session)
        gh = data.get('gh', '')
        self._gh_version_cache = gh
        if gh:
            self._github_status_label.setText(f"🐙 {gh}")
            self._github_status_label.setStyleSheet("color: #4ec94e; font-size: 13px; padding: 8px 0;")
        else:
            self._github_status_label.setText("🐙 GitHub CLI not installed")
            self._github_status_label.setStyleSheet("color: #888888; font-size: 13px; padding: 8px 0;")

        # Changes summary
        total = data.get('total', 0)
        unstaged = data.get('unstaged', 0)
        untracked = data.get('untracked', 0)
        staged = data.get('staged', 0)
        if total > 0:
            self._changes_label.setText(f"✏️ {unstaged} unstaged, {untracked} untracked, {staged} staged")
            self._changes_label.setStyleSheet("color: #e6a817; font-size: 13px; padding: 4px 0;")
        else:
            self._changes_label.setText("✏️ No changes")
            self._changes_label.setStyleSheet("color: #4ec94e; font-size: 13px; padding: 4px 0;")

        if hasattr(self, '_unstaged_count_label'):
            self._unstaged_count_label.setText(f"{unstaged} ▾")

        # Update review file list from pre-computed stats
        self._update_review_file_list_from_stats(data.get('file_stats', []), data.get('git_files', []))
    
    def _set_no_git_status(self):
        """Set all git status labels to 'no repository' state."""
        if hasattr(self, '_branch_name_label'):
            self._branch_name_label.setText("🌿 No repository")
            self._branch_name_label.setStyleSheet("color: #888888; font-size: 13px; padding: 4px 0;")
        if hasattr(self, '_github_status_label'):
            self._github_status_label.setText("🐙 GitHub CLI not installed")
            self._github_status_label.setStyleSheet("color: #888888; font-size: 13px; padding: 4px 0;")
        if hasattr(self, '_changes_label'):
            self._changes_label.setText("✏️ No changes")
            self._changes_label.setStyleSheet("color: #888888; font-size: 13px; padding: 4px 0;")
        if hasattr(self, '_unstaged_count_label'):
            self._unstaged_count_label.setText("0 ▾")

    def _check_github_cli(self):
        """Legacy stub — gh check now runs inside _GitStatusWorker."""
        pass

    def _update_review_file_list(self, git_files):
        """Legacy entry — delegates to stats-based variant."""
        # Kept for any external callers; runs with empty stats (no diff numbers)
        stats = [(gf, 0, 0) for gf in git_files[:100]]
        self._update_review_file_list_from_stats(stats, git_files)

    def _update_review_file_list_from_stats(self, file_stats, git_files):
        """Update Review tab using pre-computed diff stats (no subprocess on UI thread)."""
        # Clear existing file list
        while self._file_list_layout.count():
            child = self._file_list_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        if not git_files:
            no_changes = QLabel("No changed files")
            no_changes.setStyleSheet("color: #888888; font-size: 13px; padding: 8px 0;")
            self._file_list_layout.addWidget(no_changes)
            self._file_list_layout.addStretch()
            self._diff_notice.hide()
            return

        if len(git_files) > 50:
            self._diff_notice.show()
        else:
            self._diff_notice.hide()

        added_count = 0
        for gf, additions, deletions in file_stats:
            if additions == 0 and deletions == 0:
                continue
            file_widget = self._create_file_diff_item_from_stats(gf, additions, deletions)
            if file_widget:
                self._file_list_layout.addWidget(file_widget)
                added_count += 1

        if added_count == 0:
            no_changes = QLabel("No files with changes")
            no_changes.setStyleSheet("color: #888888; font-size: 13px; padding: 8px 0;")
            self._file_list_layout.addWidget(no_changes)

        self._file_list_layout.addStretch()

    def _create_file_diff_item(self, git_file) -> QWidget:
        """Legacy wrapper — gets diff stats inline (kept for compatibility)."""
        additions, deletions = self._get_file_diff_stats(git_file.path)
        return self._create_file_diff_item_from_stats(git_file, additions, deletions)

    def _create_file_diff_item_from_stats(self, git_file, additions: int, deletions: int) -> QWidget:
        """Create a file diff item widget from pre-computed stats."""
        if additions == 0 and deletions == 0:
            return None

        file_item = QWidget()
        file_item.setFixedHeight(28)
        file_layout = QHBoxLayout(file_item)
        file_layout.setContentsMargins(0, 2, 0, 2)
        file_layout.setSpacing(8)

        from pathlib import Path
        filename = Path(git_file.path).name
        file_name = QLabel(filename)
        file_name.setStyleSheet("color: #cccccc; font-size: 13px;")
        file_layout.addWidget(file_name)

        file_layout.addStretch()

        if additions > 0:
            additions_label = QLabel(f"+{additions}")
            additions_label.setStyleSheet("color: #4ec94e; font-size: 13px; font-weight: 500;")
            file_layout.addWidget(additions_label)

        if deletions > 0:
            deletions_label = QLabel(f"-{deletions}")
            deletions_label.setStyleSheet("color: #e05252; font-size: 13px; font-weight: 500;")
            file_layout.addWidget(deletions_label)

        expand_icon = QLabel("▾")
        expand_icon.setStyleSheet("color: #888888; font-size: 12px;")
        file_layout.addWidget(expand_icon)

        return file_item

    def _get_file_diff_stats(self, file_path: str) -> tuple:
        """Get additions and deletions for a file (checks both staged and unstaged)."""
        if not hasattr(self, '_git_manager') or not self._git_manager.is_repo():
            return 0, 0

        import subprocess
        cwd = self._git_manager._repo_path
        total_add, total_del = 0, 0

        def _parse_numstat(output: str):
            for line in output.strip().splitlines():
                parts = line.split('\t')
                if len(parts) >= 2:
                    try:
                        a = int(parts[0]) if parts[0] != '-' else 0
                        d = int(parts[1]) if parts[1] != '-' else 0
                        return a, d
                    except ValueError:
                        pass
            return 0, 0

        try:
            # Unstaged changes (working tree vs index)
            # FIX: Prevent console window popup
            kwargs = dict(cwd=cwd, capture_output=True, text=True, timeout=5)
            if sys.platform == 'win32':
                kwargs['creationflags'] = subprocess.CREATE_NO_WINDOW
            
            r = subprocess.run(
                ["git", "diff", "--numstat", file_path],
                **kwargs
            )
            if r.returncode == 0 and r.stdout.strip():
                a, d = _parse_numstat(r.stdout)
                total_add += a
                total_del += d
        except Exception:
            pass

        try:
            # Staged changes (index vs HEAD)
            # FIX: Prevent console window popup
            kwargs = dict(cwd=cwd, capture_output=True, text=True, timeout=5)
            if sys.platform == 'win32':
                kwargs['creationflags'] = subprocess.CREATE_NO_WINDOW
            
            r = subprocess.run(
                ["git", "diff", "--cached", "--numstat", file_path],
                **kwargs
            )
            if r.returncode == 0 and r.stdout.strip():
                a, d = _parse_numstat(r.stdout)
                total_add += a
                total_del += d
        except Exception:
            pass

        return total_add, total_del

    def _create_file_tree_panel(self) -> QWidget:
        """Create File Tree Panel (280px) - Project Explorer."""
        panel = QWidget()
        panel.setObjectName("fileTreePanel")
        layout = QVBoxLayout(panel)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        bg_color = "#1a1a1a"

        panel.setStyleSheet(f"""
            QWidget#fileTreePanel {{
                background-color: {bg_color};
            }}
        """)

        # Header
        header = QWidget()
        header.setFixedHeight(48)
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(12, 0, 12, 0)
        header_layout.setSpacing(8)

        title = QLabel("📁 Explore")
        title.setStyleSheet("color: #ffffff; font-size: 13px; font-weight: 500;")
        header_layout.addWidget(title)
        header_layout.addStretch()

        layout.addWidget(header)

        # Separator
        sep = QFrame()
        sep.setFrameShape(QFrame.Shape.HLine)
        sep.setStyleSheet("background-color: #2a2a2a; max-height: 1px;")
        layout.addWidget(sep)

        # Filter input
        filter_widget = QWidget()
        filter_layout = QHBoxLayout(filter_widget)
        filter_layout.setContentsMargins(12, 12, 12, 8)

        self._file_filter_input = QLineEdit()
        self._file_filter_input.setPlaceholderText("🔍 Filter files...")
        self._file_filter_input.setStyleSheet("""
            QLineEdit {
                background-color: #252525;
                color: #cccccc;
                border: none;
                border-radius: 6px;
                padding: 8px 12px;
                font-size: 13px;
            }
            QLineEdit:focus {
                background-color: #2d2d2d;
            }
        """)
        self._file_filter_input.textChanged.connect(self._filter_file_tree)
        filter_layout.addWidget(self._file_filter_input)
        layout.addWidget(filter_widget)

        # File Tree View with scrollbar
        self._file_tree_scroll = QScrollArea()
        self._file_tree_scroll.setWidgetResizable(True)
        self._file_tree_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self._file_tree_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self._file_tree_scroll.setStyleSheet("""
            QScrollArea {
                border: none;
                background-color: transparent;
            }
            QScrollBar:vertical {
                border: none;
                background: #252525;
                width: 5px;
                margin: 0px;
                border-radius: 6px;
            }
            QScrollBar::handle:vertical {
                background: #5a5a5a;
                min-height: 30px;
                border-radius: 6px;
            }
            QScrollBar::handle:vertical:hover {
                background: #6a6a6a;
            }
            QScrollBar::add-line:vertical,
            QScrollBar::sub-line:vertical {
                border: none;
                background: none;
                height: 0px;
            }
            QScrollBar::add-page:vertical,
            QScrollBar::sub-page:vertical {
                background: none;
            }
        """)

        # Create tree widget
        self._file_tree = QTreeView()
        self._file_tree.setHeaderHidden(True)
        self._file_tree.setIndentation(14)
        self._file_tree.setAnimated(True)
        self._file_tree.setExpandsOnDoubleClick(False)  # Single click expands folders
        self._file_tree.setSelectionMode(QTreeView.SelectionMode.ExtendedSelection)
        
        # Enable F2 rename (but we'll handle it manually to prevent conflicts)
        from PyQt6.QtWidgets import QAbstractItemView
        self._file_tree.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        
        # VS Code-style QSS from sidebar.py
        self._file_tree.setStyleSheet("""
            QTreeView {
                background: #1e1e1e;
                border: none;
                outline: 0;
                font-size: 12px;
                color: #cccccc;
            }
            QTreeView::item {
                height: 24px;
                border-radius: 3px;
                padding-left: 2px;
            }
            QTreeView::item:hover {
                background: #37373d;
                color: #ffffff;
            }
            QTreeView::item:selected {
                background: #094771;
                color: #ffffff;
                border: 1px solid #007acc;
            }
            QTreeView::branch {
                background: #1e1e1e;
            }
            QTreeView::branch:has-children:!has-siblings:closed,
            QTreeView::branch:closed:has-children:has-siblings {
                image: none;
                border-image: none;
            }
            QTreeView::branch:open:has-children:!has-siblings,
            QTreeView::branch:open:has-children:has-siblings {
                image: none;
                border-image: none;
            }
        """)

        # Setup file system model (from sidebar.py)
        self._file_model = QFileSystemModel()
        self._file_model.setReadOnly(False)
        self._file_model.setFilter(QDir.Filter.AllEntries | QDir.Filter.NoDotAndDotDot)
        self._file_model.setNameFilterDisables(False)
        self._file_tree.setModel(self._file_model)
        
        # Set custom delegate for chevron arrows (from sidebar.py)
        self._file_delegate = FileTreeDelegate(self._file_model)
        self._file_tree.setItemDelegate(self._file_delegate)
        
        # Hide columns except name
        for col in (1, 2, 3):
            self._file_tree.setColumnHidden(col, True)

        # Connect signals for tree updates (debounced viewport repaints)
        self._file_tree_update_timer = QTimer(self)
        self._file_tree_update_timer.setSingleShot(True)
        self._file_tree_update_timer.setInterval(50)
        self._file_tree_update_timer.timeout.connect(lambda: self._file_tree.viewport().update())
        self._file_model.directoryLoaded.connect(lambda _: self._file_tree_update_timer.start())
        self._file_tree.expanded.connect(self._on_tree_expanded)
        self._file_tree.collapsed.connect(self._on_tree_collapsed)
        
        # Enable single-click folder expansion
        self._file_tree.clicked.connect(self._on_tree_clicked)
        
        # Enable F2 key for rename
        self._file_tree.installEventFilter(self)
        
        # Connect double-click to open file
        self._file_tree.doubleClicked.connect(self._on_file_tree_double_clicked)
        
        # Override keyPressEvent for F2 rename
        original_key_press = self._file_tree.keyPressEvent
        def custom_key_press(event):
            if event.key() == Qt.Key.Key_F2:
                self._rename_file()
                return
            original_key_press(event)
        self._file_tree.keyPressEvent = custom_key_press

        self._file_tree_scroll.setWidget(self._file_tree)
        layout.addWidget(self._file_tree_scroll, 1)

        return panel

    def _filter_file_tree(self, text: str):
        """Filter files in the tree view."""
        if not text:
            self._file_model.setNameFilters([])
        else:
            # Use wildcard pattern for filtering
            self._file_model.setNameFilters([f"*{text}*"])

    def _on_file_tree_double_clicked(self, index: QModelIndex):
        """Handle file double-click to open in editor."""
        file_path = self._file_model.filePath(index)
        if Path(file_path).is_file():
            self._open_file(file_path)

    def _on_tree_expanded(self, index: QModelIndex):
        """Handle tree item expansion."""
        self._file_tree_update_timer.start()

    def _on_tree_collapsed(self, index: QModelIndex):
        """Handle tree item collapse."""
        self._file_tree_update_timer.start()

    def _on_tree_clicked(self, index: QModelIndex):
        """Handle single-click to toggle folder expansion."""
        file_path = self._file_model.filePath(index)
        if Path(file_path).is_dir():
            # Toggle expand/collapse on single click
            if self._file_tree.isExpanded(index):
                self._file_tree.collapse(index)
            else:
                self._file_tree.expand(index)

    def eventFilter(self, obj, event):
        """Handle F2 key press for renaming files/folders in Explore tree."""
        from PyQt6.QtCore import QEvent
        
        if hasattr(self, '_file_tree') and obj == self._file_tree.viewport():
            if event.type() == QEvent.Type.KeyPress:
                if event.key() == Qt.Key.Key_F2:
                    self._rename_file()
                    return True
        
        return super().eventFilter(obj, event)

    def _set_project_root(self, path: str):
        """Set the project root directory for the file tree."""
        import os
        # Normalize path (from sidebar.py line 792-807)
        normalized_path = os.path.normpath(os.path.abspath(path))
        log.info(f"[FILE TREE] Setting project root: {normalized_path}")
        
        if hasattr(self, '_file_model'):
            # Set model root and get root index
            root_idx = self._file_model.setRootPath(normalized_path)
            log.info(f"[FILE TREE] Model root index: {root_idx.row()}, {root_idx.column()}")
            log.info(f"[FILE TREE] Model root path: {self._file_model.filePath(root_idx)}")
            
            if hasattr(self, '_file_tree'):
                self._file_tree.setRootIndex(root_idx)
                self._file_tree.setVisible(True)
                log.info(f"[FILE TREE] Tree root set successfully")


    # ------------------------------------------------------------------
    # Menu
    # ------------------------------------------------------------------
    def _build_menu(self):
        mb = self.menuBar()

        # File
        file_menu = mb.addMenu("File")
        self._add_action(file_menu, "New Window", self._new_window, "Ctrl+Shift+N")
        self._add_action(file_menu, "New Chat", self._on_new_chat, "Ctrl+N")
        self._add_action(file_menu, "Quick Chat", self._quick_chat, "Alt+Ctrl+N")
        file_menu.addSeparator()
        self._add_action(file_menu, "Open File...", self._open_file_dialog, "Ctrl+Shift+O")
        self._add_action(file_menu, "Open Folder...", self._open_folder_dialog, "Ctrl+O")
        file_menu.addSeparator()
        self._add_action(file_menu, "Save", self._save_current, "Ctrl+S")
        self._add_action(file_menu, "Save All", self._save_all, "Ctrl+Shift+S")
        file_menu.addSeparator()
        self._add_action(file_menu, "Settings...", self._open_settings, "Ctrl+,")
        file_menu.addSeparator()
        self._add_action(file_menu, "Exit", self.close, "Alt+F4")

        # Edit
        edit_menu = mb.addMenu("Edit")
        self._add_action(edit_menu, "Undo", self._undo, "Ctrl+Z")
        self._add_action(edit_menu, "Redo", self._redo, "Ctrl+Y")
        edit_menu.addSeparator()
        self._add_action(edit_menu, "Cut", lambda: self._current_editor_action("cut"), "Ctrl+X")
        self._add_action(edit_menu, "Copy", lambda: self._current_editor_action("copy"), "Ctrl+C")
        self._add_action(edit_menu, "Paste", lambda: self._current_editor_action("paste"), "Ctrl+V")
        edit_menu.addSeparator()
        self._add_action(edit_menu, "Select All", lambda: self._current_editor_action("selectAll"), "Ctrl+A")
        edit_menu.addSeparator()
        self._add_action(edit_menu, "Find...", self._show_find, "Ctrl+F")
        self._add_action(edit_menu, "Find and Replace...", self._show_find_replace, "Ctrl+H")
        self._add_action(edit_menu, "Find in Files...", self._find_in_files, "Ctrl+Shift+F")
        edit_menu.addSeparator()
        self._add_action(edit_menu, "Rename...", self._rename_file, "F2")
        edit_menu.addSeparator()
        self._add_action(edit_menu, "Go to Line...", self._go_to_line, "Ctrl+G")
        edit_menu.addSeparator()
        self._add_action(edit_menu, "Toggle Comment", self._toggle_comment, "Ctrl+/")
        self._add_action(edit_menu, "Delete Line", self._delete_line, "Ctrl+Shift+K")
        self._add_action(edit_menu, "Duplicate Line", self._duplicate_line, "Ctrl+Shift+D")
        edit_menu.addSeparator()
        self._add_action(edit_menu, "Indent", self._indent_line, "Ctrl+]")
        self._add_action(edit_menu, "Outdent", self._outdent_line, "Ctrl+[")
        edit_menu.addSeparator()
        self._add_action(edit_menu, "Move Line Up", self._move_line_up, "Alt+Up")
        self._add_action(edit_menu, "Move Line Down", self._move_line_down, "Alt+Down")

        # Navigation
        nav_menu = mb.addMenu("Navigation")
        self._add_action(nav_menu, "Quick Open File...", self._quick_open, "Ctrl+P")
        self._add_action(nav_menu, "Go to Symbol...", self._go_to_symbol, "Ctrl+Shift+O")
        nav_menu.addSeparator()
        self._add_action(nav_menu, "Close Tab", self._close_current_tab, "Ctrl+W")
        self._add_action(nav_menu, "Close All Tabs", self._close_all_tabs, "Ctrl+Shift+W")
        nav_menu.addSeparator()
        self._add_action(nav_menu, "Next Tab", self._next_tab, "Ctrl+Tab")
        self._add_action(nav_menu, "Previous Tab", self._prev_tab, "Ctrl+Shift+Tab")
        nav_menu.addSeparator()
        self._add_action(nav_menu, "Previous Chat", self._previous_chat, "Ctrl+Shift+[")
        self._add_action(nav_menu, "Next Chat", self._next_chat, "Ctrl+Shift+]")
        self._add_action(nav_menu, "Back", self._navigate_back, "Ctrl+[")
        self._add_action(nav_menu, "Forward", self._navigate_forward, "Ctrl+]")
        nav_menu.addSeparator()
        self._add_action(nav_menu, "Keyboard Shortcuts Help", self._show_shortcuts_help, "Ctrl+Alt+K")

        # View
        view_menu = mb.addMenu("View")
        self._add_action(view_menu, "Toggle Theme", self._toggle_theme, "Ctrl+Shift+T")
        self._add_action(view_menu, "Toggle Terminal", self._toggle_terminal, "Ctrl+J")
        view_menu.addSeparator()
        self._add_action(view_menu, "Toggle Sidebar", self._toggle_sidebar, "Ctrl+B")
        self._add_action(view_menu, "Toggle File Tree", self._toggle_file_tree, "Ctrl+Shift+I")
        self._add_action(view_menu, "Toggle Review Panel", self._toggle_review_panel_menu, "Alt+Ctrl+B")
        view_menu.addSeparator()
        self._add_action(view_menu, "Zoom In", self._zoom_in, "Ctrl++")
        self._add_action(view_menu, "Zoom Out", self._zoom_out, "Ctrl+-")
        self._add_action(view_menu, "Reset Zoom", self._zoom_reset, "Ctrl+0")
        view_menu.addSeparator()
        self._add_action(view_menu, "Toggle Full Screen", self._toggle_fullscreen, "F11")

        # AI
        ai_menu = mb.addMenu("AI")
        self._add_action(ai_menu, "Explain Code", lambda: self._ai_action("explain"), "Ctrl+Shift+E")
        self._add_action(ai_menu, "Refactor Code", lambda: self._ai_action("refactor"), "Ctrl+Shift+R")
        self._add_action(ai_menu, "Write Tests", lambda: self._ai_action("tests"), "Ctrl+Shift+U")
        self._add_action(ai_menu, "Debug Help", lambda: self._ai_action("debug"), "Ctrl+Shift+D")
        ai_menu.addSeparator()
        
        # Phase 1, 2, 3 Integration: Agent Mode submenu
        mode_menu = ai_menu.addMenu("Agent Mode")
        self._add_action(mode_menu, "Build Mode", lambda: self._set_agent_mode("build"), "")
        self._add_action(mode_menu, "Explore Mode", lambda: self._set_agent_mode("explore"), "")
        self._add_action(mode_menu, "Debug Mode", lambda: self._set_agent_mode("debug"), "")
        self._add_action(mode_menu, "Plan Mode", lambda: self._set_agent_mode("plan"), "")
        ai_menu.addSeparator()
        
        # Phase 3 Integration: Skills and MCP
        self._add_action(ai_menu, "Browse Skills...", self._show_skills_browser, "")
        self._add_action(ai_menu, "MCP Connections...", self._show_mcp_connections, "")
        ai_menu.addSeparator()
        
        # Phase 4 Integration: TODO, Permission, and GitHub
        todo_menu = ai_menu.addMenu("Tasks & TODOs")
        self._add_action(todo_menu, "View Tasks...", self._show_todo_manager, "")
        self._add_action(todo_menu, "Add Task...", self._add_todo_task, "")
        todo_menu.addSeparator()
        self._add_action(todo_menu, "Complete Task", self._complete_todo_task, "")
        
        self._add_action(ai_menu, "Permission Settings...", self._show_permission_settings, "")
        self._add_action(ai_menu, "GitHub Integration...", self._show_github_integration, "")
        self._add_action(ai_menu, "Memory Manager...", self._show_memory_manager, "Ctrl+Shift+M")
        ai_menu.addSeparator()
        
        self._add_action(ai_menu, "AI Chat Focus", self._focus_ai_chat, "Ctrl+Shift+A")
        
        # Command Palette - AI-First mode uses Ctrl+K
        self._add_action(file_menu, "Command Palette...", self._show_command_palette, "Ctrl+K")
        ai_menu.addSeparator()
        self._add_action(ai_menu, "Clear Chat", self._ai_chat.clear_chat, "")

        # Terminal
        term_menu = mb.addMenu("Terminal")
        self._add_action(term_menu, "New Terminal", lambda: self._new_terminal(show_panel=True), "Ctrl+Shift+`")
        self._add_action(term_menu, "Kill Terminal", self._kill_current_terminal, "")
        term_menu.addSeparator()
        self._add_action(term_menu, "Toggle Terminal Panel", self._toggle_terminal, "Ctrl+J")

        # Window
        window_menu = mb.addMenu("Window")
        self._add_action(window_menu, "Minimize", self._minimize_window, "Ctrl+M")
        self._add_action(window_menu, "Zoom", self._zoom_window, "")
        self._add_action(window_menu, "Close", self._close_window, "Ctrl+W")

        # Help
        help_menu = mb.addMenu("Help")
        self._add_action(help_menu, "Cortex Documentation", self._open_documentation, "")
        self._add_action(help_menu, "What's New", self._show_whats_new, "")
        self._add_action(help_menu, "Automations", self._show_automations, "")
        self._add_action(help_menu, "Local Environments", self._show_local_envs, "")
        self._add_action(help_menu, "Worktrees", self._show_worktrees, "")
        self._add_action(help_menu, "Skills", self._show_skills_help, "")
        self._add_action(help_menu, "Model Context Protocol", self._show_mcp_help, "")
        self._add_action(help_menu, "Troubleshooting", self._show_troubleshooting, "")
        help_menu.addSeparator()
        self._add_action(help_menu, "Send Feedback", self._send_feedback, "")
        self._add_action(help_menu, "Start Trace Recording", self._start_trace, "")
        help_menu.addSeparator()
        self._add_action(help_menu, "Keyboard Shortcuts", self._show_keyboard_shortcuts, "F1")
        help_menu.addSeparator()
        self._add_action(help_menu, "About Cortex", self._show_about, "")

        # ═══════════════════════════════════════════════════════════════════════
        # Panel Toggle Button Group — right corner of menu bar
        # ═══════════════════════════════════════════════════════════════════════
        self._panel_toggle_bar = self._build_panel_toggle_bar()

        # QMenuBar doesn't support custom widgets in the same way, so we add
        # them as a QAction with a QWidget. We use a spacer trick:
        # Create a spacer action to push buttons right.
        spacer_action = QAction(self)
        spacer_action.setVisible(False)  # Won't show but allows layout control
        # Instead, we'll place buttons in a QWidget that sits on the right
        # by using a custom approach: a QWidget placed next to menu bar via layout.
        # For simplicity, add directly to menuBar's cornerWidget.
        self.menuBar().setCornerWidget(self._panel_toggle_bar, Qt.Corner.TopRightCorner)

    def _add_action(self, menu, text, slot, shortcut=""):
        action = QAction(text, self)
        if shortcut:
            action.setShortcut(QKeySequence(shortcut))
        action.triggered.connect(slot)
        menu.addAction(action)
        return action

    # ------------------------------------------------------------------
    # Toolbar
    # ------------------------------------------------------------------
    # Toolbar removed in AI-first mode - replaced by TopNavBar component
    # def _build_toolbar(self):


    # ------------------------------------------------------------------
    # Panel Toggle Buttons
    # ------------------------------------------------------------------
    def _build_panel_toggle_bar(self) -> QWidget:
        """Build a horizontal bar of 4 toggle buttons — one per panel group."""
        bar = QWidget()
        bar.setObjectName("panelToggleBar")
        bar.setStyleSheet("""
            QWidget#panelToggleBar {
                background: transparent;
                padding: 0px 8px;
            }
            QPushButton {
                background: transparent;
                border: none;
                border-radius: 4px;
                padding: 2px;
                margin: 1px 0px;
            }
            QPushButton:hover {
                background: rgba(255, 255, 255, 0.08);
            }
            QPushButton:pressed {
                background: rgba(255, 255, 255, 0.15);
            }
        """)

        layout = QHBoxLayout(bar)
        layout.setContentsMargins(0, 0, 4, 0)
        layout.setSpacing(1)

        from src.utils.icons import make_icon

        icon_size = 22
        btn_size = 26

        def _make_toggle(visible_icon: str, hidden_icon: str, tooltip_v: str, tooltip_h: str,
                         is_visible_getter, toggle_fn):
            """Single toggle button — switches icon/tooltip when panel visibility changes."""
            btn = QPushButton()
            btn.setFixedSize(btn_size, btn_size)
            btn.setCursor(Qt.CursorShape.PointingHandCursor)

            is_dark = self._settings.theme == "dark"
            icon_color = "#c8c8c8" if is_dark else "#444444"

            _visible = is_visible_getter()
            btn.setIcon(make_icon(visible_icon if _visible else hidden_icon, icon_color, icon_size))
            btn.setToolTip(tooltip_v if _visible else tooltip_h)

            def on_click():
                nonlocal _visible
                _visible = not _visible
                btn.setIcon(make_icon(visible_icon if _visible else hidden_icon, icon_color, icon_size))
                btn.setToolTip(tooltip_v if _visible else tooltip_h)
                toggle_fn(_visible)

            btn.clicked.connect(on_click)
            return btn

        # 1. Left Sidebar toggle
        layout.addWidget(_make_toggle(
            "panel-left-sidebar-visible", "panel-left-sidebar-hidden",
            "Hide Left Sidebar", "Show Left Sidebar",
            lambda: not getattr(self, '_left_sidebar_hidden', False),
            lambda v: self._toggle_left_sidebar(v)
        ))

        # 2. Right Sidebar (Explore/File Tree) toggle
        layout.addWidget(_make_toggle(
            "panel-right-sidebar-visible", "panel-right-sidebar-hidden",
            "Hide Explore Panel", "Show Explore Panel",
            lambda: not getattr(self, '_file_tree_hidden', False),
            lambda v: self._toggle_file_tree_panel(v)
        ))

        # 3. AI Chat toggle
        layout.addWidget(_make_toggle(
            "panel-ai-chat-visible", "panel-ai-chat-hidden",
            "Hide AI Chat", "Show AI Chat",
            lambda: not getattr(self, '_chat_panel_hidden', False),
            lambda v: self._toggle_ai_chat_panel(v)
        ))

        # 4. Review/Summary/Git panel toggle (all 3 tabs share one panel)
        layout.addWidget(_make_toggle(
            "panel-review-visible", "panel-review-hidden",
            "Hide Review Panel", "Show Review Panel",
            lambda: not getattr(self, '_review_panel_hidden', False),
            lambda v: self._toggle_review_panel(v)
        ))

        return bar

    def _make_spacer(self) -> QWidget:
        spacer = QWidget()
        spacer.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        return spacer

    # ------------------------------------------------------------------
    # Status Bar
    # ------------------------------------------------------------------
    def _build_status_bar(self):
        sb = self.statusBar()
        sb.setStyleSheet("""
            QStatusBar {
                background-color: #1e1e1e;
                color: #cccccc;
                border-top: 1px solid #2a2a2a;
            }
            QLabel {
                color: #cccccc;
            }
        """)
        self._status_file = QLabel("  No file open")
        self._status_cursor = QLabel("Ln 1, Col 1")
        self._status_lang = QLabel("Plain Text")
        self._status_ai = QLabel("AI: Ready")

        for lbl in [self._status_file, self._status_cursor, self._status_lang, self._status_ai]:
            sb.addWidget(lbl)

        sb.addPermanentWidget(QLabel("  Cortex AI Agent v1.0.15  "))

    def _update_status_cursor(self, line: int, col: int):
        self._status_cursor.setText(f"Ln {line}, Col {col}")

    def _update_status_file(self, filepath: str = None):
        if filepath:
            self._status_file.setText(f"  {shorten_path(filepath)}")
            self._status_lang.setText(detect_language(filepath).title())
        else:
            self._status_file.setText("  No file open")
            self._status_lang.setText("Plain Text")

    def _apply_initial_theme(self):
        """Apply the saved (or default) theme to all panels at startup."""
        from PyQt6.QtWidgets import QApplication as _App
        saved = self._settings.theme if hasattr(self._settings, 'theme') else "dark"
        theme_name = saved if isinstance(saved, str) and saved in ("dark", "light") else "dark"
        # Apply theme
        self._theme_manager.apply(theme_name, _App.instance())
        is_dark = theme_name == "dark"
        # Propagate to all panels
        self._ai_chat.set_theme(is_dark)
        self._sidebar.set_theme(is_dark)
        self._apply_welcome_theme(is_dark)
        # Theme button label
        if hasattr(self, '_theme_btn') and self._theme_btn:
            self._theme_btn.setText("☀️" if not is_dark else "🌙")
        
        # Apply to editor tabs
        self._editor_tabs.update_theme(is_dark)
        
        # Apply to terminal tab bar
        if isinstance(self._terminal_tabs.tabBar(), CleanTabBar):
            self._terminal_tabs.tabBar().set_dark(is_dark)
        
        # Style the tab widget panels
        tab_bg = "#1e1e1e" if is_dark else "#ffffff"
        tab_border = "#3e3e42" if is_dark else "#dee2e6"
        tab_style = f"""
            QTabWidget::pane {{ border: 1px solid {tab_border}; background: {tab_bg}; }}
            QTabWidget::tab-bar {{ left: 0px; }}
        """
        self._editor_tabs.setStyleSheet(tab_style)
        self._terminal_tabs.setStyleSheet(tab_style)
        
        # Apply to all terminal widgets
        for i in range(self._terminal_tabs.count()):
            term = self._terminal_tabs.widget(i)
            if isinstance(term, XTermWidget):
                term.set_theme(is_dark)
        
        # Apply global syntax highlighting fonts and colors
        self._apply_syntax_highlighting_fonts()

    def _apply_syntax_highlighting_fonts(self):
        """Apply Dracula-themed fonts and colors globally to code displays."""
        if not HAS_SYNTAX_HIGHLIGHTING:
            return
        
        try:
            # Initialize global code colorizer
            self._code_colorizer = UniversalCodeColorizer()
            
            # Apply monospace font (for code editors, terminal)
            mono_fonts = FONTS['mono']
            
            # Apply sans-serif font (for UI elements)
            sans_fonts = FONTS['sans']
            
            # Apply to all code editors
            if hasattr(self, '_code_editor') and self._code_editor:
                code_font = QFont()
                code_font.setFamily(mono_fonts[0])  # Use first preferred monospace font
                code_font.setPointSize(10)
                code_font.setFixedPitch(True)
                self._code_editor.setFont(code_font)
                log.info(f"[FONTS] Applied monospace font to editor: {mono_fonts[0]}")
            
            # Apply to terminal if available
            if hasattr(self, '_terminal') and self._terminal:
                term_font = QFont()
                term_font.setFamily(mono_fonts[0])
                term_font.setPointSize(9)
                term_font.setFixedPitch(True)
                self._terminal.setFont(term_font)
                log.info(f"[FONTS] Applied terminal font: {mono_fonts[0]}")
            
            # Log Dracula theme application
            log.info(f"[COLORS] Dracula theme applied with {len(DRACULA_COLORS)} color definitions")
            log.info(f"[LANGUAGES] 100+ programming languages supported")
            log.info(f"[FRAMEWORKS] React, Vue, Angular, Django, Flask, FastAPI, and 20+ frameworks")
            log.info(f"[MARKDOWN] Blue headings (#0047AB), White text (#ffffff)")
            
            # Display font stack information
            mono_stack = " -> ".join(mono_fonts)
            sans_stack = " -> ".join(sans_fonts)
            log.info(f"[FONT_STACK_MONO] {mono_stack}")
            log.info(f"[FONT_STACK_SANS] {sans_stack}")
            
        except Exception as e:
            log.warning(f"Error applying syntax highlighting fonts: {e}")
    
    def colorize_code(self, code, language='plaintext'):
        """
        Public method to colorize code with Dracula theme.
        Every language gets colors - NO white text fallback.
        """
        if not HAS_SYNTAX_HIGHLIGHTING or not hasattr(self, '_code_colorizer'):
            return code
        
        try:
            return self._code_colorizer.colorize(code, language)
        except Exception as e:
            log.warning(f"Error colorizing code for language '{language}': {e}")
            return code
    
    def colorize_markdown(self, markdown_text):
        """
        Public method to colorize Markdown with blue headings and white text.
        """
        if not HAS_SYNTAX_HIGHLIGHTING:
            return markdown_text
        
        try:
            return MarkdownColorizer.colorize(markdown_text)
        except Exception as e:
            log.warning(f"Error colorizing markdown: {e}")
            return markdown_text

    # ------------------------------------------------------------------
    # Signal Connections
    # ------------------------------------------------------------------
    def _connect_signals(self):
        # Sidebar signals
        self._sidebar.file_opened.connect(self._open_file)
        self._sidebar.file_search_opened.connect(self._open_file_at_line)
        self._sidebar.ai_action_requested.connect(self._ai_action)
        self._sidebar.file_renamed.connect(self._on_sidebar_file_renamed)
        self._sidebar.file_deleted.connect(self._on_sidebar_file_deleted)
        
        # Changed files panel signals
        self._sidebar.file_accepted.connect(self._on_accept_file_edit)
        self._sidebar.file_rejected.connect(self._on_reject_file_edit)
        self._sidebar.accept_all_requested.connect(self._on_accept_all_files)
        self._sidebar.reject_all_requested.connect(self._on_reject_all_files)

        # Sidebar footer gear button → Memory Manager
        self._sidebar.settings_requested.connect(self._show_memory_manager)

        # Project manager
        self._project_manager.project_opened.connect(self._on_project_opened)
        self._project_manager.project_closed.connect(self._on_project_closed)


        # Editor tab changes
        self._editor_tabs.currentChanged.connect(self._on_tab_changed)
        
        # File manager undo/redo signals
        if hasattr(self, '_file_manager'):
            self._file_manager.file_deleted.connect(self._on_file_deleted_for_undo)
            self._file_manager.file_restored.connect(self._on_file_restored_for_redo)

        # AI chat - ONLY connect signals here to avoid duplicates
        self._ai_chat.message_sent.connect(self._on_ai_chat_message)
        self._ai_chat.model_changed.connect(self._on_model_changed)
        self._ai_chat.vision_history_sync.connect(self._ai_agent.inject_vision_history)
        self._ai_agent.response_chunk.connect(self._ai_chat.on_chunk)
        self._ai_agent.response_complete.connect(self._ai_chat.on_complete)
        self._ai_agent.response_complete.connect(self._on_ai_task_complete)
        self._ai_agent.request_error.connect(self._ai_chat.on_error)
        self._ai_agent.file_generated.connect(self._open_file)
        # TEMPORARILY DISABLED - file_tracker was part of deleted agentic code
        # self._ai_agent.file_edited_diff.connect(self._file_tracker.add_edit)
        self._ai_agent.file_edited_diff.connect(self._on_file_edited_diff_for_js)
        self._ai_agent.file_edited_diff.connect(self._on_inline_edit_diff)
        self._ai_agent.file_edited_diff.connect(self._ai_chat.on_file_edited_diff)  # populate Changed Files panel with +/- counts
        self._ai_agent.file_edited_diff.connect(self._ai_chat.show_diff_card)       # show diff viewer card in chat
        # File operation cards — animated create/edit cards
        self._ai_agent.file_creating_started.connect(self._on_file_creating_started)
        self._ai_agent.file_editing_started.connect(self._on_file_editing_started)
        self._ai_agent.file_operation_completed.connect(self._on_file_operation_completed)
        self._ai_agent.tool_activity.connect(self._ai_chat.show_tool_activity)
        self._ai_agent.directory_contents.connect(self._ai_chat.show_directory_contents)
        self._ai_agent.directory_contents.connect(self._on_directory_contents_for_tree)
        self._ai_agent.thinking_started.connect(self._ai_chat.show_thinking)
        self._ai_agent.thinking_stopped.connect(self._ai_chat.hide_thinking)
        self._ai_agent.todos_updated.connect(self._ai_chat.update_todos)
        self._ai_agent.tool_summary_ready.connect(self._ai_chat.show_tool_summary)
        # Recovery: context compaction status + turn-limit continuation
        self._ai_agent.agent_status_update.connect(self._ai_chat.on_agent_status_update)
        self._ai_agent.turn_limit_hit.connect(self._ai_chat.on_turn_limit_hit)
        # Token budget: real-time context usage bar
        self._ai_agent.context_budget_update.connect(self._ai_chat.on_context_budget_update)
        # Permission gate: agent → chat UI shows card; user response → agent continues
        self._ai_agent.permission_requested.connect(self._ai_chat._on_permission_request)
        self._ai_chat.permission_decided.connect(self._ai_agent.on_permission_respond)
        
        # Active signal connections
        self._ai_chat.generate_plan_requested.connect(self._on_generate_plan)
        self._ai_chat.open_file_requested.connect(self._open_file)
        self._ai_chat.open_file_at_line_requested.connect(self._open_file_at_line)
        log.info(f"[Diff-Debug] Connecting show_diff_requested to _on_show_diff. Signal exists: {hasattr(self._ai_chat, 'show_diff_requested')}")
        self._ai_chat.show_diff_requested.connect(self._on_show_diff)
        self._ai_chat.answer_question_requested.connect(self._ai_agent.user_responded)
        self._ai_chat.smart_paste_check_requested.connect(self._on_smart_paste_check)
        
        # Todo toggle (logs only - state managed by bridge/JS)
        self._ai_chat.toggle_todo_requested.connect(self._on_toggle_todo)
        
        # Interactive questions from agent to user
        self._ai_agent.user_question_requested.connect(self._on_ai_question_requested)

        # ========== CortexDiffBridge: wire accept/reject signals ==========
        # This connects useDiffInIDE.py's CortexDiffBridge to Cortex's FEC card
        # accept/reject signals so the agent can await user confirmation of edits.
        try:
            import importlib as _il
            _diff_ide_mod = _il.import_module("agent.src.hooks.useDiffInIDE")
            _cdb = _diff_ide_mod.CortexDiffBridge.instance()
            _cdb.register_accept_signal(self._ai_chat.accept_file_edit_requested)
            _cdb.register_reject_signal(self._ai_chat.reject_file_edit_requested)
            log.info("[CortexDiffBridge] Accept/Reject signals wired")
        except Exception as _cdb_err:
            log.warning(f"[CortexDiffBridge] Signal wiring skipped: {_cdb_err}")

        # Terminal tab changes
        self._terminal_tabs.currentChanged.connect(self._on_terminal_tab_changed)
        
        # ========== PERMISSION SYSTEM CONNECTION (NEW) ==========
        # Connect AI agent to UI for permission dialogs
        self._ai_agent.set_ui_parent(self)
        log.info("Permission system initialized and connected to main window")

    # ------------------------------------------------------------------
    # Actions
    # ------------------------------------------------------------------
    def _new_file(self):
        # Get current theme state and apply it to the new editor
        is_dark = self._theme_manager.is_dark
        editor = CodeEditor(language="python")
        if is_dark:
            editor.apply_dark_theme()
        else:
            editor.apply_light_theme()
        idx = self._editor_tabs.addTab(editor, "untitled.py")
        self._editor_tabs.setCurrentIndex(idx)
        editor.cursor_position_changed.connect(self._update_status_cursor)
        editor.inline_edit_submitted.connect(self._on_inline_edit_submitted)
        editor.inline_edit_cancelled.connect(self._on_inline_edit_cancelled)
        editor.inline_diff_requested.connect(self._on_inline_diff_requested)
        editor.code_copied.connect(self._on_code_copied)

    def _open_file_dialog(self):
        path, _ = QFileDialog.getOpenFileName(self, "Open File",
                                               str(self._project_manager.root or Path.home()))
        if path:
            self._open_file(path)
    
    def _find_file_in_project(self, filename: str) -> str | None:
        """Search for a file by name in the project directory (recursive)."""
        if not self._project_manager.root:
            return None
        
        # Extract just the filename (remove any directory components)
        from pathlib import Path as PathLib
        clean_filename = PathLib(filename).name
        
        root = Path(self._project_manager.root)
        
        # Search recursively for the file
        try:
            for file_path in root.rglob(clean_filename):
                if file_path.is_file():
                    return str(file_path)
        except Exception:
            pass
        
        return None

    def _open_folder_dialog(self):
        folder = QFileDialog.getExistingDirectory(self, "Open Folder",
                                                   str(Path.home()))
        if folder:
            self._open_folder_programmatic(folder)

    def _open_folder_programmatic(self, folder: str):
        """Open a folder as the active project (no dialog, usable from argv/drag-drop)."""
        self._project_manager.open(folder)
        
        # Set project root for file tree (with delay to ensure UI is ready)
        QTimer.singleShot(100, lambda: self._set_project_root(folder))
        
        # Initialize Git repository
        if hasattr(self, '_git_manager'):
            self._git_manager.set_repository(folder)
            log.info(f"[GIT] Repository set to: {folder}")
            # Update Git summary after a short delay
            QTimer.singleShot(300, self._update_git_summary)
        
        try:
            from src.core.lsp_manager import get_lsp_manager
            get_lsp_manager().set_project_root(folder)
        except Exception:
            pass

    def _new_project(self):
        folder = QFileDialog.getExistingDirectory(self, "Select Folder for New Project",
                                                   str(Path.home()))
        if folder:
            # For now, just open it as a project. 
            # Could add a template/scaffolding step here later.
            self._project_manager.open(folder)
            
            # Set project root for file tree (with delay to ensure UI is ready)
            QTimer.singleShot(100, lambda: self._set_project_root(folder))
            
            # Initialize Git repository
            if hasattr(self, '_git_manager'):
                self._git_manager.set_repository(folder)
                log.info(f"[GIT] Repository set to: {folder}")
                QTimer.singleShot(300, self._update_git_summary)

    def _open_file(self, filepath: str):
        # Normalize path (convert forward slashes to backslashes on Windows)
        filepath = os.path.normpath(filepath)
        path = Path(filepath)
        log.info(f"Opening file: {filepath}")
        
        # If file doesn't exist, try to find it in the project
        if not path.exists() or not path.is_file():
            # Try searching in project directory
            found_path = self._find_file_in_project(path.name)
            if found_path:
                log.info(f"Found file in project: {found_path}")
                path = Path(found_path)
                filepath = str(found_path)
            else:
                log.warning(f"File skip (not found or dir): {filepath}")
                return
        
        # Check file extension for images and documents
        file_ext = path.suffix.lower()
        
        # Handle image files
        image_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.svg', '.webp', '.ico', '.tiff', '.tif'}
        if file_ext in image_extensions:
            self._open_image_file(filepath)
            return
        
        # Handle PDF files
        if file_ext == '.pdf':
            self._open_pdf_file(filepath)
            return
        
        # Handle Office documents
        office_extensions = {'.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx'}
        if file_ext in office_extensions:
            self._open_office_file(filepath)
            return
            
        # Initialize file snapshots dict for diff generation
        if not hasattr(self, '_file_snapshots'):
            self._file_snapshots = {}
        
        if self._file_manager.is_binary(filepath):
            log.info(f"File skip (binary): {filepath}")
            QMessageBox.information(self, "Binary File",
                                    f"'{path.name}' is a binary file and cannot be edited.")
            return
            
        try:
            content = self._file_manager.read(filepath)
            if content is None:
                log.error(f"Failed to read content: {filepath}")
                return
            
            # Store original snapshot for diff generation (only if not already stored)
            # This preserves the FIRST version opened, allowing multiple diffs
            if filepath not in self._file_snapshots:
                self._file_snapshots[filepath] = content
                log.info(f"[Snapshot] Stored initial snapshot: {filepath} ({len(content)} chars)")
            else:
                log.info(f"[Snapshot] Keeping existing snapshot for: {filepath}")
                
            log.info(f"Content read ({len(content)} chars). Detecting language...")
            language = detect_language(filepath)
            log.info(f"Language detected: {language}. Opening index in tabs...")
            
            # Get current theme state and pass it to the editor
            is_dark = self._theme_manager.is_dark
            idx = self._editor_tabs.open_file(filepath, content, language, is_dark)
            
            # Connect cursor signal for the new editor
            editor = self._editor_tabs.widget(idx)
            if isinstance(editor, CodeEditor):
                editor.cursor_position_changed.connect(self._update_status_cursor)
                editor.inline_edit_submitted.connect(self._on_inline_edit_submitted)
                editor.inline_edit_cancelled.connect(self._on_inline_edit_cancelled)
                editor.inline_diff_requested.connect(self._on_inline_diff_requested)
                editor.code_copied.connect(self._on_code_copied)

            self._update_status_file(filepath)
            is_dark = self._theme_manager.is_dark
            if isinstance(editor, CodeEditor):
                # Block content_modified signal during theme set (rehighlight triggers it)
                from PyQt6.QtCore import QSignalBlocker
                with QSignalBlocker(editor.document()):
                    editor.set_theme(is_dark)
            log.info(f"File opened successfully: {filepath}")
        except Exception as e:
            log.error(f"Error opening file {filepath}: {e}", exc_info=True)

    def _open_image_file(self, filepath: str):
        """Open an image file in a viewer tab, scaled to fit window."""
        from PyQt6.QtWidgets import QLabel, QScrollArea
        from PyQt6.QtCore import Qt
        from PyQt6.QtGui import QPixmap
        
        try:
            log.info(f"Opening image file: {filepath}")
            path = Path(filepath)
            
            # Load image
            pixmap = QPixmap(filepath)
            
            if pixmap.isNull():
                log.error(f"Failed to load image: {filepath}")
                QMessageBox.warning(self, "Error", f"Could not load image: {path.name}")
                return
            
            # Get available size (tab widget size minus some padding)
            tab_size = self._editor_tabs.size()
            max_width = max(tab_size.width() - 40, 400)  # Min 400px width
            max_height = max(tab_size.height() - 80, 300)  # Min 300px height
            
            # Scale image to fit while maintaining aspect ratio
            scaled_pixmap = pixmap.scaled(
                max_width, 
                max_height,
                Qt.AspectRatioMode.KeepAspectRatio,
                Qt.TransformationMode.SmoothTransformation
            )
            
            # Create label with scaled image
            image_label = QLabel()
            image_label.setPixmap(scaled_pixmap)
            image_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            
            # Create scroll area
            scroll_area = QScrollArea()
            scroll_area.setWidgetResizable(True)
            scroll_area.setAlignment(Qt.AlignmentFlag.AlignCenter)
            scroll_area.setWidget(image_label)
            
            # Add to tabs
            idx = self._editor_tabs.addTab(scroll_area, path.name)
            self._editor_tabs.setTabToolTip(idx, filepath)
            self._editor_tabs.setCurrentIndex(idx)
            
            self._update_status_file(filepath)
            log.info(f"Image file opened successfully: {filepath} (scaled to {scaled_pixmap.width()}x{scaled_pixmap.height()})")
            
        except Exception as e:
            log.error(f"Error opening image file {filepath}: {e}", exc_info=True)
            QMessageBox.warning(self, "Error", f"Could not open image: {e}")

    def _open_pdf_file(self, filepath: str):
        """Open a PDF file by rendering pages as images."""
        from PyQt6.QtWidgets import QLabel, QScrollArea, QVBoxLayout, QWidget
        from PyQt6.QtCore import Qt
        from PyQt6.QtGui import QPixmap, QImage
        import fitz  # PyMuPDF
        
        try:
            log.info(f"Opening PDF file: {filepath}")
            path = Path(filepath)
            
            # Open PDF with PyMuPDF
            doc = fitz.open(filepath)
            
            # Store page count immediately
            total_pages = doc.page_count
            
            if total_pages == 0:
                QMessageBox.warning(self, "Error", "PDF has no pages")
                return
            
            # Create scrollable container for all pages
            container = QWidget()
            layout = QVBoxLayout(container)
            layout.setSpacing(10)
            layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
            
            # Get available width
            tab_width = self._editor_tabs.width() - 60
            
            # Render each page as an image
            for page_num in range(min(total_pages, 50)):  # Limit to 50 pages
                page = doc[page_num]
                
                # Calculate zoom to fit width
                zoom = tab_width / page.rect.width
                mat = fitz.Matrix(zoom, zoom)
                
                # Render page to pixmap
                pix = page.get_pixmap(matrix=mat)
                
                # Convert to QImage
                img = QImage(pix.samples, pix.width, pix.height, pix.stride, QImage.Format.Format_RGB888)
                pixmap = QPixmap.fromImage(img)
                
                # Create label for page
                label = QLabel()
                label.setPixmap(pixmap)
                label.setAlignment(Qt.AlignmentFlag.AlignCenter)
                label.setStyleSheet("background-color: white; border: 1px solid #ccc;")
                
                layout.addWidget(label)
            
            doc.close()
            
            # Create scroll area
            scroll = QScrollArea()
            scroll.setWidgetResizable(True)
            scroll.setWidget(container)
            scroll.setStyleSheet("background-color: #f0f0f0;")
            
            # Add to tabs
            idx = self._editor_tabs.addTab(scroll, path.name)
            self._editor_tabs.setTabToolTip(idx, f"{filepath} ({total_pages} pages)")
            self._editor_tabs.setCurrentIndex(idx)
            
            self._update_status_file(filepath)
            log.info(f"PDF file opened successfully: {filepath} ({total_pages} pages)")
            
        except ImportError:
            log.error("PyMuPDF (fitz) not installed")
            QMessageBox.warning(self, "Error", "PyMuPDF not installed. Run: pip install PyMuPDF")
        except Exception as e:
            log.error(f"Error opening PDF file {filepath}: {e}", exc_info=True)
            QMessageBox.warning(self, "Error", f"Could not open PDF: {e}")

    def _open_office_file(self, filepath: str):
        """Open Office documents (Word, Excel, PowerPoint) inside the IDE as formatted text."""
        from PyQt6.QtWidgets import QTextEdit
        from PyQt6.QtCore import Qt
        
        try:
            log.info(f"Opening Office file: {filepath}")
            path = Path(filepath)
            file_ext = path.suffix.lower()
            
            # Create text viewer
            text_viewer = QTextEdit()
            text_viewer.setReadOnly(True)
            text_viewer.setLineWrapMode(QTextEdit.LineWrapMode.WidgetWidth)
            
            content_html = ""
            
            if file_ext in ['.docx']:
                # Read Word document
                try:
                    from docx import Document
                    doc = Document(filepath)
                    
                    content_html = f"<h2>{path.name}</h2><hr>"
                    
                    for para in doc.paragraphs:
                        if para.text.strip():
                            # Check if it's a heading based on style
                            if para.style.name.startswith('Heading'):
                                content_html += f"<h3>{para.text}</h3>"
                            else:
                                content_html += f"<p>{para.text}</p>"
                    
                    # Add tables
                    for table in doc.tables:
                        content_html += "<table border='1' cellpadding='5'>"
                        for row in table.rows:
                            content_html += "<tr>"
                            for cell in row.cells:
                                content_html += f"<td>{cell.text}</td>"
                            content_html += "</tr>"
                        content_html += "</table><br>"
                    
                except ImportError:
                    content_html = f"<p style='color:red'>Error: python-docx library not installed.<br>Install with: pip install python-docx</p>"
                    log.error("python-docx library not installed")
                except Exception as e:
                    content_html = f"<p style='color:red'>Error reading document: {e}</p>"
                    log.error(f"Error reading docx: {e}")
            
            elif file_ext in ['.xlsx', '.xls']:
                # Read Excel document
                try:
                    if file_ext == '.xlsx':
                        import openpyxl
                        wb = openpyxl.load_workbook(filepath, data_only=True)
                        sheetnames = wb.sheetnames
                        
                        content_html = f"<h2>{path.name}</h2><hr>"
                        
                        for sheet_name in sheetnames:
                            sheet = wb[sheet_name]
                            content_html += f"<h3>Sheet: {sheet_name}</h3>"
                            content_html += "<table border='1' cellpadding='5' style='border-collapse:collapse'>"
                            
                            # Read first 100 rows max
                            row_count = 0
                            for row in sheet.iter_rows(max_row=100):
                                content_html += "<tr>"
                                for cell in row:
                                    value = cell.value if cell.value is not None else ""
                                    content_html += f"<td>{value}</td>"
                                content_html += "</tr>"
                                row_count += 1
                                if row_count >= 100:
                                    content_html += "<tr><td colspan='100'>... (showing first 100 rows)</td></tr>"
                                    break
                            
                            content_html += "</table><br>"
                    
                    elif file_ext == '.xls':
                        import xlrd
                        wb = xlrd.open_workbook(filepath)
                        
                        content_html = f"<h2>{path.name}</h2><hr>"
                        
                        for sheet_idx in range(wb.nsheets):
                            sheet = wb.sheet_by_index(sheet_idx)
                            content_html += f"<h3>Sheet: {sheet.name}</h3>"
                            content_html += "<table border='1' cellpadding='5' style='border-collapse:collapse'>"
                            
                            # Read first 100 rows max
                            for row_idx in range(min(sheet.nrows, 100)):
                                content_html += "<tr>"
                                for col_idx in range(sheet.ncols):
                                    value = sheet.cell_value(row_idx, col_idx)
                                    content_html += f"<td>{value}</td>"
                                content_html += "</tr>"
                            
                            if sheet.nrows > 100:
                                content_html += "<tr><td colspan='100'>... (showing first 100 rows)</td></tr>"
                            
                            content_html += "</table><br>"
                    
                except ImportError as e:
                    content_html = f"<p style='color:red'>Error: Required library not installed.<br>Install with: pip install openpyxl xlrd</p>"
                    log.error(f"Library not installed: {e}")
                except Exception as e:
                    content_html = f"<p style='color:red'>Error reading spreadsheet: {e}</p>"
                    log.error(f"Error reading xlsx/xls: {e}")
            
            elif file_ext == '.doc':
                # Old Word format - try to extract text
                try:
                    # Try using textract if available
                    import textract
                    text = textract.process(filepath).decode('utf-8', errors='ignore')
                    content_html = f"<h2>{path.name}</h2><hr>"
                    content_html += f"<pre style='white-space:pre-wrap;font-family:Arial,sans-serif'>{text}</pre>"
                except ImportError:
                    content_html = f"""<p style='color:orange'>
                        <b>Old Word Format (.doc)</b><br><br>
                        This file uses the older .doc format which requires additional libraries.<br><br>
                        <b>Options:</b><br>
                        1. Convert to .docx format (open in Word and Save As .docx)<br>
                        2. Install textract: <code>pip install textract</code><br>
                        3. Open externally with Microsoft Word
                    </p>"""
                    log.warning(f"Old .doc format not supported without textract: {filepath}")
                except Exception as e:
                    content_html = f"<p style='color:red'>Error reading .doc file: {e}</p>"
                    log.error(f"Error reading doc: {e}")
            
            else:
                content_html = f"<p>File type '{file_ext}' is not supported for internal viewing.</p>"
            
            # Set content with styling
            text_viewer.setHtml(f"""
            <html>
            <head>
                <style>
                    body {{ font-family: Arial, sans-serif; padding: 20px; line-height: 1.6; }}
                    h2 {{ color: #333; border-bottom: 2px solid #0078d4; padding-bottom: 10px; }}
                    h3 {{ color: #555; margin-top: 20px; }}
                    table {{ margin: 10px 0; border: 1px solid #ddd; }}
                    td {{ padding: 8px; border: 1px solid #ddd; }}
                    p {{ margin: 10px 0; }}
                </style>
            </head>
            <body>
                {content_html}
            </body>
            </html>
            """)
            
            # Add to tabs
            idx = self._editor_tabs.addTab(text_viewer, path.name)
            self._editor_tabs.setTabToolTip(idx, filepath)
            self._editor_tabs.setCurrentIndex(idx)
            
            self._update_status_file(filepath)
            log.info(f"Office file opened internally: {filepath}")
            
        except Exception as e:
            log.error(f"Error opening Office file {filepath}: {e}", exc_info=True)
            QMessageBox.warning(self, "Error", f"Could not open document: {e}")


    def _diff_cache_path(self):
        try:
            from pathlib import Path
            return Path.home() / ".cortex" / "diff_cache.json"
        except Exception:
            return None

    def _load_diff_cache(self):
        if hasattr(self, '_diff_cache'):
            return self._diff_cache
        self._diff_cache = {}
        try:
            cache_path = self._diff_cache_path()
            if cache_path and cache_path.exists():
                import json
                self._diff_cache = json.loads(cache_path.read_text(encoding='utf-8')) or {}
        except Exception:
            self._diff_cache = {}
        return self._diff_cache

    def _save_diff_cache(self):
        try:
            cache_path = self._diff_cache_path()
            if not cache_path:
                return
            cache_path.parent.mkdir(parents=True, exist_ok=True)
            import json
            cache_path.write_text(json.dumps(self._diff_cache), encoding='utf-8')
        except Exception:
            pass

    def _on_show_diff(self, file_path: str):
        """Show diff in Qt dialog window — triggered by Diff button click in chat."""
        log.info(f"[Diff] _on_show_diff called with: {file_path}")
        original, modified = '', ''

        # 1. Try the Python diff data store (most reliable)
        import os
        import subprocess
        from pathlib import Path
        normalized_requested = os.path.normcase(os.path.normpath(file_path))
        
        if hasattr(self, '_diff_data_store'):
            # Direct check
            if file_path in self._diff_data_store:
                original, modified = self._diff_data_store[file_path]
                log.info(f"[Diff] Found in _diff_data_store (direct): {file_path}")
            else:
                # Iterative normalized check
                log.debug(f"[Diff] Checking normalized path for: {normalized_requested}")
                for k, v in self._diff_data_store.items():
                    if os.path.normcase(os.path.normpath(k)) == normalized_requested:
                        original, modified = v
                        log.info(f"[Diff] Found in _diff_data_store (normalized): {k}")
                        break

        # 2b. Fallback to persisted diff cache
        if not modified:
            cache = self._load_diff_cache()
            cached = None
            if cache:
                norm = os.path.normcase(os.path.normpath(file_path))
                cached = cache.get(file_path) or cache.get(norm)
                if not cached:
                    # try to find by normalized keys
                    for k, v in cache.items():
                        if os.path.normcase(os.path.normpath(k)) == norm:
                            cached = v
                            break
            if cached:
                original = cached.get('original', '')
                modified = cached.get('modified', '')
                log.info(f"[Diff] Found in diff cache: {file_path}")
        # 3. Fallback to file_tracker
        if not modified:
            edit_info = self._file_tracker.get_edit(file_path)
            if edit_info:
                original = edit_info.original_content if edit_info.edit_type != 'C' else ''
                modified = edit_info.new_content
                log.info(f"[Diff] Found in file_tracker: {file_path}")

        if not modified:
            # Fallback: try Git diff against HEAD if repository available
            try:
                project_root = getattr(self, '_project_manager', None).root if hasattr(self, '_project_manager') else None
                if project_root and os.path.isdir(os.path.join(project_root, '.git')) and os.path.exists(file_path):
                    rel_path = os.path.relpath(file_path, project_root)
                    # Try to load original from HEAD (tracked files)
                    git_show = subprocess.run(
                        ['git', '-C', project_root, 'show', f'HEAD:{rel_path}'],
                        capture_output=True, text=True,
                        encoding='utf-8', errors='replace',
                        creationflags=subprocess.CREATE_NO_WINDOW if os.name == 'nt' else 0
                    )
                    if git_show.returncode == 0:
                        original = git_show.stdout
                        modified = Path(file_path).read_text(encoding='utf-8', errors='replace')
                        log.info(f"[Diff] Loaded diff via git for: {file_path}")
                    else:
                        # If untracked, treat original as empty
                        git_ls = subprocess.run(
                            ['git', '-C', project_root, 'ls-files', '--error-unmatch', rel_path],
                            capture_output=True, text=True,
                            encoding='utf-8', errors='replace',
                            creationflags=subprocess.CREATE_NO_WINDOW if os.name == 'nt' else 0
                        )
                        if git_ls.returncode != 0:
                            original = ''
                            modified = Path(file_path).read_text(encoding='utf-8', errors='replace')
                            log.info(f"[Diff] Loaded diff for untracked file: {file_path}")
            except Exception as _ge:
                log.debug(f"[Diff] Git fallback failed: {_ge}")

        # Final fallback: Use file snapshot taken when opened
        if not modified and hasattr(self, '_file_snapshots') and file_path in self._file_snapshots:
            original = self._file_snapshots[file_path]
            try:
                modified = Path(file_path).read_text(encoding='utf-8', errors='replace')
                if original != modified:
                    log.info(f"[Diff] Using file snapshot for diff: {file_path}")
                else:
                    log.info(f"[Diff] File unchanged since opened: {file_path}")
                    original = ''
                    modified = ''
            except Exception as e:
                log.warning(f"[Diff] Failed to read current file content: {e}")
                original = ''
                modified = ''

        if not modified:
            log.warning(f"[Diff] No diff data found for {file_path}")
            log.debug(f"[Diff] _diff_data_store keys: {list(getattr(self, '_diff_data_store', {}).keys())}")
            if hasattr(self, '_ai_chat'):
                import os
                filename = os.path.basename(file_path)
                self._ai_chat.add_system_message(f"⚠️ No diff data available for {filename}. It was not edited in this session.")
            return

        log.info(f"[Diff] Opening diff tab for {file_path} (+{len(modified)} chars)")
        is_dark = self._theme_manager.is_dark
        self._editor_tabs.open_diff_tab(file_path, original, modified, is_dark)

    def _on_inline_edit_submitted(self, prompt: str, selection_text: str, line_range: tuple):
        editor = self._editor_tabs.currentWidget()
        if not isinstance(editor, CodeEditor):
            self._ai_chat.add_system_message("Open a file to use inline edit.")
            return

        file_path = self._editor_tabs.current_filepath()
        if not file_path:
            self._ai_chat.add_system_message("Open a file to use inline edit.")
            return

        start_line, end_line = line_range
        if start_line == end_line:
            line_range_text = f"{start_line}"
        else:
            line_range_text = f"{start_line}-{end_line}"

        rel_path = file_path
        project_root = self._project_manager.root
        if project_root:
            try:
                rel_path = str(Path(file_path).resolve().relative_to(Path(project_root).resolve()))
            except Exception:
                rel_path = file_path

        self._inline_edit_context = {
            "file_path": os.path.normpath(file_path),
            "relative_path": rel_path,
            "line_range": line_range,
            "selection_text": selection_text,
            "editor": editor,
            "diff": None,
        }

        self._ai_chat.add_system_message(
            f"Inline edit: `{rel_path}` lines {line_range_text}"
        )
        self._ai_chat._add_ai_bubble_streaming()

        inline_prompt = (
            "Inline edit request.\n"
            f"File (project-relative): {rel_path}\n"
            f"Selection lines: {line_range_text}\n"
            "Selected code:\n"
            "```\n"
            f"{selection_text}\n"
            "```\n"
            "Instruction:\n"
            f"{prompt}\n\n"
            "Constraints:\n"
            "- Apply changes only within the selection unless absolutely required.\n"
            "- Use file editing tools (prefer replace_lines) to update the file.\n"
            "- Do not modify other files.\n"
            "- Use the project-relative path above.\n"
        )
        self._ai_agent.chat(inline_prompt)

    def _on_inline_edit_cancelled(self):
        self._inline_edit_context = None

    def _on_inline_diff_requested(self):
        context = self._inline_edit_context or {}
        file_path = context.get("file_path")
        if not file_path:
            return

        diff_pair = context.get("diff")
        if diff_pair:
            original, modified = diff_pair
            is_dark = self._theme_manager.is_dark
            self._editor_tabs.open_diff_tab(file_path, original, modified, is_dark)
            return

        self._on_show_diff(file_path)

    def _on_inline_edit_diff(self, file_path: str, original: str, new_content: str):
        context = self._inline_edit_context
        if not context:
            return

        target = context.get("file_path")
        if not target:
            return

        if os.path.normcase(os.path.normpath(file_path)) != os.path.normcase(os.path.normpath(target)):
            return

        editor = context.get("editor")
        if not isinstance(editor, CodeEditor):
            return

        import difflib

        diff_lines = list(difflib.unified_diff(
            original.splitlines(),
            new_content.splitlines(),
            fromfile="Original",
            tofile="Modified",
            lineterm=""
        ))
        if not diff_lines:
            diff_text = "No changes detected."
        else:
            max_lines = 200
            if len(diff_lines) > max_lines:
                diff_lines = diff_lines[:max_lines]
                diff_lines.append("... (diff truncated)")
            diff_text = "\n".join(diff_lines)

        context["diff"] = (original, new_content)
        editor.show_inline_diff(diff_text)

    def _on_file_edited_diff_for_js(self, file_path: str, original: str, new_content: str):
        """Store diff data in Python dict for the Qt dialog viewer."""
        if not hasattr(self, '_diff_data_store'):
            self._diff_data_store = {}
        norm_path = os.path.normcase(os.path.normpath(file_path))
        self._diff_data_store[file_path] = (original, new_content)
        self._diff_data_store[norm_path] = (original, new_content)

        # ── Invalidate file_manager cache so _open_file reads fresh content ──
        try:
            resolved = str(Path(file_path).resolve())
            self._file_manager._file_cache.put(resolved, new_content)
            self._file_manager._hash_cache[resolved] = self._file_manager._compute_hash(new_content)
            if resolved in self._file_manager._open_files:
                self._file_manager._open_files[resolved] = new_content
            log.debug(f"[Diff] Updated file_manager cache for: {file_path} ({len(new_content)} chars)")
        except Exception as e:
            log.debug(f"[Diff] Cache update failed: {e}")
        # Persist to diff cache for cross-session diff viewing
        cache = self._load_diff_cache()
        cache[file_path] = {
            'original': original,
            'modified': new_content,
            'ts': int(__import__('time').time())
        }
        # Prune cache to last 100 entries
        if len(cache) > 100:
            items = sorted(cache.items(), key=lambda kv: kv[1].get('ts', 0), reverse=True)
            self._diff_cache = dict(items[:100])
        else:
            self._diff_cache = cache
        self._save_diff_cache()
        log.info(f"[Diff] Stored diff data for: {file_path} (original: {len(original)} chars, new: {len(new_content)} chars)")
        log.debug(f"[Diff] _diff_data_store now has {len(self._diff_data_store)} entries")

        # Update sidebar changed files panel
        try:
            edit_type = "C" if not original else "M"
            self._sidebar.add_changed_file(file_path, edit_type)
        except Exception as e:
            log.debug(f"Sidebar update skipped: {e}")

    # ============================================================
    # FILE OPERATION CARDS (Create/Edit with animation)
    # ============================================================
    
    def _on_file_creating_started(self, file_path: str):
        """Show 'Creating file...' card with pulse animation."""
        try:
            card_id = self._ai_chat.show_file_creating_card(file_path)
            if not hasattr(self, '_file_op_cards'):
                self._file_op_cards = {}
            self._file_op_cards[file_path] = card_id
            log.debug(f"[FileOp] Started creating card ({card_id}) for: {file_path}")
        except Exception as e:
            log.debug(f"[FileOp] Failed to show creating card: {e}")

    def _on_file_editing_started(self, file_path: str):
        """Show 'Editing file...' card with pulse animation."""
        try:
            if not hasattr(self, '_file_op_cards'):
                self._file_op_cards = {}
            # If there's already an active card for this file, remove/complete it first
            old_card_id = self._file_op_cards.get(file_path)
            if old_card_id:
                self._ai_chat.dismiss_file_op_card(old_card_id)
                self._file_op_cards.pop(file_path, None)
            card_id = self._ai_chat.show_file_editing_card(file_path)
            self._file_op_cards[file_path] = card_id
            log.debug(f"[FileOp] Started editing card ({card_id}) for: {file_path}")
        except Exception as e:
            log.debug(f"[FileOp] Failed to show editing card: {e}")

    def _on_file_operation_completed(self, _unused_card_id: str, file_path: str, content: str, op_type: str):
        """Transform operation card to show completed file."""
        try:
            # Use the card_id stored when the card was first created (from show_file_*_card)
            # The card_id from agent_bridge is different — it was generated before the JS card was made.
            real_card_id = getattr(self, '_file_op_cards', {}).get(file_path)
            if not real_card_id:
                log.debug(f"[FileOp] No stored card_id for {file_path}, skipping completion")
                return
            if op_type == "create":
                self._ai_chat.complete_file_creating_card(real_card_id, file_path, content)
            else:
                original = ""
                if hasattr(self, '_diff_data_store') and file_path in self._diff_data_store:
                    original, _ = self._diff_data_store[file_path]
                self._ai_chat.complete_file_editing_card(real_card_id, file_path, original, content)
            # Clean up stored card_id
            self._file_op_cards.pop(file_path, None)
            log.debug(f"[FileOp] Completed {op_type} card for: {file_path}")
        except Exception as e:
            log.debug(f"[FileOp] Failed to complete operation card: {e}")

    def _on_ai_question_requested(self, question_payload: dict):
        """Handle AI asking a question that requires user response in chat."""
        log.info(f"AI requested user input: {question_payload.get('text', question_payload.get('question', ''))[:50]}...")
        # Structuring the question info for the JS UI
        # CRITICAL: Use permission_request_id if available (for permission cards),
        # otherwise fall back to tool_call_id (for general questions)
        request_id = question_payload.get("permission_request_id", question_payload.get("id", str(_uuid.uuid4())))
        info = {
            "id": request_id,
            "text": question_payload["text"],
            "type": question_payload.get("type", "text"),
            "choices": question_payload.get("choices", []),
            "default": question_payload.get("default", ""),
            "details": question_payload.get("details", ""),
            "scope": question_payload.get("scope", "user"),
            "tool_name": question_payload.get("tool_name", "AskUserQuestion")
        }
        self._ai_chat.show_question(info)

    def _open_file_at_line(self, file_path: str, line_number: int):
        """Open file and navigate to specific line."""
        # First open the file
        self._open_file(file_path)
        
        # Get the current editor
        editor = self._editor_tabs.currentWidget()
        if isinstance(editor, CodeEditor):
            # Navigate to line (0-indexed in Scintilla, but 1-indexed for users)
            line_index = max(0, line_number - 1)
            editor.setCursorPosition(line_index, 0)
            editor.ensureLineVisible(line_index)
            log.info(f"Navigated to line {line_number} in {file_path}")

    def _on_accept_file_edit(self, file_path: str):
        """Accept AI edit — the file is already written to disk, just acknowledge."""
        if file_path == "__ALL__":
            self._on_accept_all_files()
            return
        file_path = os.path.normpath(file_path)
        log.info(f"[Accept] User accepted edit: {file_path}")

        # Open/refresh the file in the editor so the user sees the accepted state
        self._open_file(file_path)
        self.statusBar().showMessage(f"✓ Accepted changes to {Path(file_path).name}", 3000)

        # Clean up tracking state
        if hasattr(self, '_diff_data_store'):
            norm = os.path.normcase(file_path)
            self._diff_data_store.pop(file_path, None)
            self._diff_data_store.pop(norm, None)
        if hasattr(self, '_file_snapshots'):
            self._file_snapshots.pop(file_path, None)

        self._sidebar.remove_changed_file(file_path)

    def _on_reject_file_edit(self, file_path: str):
        """Reject AI edit — write original content back to disk and reload editor."""
        if file_path == "__ALL__":
            self._on_reject_all_files()
            return
        file_path = os.path.normpath(file_path)
        log.info(f"[Reject] User rejected edit: {file_path}")

        original = self._get_original_content(file_path)

        if original is not None:
            try:
                Path(file_path).write_text(original, encoding='utf-8')
                log.info(f"[Reject] Reverted {file_path} ({len(original)} chars)")
            except Exception as e:
                log.error(f"[Reject] Failed to revert {file_path}: {e}")
                self.statusBar().showMessage(f"✗ Revert failed for {Path(file_path).name}: {e}", 5000)
                return

            # Reload in editor so editor shows the reverted content
            self._open_file(file_path)
            self.statusBar().showMessage(f"↩ Reverted {Path(file_path).name} to original", 3000)
        else:
            # No original found — just open the file so user can review manually
            log.warning(f"[Reject] No original content for {file_path} — opening for review")
            self._open_file(file_path)
            self.statusBar().showMessage(
                f"⚠ No original content found for {Path(file_path).name} — review manually", 5000
            )

        # Clean up tracking state
        if hasattr(self, '_diff_data_store'):
            norm = os.path.normcase(file_path)
            self._diff_data_store.pop(file_path, None)
            self._diff_data_store.pop(norm, None)
        if hasattr(self, '_file_snapshots'):
            self._file_snapshots.pop(file_path, None)

        self._sidebar.remove_changed_file(file_path)

    def _get_original_content(self, file_path: str) -> Optional[str]:
        """
        Look up the original (pre-AI-edit) content for a file.
        Priority: _diff_data_store → _file_snapshots → diff_cache.json
        """
        norm = os.path.normcase(os.path.normpath(file_path))

        # 1. In-memory diff store (most reliable — set by _on_file_edited_diff_for_js)
        if hasattr(self, '_diff_data_store'):
            for key, (original, _modified) in self._diff_data_store.items():
                if os.path.normcase(os.path.normpath(key)) == norm:
                    if original:  # empty string means new file
                        return original

        # 2. File snapshot taken when file was first opened
        if hasattr(self, '_file_snapshots'):
            for key, content in self._file_snapshots.items():
                if os.path.normcase(os.path.normpath(key)) == norm:
                    return content

        # 3. Persisted diff cache
        try:
            cache = self._load_diff_cache()
            for key, entry in cache.items():
                if os.path.normcase(os.path.normpath(key)) == norm:
                    orig = entry.get('original', '')
                    if orig:
                        return orig
        except Exception:
            pass

        return None
    
    def _on_load_full_chat_requested(self, conversation_id: str):
        """Load full chat messages from SQLite database."""
        log.info(f"Loading full chat: {conversation_id}")
        
        try:
            # Load from SQLite via bridge
            full_chat_json_str = self._ai_chat.load_full_chat_from_sqlite(conversation_id)
            
            if full_chat_json_str and full_chat_json_str != "[]":
                log.info(f"Loaded {len(full_chat_json_str)} chars of chat data")
                # Send back to JavaScript
                self._ai_chat._view.page().runJavaScript(
                    f"window.handleFullChatLoad('{conversation_id}', {full_chat_json_str});"
                )
                self.statusBar().showMessage(f"✓ Loaded chat history", 2000)
            else:
                log.warning(f"No chat data found for: {conversation_id}")
                self._ai_chat._view.page().runJavaScript(
                    f"window.handleFullChatLoad('{conversation_id}', null);"
                )
        except Exception as e:
            log.error(f"Failed to load full chat: {e}")
            self._ai_chat._view.page().runJavaScript(
                f"window.handleFullChatLoad('{conversation_id}, null);"
            )

    def _on_accept_all_files(self):
        """Accept all pending AI edits — files already on disk, just clean up state."""
        log.info("[Accept All] User accepted all file edits")
        # Deduplicate keys — _diff_data_store has both raw and normcase keys per file
        seen = set()
        for file_path in list(getattr(self, '_diff_data_store', {}).keys()):
            try:
                norm = os.path.normcase(os.path.normpath(file_path))
                if norm in seen:
                    continue
                seen.add(norm)
                if os.path.isfile(norm):
                    self._open_file(norm)
            except Exception:
                pass
        if hasattr(self, '_diff_data_store'):
            self._diff_data_store.clear()
        if hasattr(self, '_file_snapshots'):
            self._file_snapshots.clear()
        self.statusBar().showMessage("✓ Accepted all changes", 3000)
        self._sidebar.clear_changed_files()

    def _on_reject_all_files(self):
        """Reject all pending AI edits — revert each file to its original content."""
        log.info("[Reject All] User rejected all file edits")
        reverted, failed = 0, 0
        # Deduplicate keys — _diff_data_store has both raw and normcase keys per file
        seen = set()
        for file_path in list(getattr(self, '_diff_data_store', {}).keys()):
            try:
                norm = os.path.normcase(os.path.normpath(file_path))
                if norm in seen:
                    continue
                seen.add(norm)
                original = self._get_original_content(norm)
                if original is not None and os.path.isfile(norm):
                    Path(norm).write_text(original, encoding='utf-8')
                    self._open_file(norm)
                    reverted += 1
                    log.info(f"[Reject All] Reverted {norm}")
                else:
                    failed += 1
            except Exception as e:
                log.error(f"[Reject All] Failed to revert {file_path}: {e}")
                failed += 1
        if hasattr(self, '_diff_data_store'):
            self._diff_data_store.clear()
        if hasattr(self, '_file_snapshots'):
            self._file_snapshots.clear()
        msg = f"↩ Reverted {reverted} file(s)"
        if failed:
            msg += f" ({failed} could not be reverted)"
        self.statusBar().showMessage(msg, 5000)
        self._sidebar.clear_changed_files()

    def _on_code_copied(self, text: str, file_path: str, start_line: int, end_line: int):
        """Store copy metadata so smart paste can use it after focus changes."""
        self._last_copy_info = {
            'text': text,
            'file_path': file_path,
            'start_line': start_line,
            'end_line': end_line,
        }

    def _on_smart_paste_check(self, pasted_text: str):
        """Check if pasted text matches last editor copy. Send result to chat."""
        try:
            import json as _json

            # Use stored copy metadata (captured at Ctrl+C time, before focus changed)
            copy_info = getattr(self, '_last_copy_info', None)
            if not copy_info or not copy_info.get('text'):
                self._ai_chat._view.page().runJavaScript(
                    "handleSmartPasteResult({isMatch: false});"
                )
                return

            def normalize(text):
                return '\n'.join(line.strip() for line in text.strip().split('\n') if line.strip())

            pasted_norm = normalize(pasted_text)
            copied_norm = normalize(copy_info['text'])

            if not copied_norm or not pasted_norm:
                self._ai_chat._view.page().runJavaScript(
                    "handleSmartPasteResult({isMatch: false});"
                )
                return

            is_match = (pasted_norm == copied_norm or
                        pasted_norm in copied_norm or
                        copied_norm in pasted_norm)

            if is_match and copy_info.get('file_path'):
                file_path  = copy_info['file_path']
                start_line = copy_info['start_line']
                end_line   = copy_info['end_line']
                file_name  = os.path.basename(file_path)
                ext        = os.path.splitext(file_path)[1].lstrip('.')
                line_range = str(start_line) if start_line == end_line else f"{start_line}-{end_line}"

                self._ai_chat._view.page().runJavaScript(
                    f"handleSmartPasteResult({{isMatch: true, "
                    f"filePath: {_json.dumps(file_path)}, "
                    f"fileName: {_json.dumps(file_name)}, "
                    f"lineRange: {_json.dumps(line_range)}, "
                    f"code: {_json.dumps(pasted_text)}, "
                    f"language: {_json.dumps(ext)}}});"
                )
                log.info(f"Smart paste matched: {file_name} lines {line_range}")
                # Clear after use so next paste starts fresh
                self._last_copy_info = None
                return

            self._ai_chat._view.page().runJavaScript(
                "handleSmartPasteResult({isMatch: false});"
            )

        except Exception as e:
            log.error(f"Smart paste check error: {e}")
            self._ai_chat._view.page().runJavaScript(
                "handleSmartPasteResult({isMatch: false});"
            )

    # ============================================================================
    # NEW: OpenCode Enhancement Integration Handlers
    # ============================================================================
    
    def _on_intent_classified(self, message: str, intent: str, confidence: float):
        """Handle intent classification from AI Integration Layer."""
        log.info(f"[Intent] {intent} (confidence: {confidence:.2f}): {message[:50]}...")
        # Could update UI to show detected intent
        
    def _on_agent_selected(self, agent_type: str, reason: str, confidence: float):
        """Handle agent selection from AI Integration Layer."""
        log.info(f"[Agent] Selected {agent_type} (confidence: {confidence:.2f}): {reason}")
        # Could show agent indicator in UI
        
    def _on_tools_selected(self, tool_names: list):
        """Handle tool selection from AI Integration Layer."""
        log.info(f"[Tools] Selected: {', '.join(tool_names)}")
        # Could show tool indicators in UI
        
    def _on_permission_requested(self, request_id: str, html_card: str):
        """Handle permission request - show permission card in chat."""
        log.info(f"[Permission] Request {request_id} - showing permission card")
        
        # Show permission card in AI chat
        if hasattr(self._ai_chat, 'show_permission_card'):
            self._ai_chat.show_permission_card(request_id, html_card)
        else:
            # Fallback: add as system message
            self._ai_chat.add_system_message("🔒 Permission required. Please check the chat interface.")
            
    def _on_permission_granted(self, request_id: str, scope: str):
        """Handle permission grant."""
        log.info(f"[Permission] Granted {request_id} with scope {scope}")
        
        # Retry the AI processing now that permission is granted
        # Get the last user message and retry
        if hasattr(self._ai_chat, '_last_user_message'):
            message = self._ai_chat._last_user_message
            context = []
            
            if self._project_manager.root:
                context.append(f"Project path: {self._project_manager.root}")
            
            editor = self._editor_tabs.current_editor()
            if editor:
                fp = self._editor_tabs.current_filepath()
                if fp:
                    name = Path(fp).name
                    content = editor.get_all_text()
                    if len(content) > 5000:
                        content = content[:5000] + "... (truncated)"
                    context.append(f"Current file ({name}):\n```\n{content}\n```")
            
            full_context = "\n\n".join(context)
            
            # NEW: Check if we have enhancement data stored
            enhancement_data = self._ai_agent.get_last_enhancement_data()
            if enhancement_data.get("intent"):
                log.info("Retrying with chat_with_enhancement after permission grant")
                self._ai_agent.chat_with_enhancement(
                    message,
                    intent=enhancement_data["intent"],
                    route=enhancement_data["route"],
                    tools=enhancement_data["tools"],
                    code_context=full_context
                )
            else:
                self._ai_agent.chat(message, full_context)
            
    def _on_permission_denied(self, request_id: str, reason: str):
        """Handle permission denial."""
        log.info(f"[Permission] Denied {request_id}: {reason}")
        self._ai_chat.add_system_message(f"❌ Permission denied: {reason}")
        
    def _on_chat_permission_response(self, request_id: str, approved: bool, scope: str = "session", remember: bool = False):
        """Handle permission response from chat UI.
        
        NOTE: This method is now DEPRECATED. The active permission flow uses
        permission_decided signal connected directly to bridge.on_permission_respond().
        This handler is kept for backwards compatibility but the permission_response
        signal is no longer connected.
        """
        log.info(f"[DEPRECATED] _on_chat_permission_response called: {request_id}, approved={approved}")
        # The real permission flow now goes through:
        #   _ai_chat.permission_decided -> _ai_agent.on_permission_respond(decision)
        # This method is no longer in the active path.
    
    def _on_user_denied_workflow(self, tool_name: str):
        """Handle user denying workflow twice - stop AI agent."""
        log.warning(f"User denied {tool_name} twice - stopping AI agent")
        
        # Stop the AI agent immediately
        self._ai_agent.stop()
        
        # Add system message explaining what happened
        self._ai_chat.add_system_message(
            f"⏹️ **Workflow Stopped**\n\n"
            f"You denied `{tool_name}` twice. The AI agent has stopped its current work.\n\n"
            f"If you'd like to continue with a different approach, please send a new message."
        )
        
        # Hide thinking indicator
        self._ai_chat.hide_thinking()
        
        # Reset UI - show send button again
        self._view.page().runJavaScript("if(window._onGenerationComplete) window._onGenerationComplete();")

    # ========== TESTING WORKFLOW HANDLERS (NEW) ==========
    
    def _on_testing_decision(self, decision: str, priority: str, trigger: str):
        """Handle testing decision signal."""
        log.info(f"[Testing] Decision: {decision} (priority: {priority}, trigger: {trigger})")
        
        # Show UI notification based on decision
        if decision == 'write_tests':
            self._ai_chat.add_system_message(
                f"🧪 **Testing Mode Activated**\n"
                f"Priority: {priority.upper()} | Trigger: {trigger}\n"
                f"The AI will analyze your code and suggest appropriate tests."
            )
        elif decision == 'skip_tests':
            log.debug("Testing skipped - no triggers detected")
    
    def _on_test_tools_selected(self, tools: list):
        """Handle test tools selection signal."""
        log.info(f"[Testing] Tools selected: {tools}")
        
        if tools:
            self._ai_chat.add_system_message(
                f"🔧 **Test Framework:** {', '.join(tools)}"
            )
    
    def _on_test_execution_started(self, test_type: str):
        """Handle test execution start signal."""
        log.info(f"[Testing] Execution started: {test_type}")
        self._ai_chat.add_system_message(f"▶️ Running {test_type} tests...")
    
    def _on_test_execution_completed(self, all_passed: bool, passed_count: int, failed_count: int):
        """Handle test execution completion signal."""
        log.info(f"[Testing] Execution completed: {passed_count} passed, {failed_count} failed")
        
        if all_passed:
            self._ai_chat.add_system_message(
                f"✅ **All Tests Passed!** ({passed_count} tests)"
            )
        else:
            self._ai_chat.add_system_message(
                f"⚠️ **Tests Completed:** {passed_count} passed, {failed_count} failed"
            )
    
    def _on_test_analysis_ready(self, analysis: dict):
        """Handle test analysis results signal."""
        log.info(f"[Testing] Analysis ready: {analysis.get('all_passed', False)}")
        
        # Display failure patterns if any
        patterns = analysis.get('patterns', [])
        if patterns:
            pattern_text = '\n'.join([f"- {p.get('type', 'unknown')}: {p.get('description', '')}" 
                                     for p in patterns[:3]])
            self._ai_chat.add_system_message(
                f"📊 **Failure Analysis:**\n{pattern_text}"
            )

    def _open_file_at_line_duplicate(self, filepath: str, line: int):
        self._open_file(filepath)
        editor = self._editor_tabs.current_editor()
        if editor:
            # Move cursor to line
            cursor = editor.textCursor()
            block = editor.document().findBlockByLineNumber(line - 1)
            cursor.setPosition(block.position())
            editor.setTextCursor(cursor)
            editor.centerCursor()

    def _save_current(self):
        ok = self._editor_tabs.save_current(self._file_manager)
        if ok:
            self._status_file.setText(f"  Saved ✓  {self._status_file.text().strip()}")
            QTimer.singleShot(2000, lambda: self._update_status_file(
                self._editor_tabs.current_filepath()))

    def _save_all(self):
        for i in range(self._editor_tabs.count()):
            editor = self._editor_tabs.widget(i)
            fp = self._editor_tabs._files.get(i)
            if isinstance(editor, CodeEditor) and fp:
                content = editor.get_all_text()
                self._file_manager.write(fp, content)

    def _run_file(self):
        fp = self._editor_tabs.current_filepath()
        if not fp or not Path(fp).exists():
            return
        # HTML files → built-in Live Server (no terminal needed)
        if Path(fp).suffix.lower() in {".html", ".htm"}:
            self._save_current()
            self._run_live_server(fp)
            return
        self._save_current()
        
        self._terminal_tabs.setVisible(True)
        term = self._current_terminal()
        if not term:
            term = self._new_terminal()
        
        term.setFocus()
        lang = detect_language(fp)
        command = self._build_run_command(fp, lang)
        if command:
            term.execute_command(command)
        else:
            QMessageBox.information(self, "Run", f"Running {lang} is not yet supported.")

    def _run_live_server(self, file_path: str):
        """Live Server removed in AI-first mode - AI handles preview"""
        QMessageBox.information(self, "Info", 
            "Live Server is not available in AI-first mode.\n"
            "Ask AI to preview your HTML code instead.")
        return

        import webbrowser
        webbrowser.open(url)

        # Show status in status bar
        if hasattr(self, '_statusbar_label'):
            self._statusbar_label.setText(
                f"Live Server  \u25cf  http://localhost:{port}   —   click \u25b6 to restart"
            )

    def _build_run_command(self, file_path: str, lang: str) -> str | None:
        """Build a run command for the current file based on language."""
        is_windows = platform.system() == "Windows"
        root = self._project_manager.root or str(Path(file_path).parent)
        build_dir = os.path.join(root, ".cortex_build")
        stem = Path(file_path).stem

        if is_windows:
            quote = lambda p: f'"{p}"'
            mkdir_cmd = f'New-Item -ItemType Directory -Force -Path {quote(build_dir)} | Out-Null'
        else:
            import shlex
            quote = shlex.quote
            mkdir_cmd = f'mkdir -p {quote(build_dir)}'

        if lang == "python":
            return f'python {quote(file_path)}'
        if lang in {"javascript", "jsx"}:
            return f'node {quote(file_path)}'
        if lang in {"typescript", "tsx"}:
            # Compile with tsc and run with node (faster than npx ts-node)
            js_out = os.path.join(build_dir, stem + ".js")
            if is_windows:
                return f'tsc {quote(file_path)} --outDir {quote(build_dir)}; if ($LASTEXITCODE -eq 0) {{ node {quote(js_out)} }}'
            else:
                return f'tsc {quote(file_path)} --outDir {quote(build_dir)} && node {quote(js_out)}'
        if lang == "bash":
            return f'bash {quote(file_path)}'
        if lang == "batch":
            return f'& {quote(file_path)}' if is_windows else None
        if lang == "powershell":
            return f'& {quote(file_path)}' if is_windows else f'pwsh {quote(file_path)}'
        if lang == "ruby":
            return f'ruby {quote(file_path)}'
        if lang == "php":
            return f'php {quote(file_path)}'
        if lang == "go":
            return f'go run {quote(file_path)}'
        if lang == "rust":
            exe_path = os.path.join(build_dir, stem + (".exe" if is_windows else ""))
            compile_cmd = f'rustc {quote(file_path)} -o {quote(exe_path)}'
            run_cmd = f'& {quote(exe_path)}' if is_windows else f'{quote(exe_path)}'
            return f'{mkdir_cmd}; {compile_cmd}; if ($LASTEXITCODE -eq 0) {{ {run_cmd} }}' if is_windows else f'{mkdir_cmd} && {compile_cmd} && {run_cmd}'
        if lang == "c":
            exe_path = os.path.join(build_dir, stem + (".exe" if is_windows else ""))
            compile_cmd = f'gcc {quote(file_path)} -o {quote(exe_path)}'
            run_cmd = f'& {quote(exe_path)}' if is_windows else f'{quote(exe_path)}'
            return f'{mkdir_cmd}; {compile_cmd}; if ($LASTEXITCODE -eq 0) {{ {run_cmd} }}' if is_windows else f'{mkdir_cmd} && {compile_cmd} && {run_cmd}'
        if lang == "cpp":
            exe_path = os.path.join(build_dir, stem + (".exe" if is_windows else ""))
            compile_cmd = f'g++ {quote(file_path)} -o {quote(exe_path)}'
            run_cmd = f'& {quote(exe_path)}' if is_windows else f'{quote(exe_path)}'
            return f'{mkdir_cmd}; {compile_cmd}; if ($LASTEXITCODE -eq 0) {{ {run_cmd} }}' if is_windows else f'{mkdir_cmd} && {compile_cmd} && {run_cmd}'
        if lang == "java":
            package_name = ""
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    for line in f:
                        line = line.strip()
                        if line.startswith("package ") and line.endswith(";"):
                            package_name = line[len("package "):-1].strip()
                            break
            except Exception:
                package_name = ""
            class_name = f"{package_name}.{stem}" if package_name else stem
            compile_cmd = f'javac -d {quote(build_dir)} {quote(file_path)}'
            run_cmd = f'java -cp {quote(build_dir)} {class_name}'
            return f'{mkdir_cmd}; {compile_cmd}; if ($LASTEXITCODE -eq 0) {{ {run_cmd} }}' if is_windows else f'{mkdir_cmd} && {compile_cmd} && {run_cmd}'
        return None

    def _new_terminal(self, show_panel: bool = True) -> XTermWidget:
        # If call comes from a signal (like clicked), show_panel might be the 'checked' state (False usually)
        # So we force it to True if it's not explicitly False from our internal calls
        if not isinstance(show_panel, bool): show_panel = True
        
        term = XTermWidget()
        term.set_theme(self._theme_manager.is_dark)
        
        # Initialize with current project directory if available
        if self._project_manager.root:
            term.set_cwd(str(self._project_manager.root))
            
        idx = self._terminal_tabs.addTab(term, f"Terminal {self._terminal_tabs.count() + 1}")
        self._terminal_tabs.setCurrentIndex(idx)
        
        if show_panel:
            self._terminal_tabs.setVisible(True)
            term.setFocus()
            
        # Hook up "New Terminal" button from within the terminal
        term.new_terminal_requested.connect(lambda: self._new_terminal(show_panel=True))
        
       
       
        # Connect file operations to AI chat
        term.file_operation_detected.connect(self._on_terminal_file_operation)
        
        return term
        
    def _on_terminal_file_operation(self, operation_type: str, file_path: str, status: str):
        """Handle file operations from terminal and show in AI chat."""
        # Map operation types to display format
        op_labels = {
            'create': 'Creating file',
            'create_dir': 'Creating directory',
            'delete': 'Deleting file',
            'delete_dir': 'Deleting directory',
            'move': 'Moving file',
            'copy': 'Copying file',
            'rename': 'Renaming file'
        }
        
        label = op_labels.get(operation_type, operation_type)
        display_info = f"{label}: {file_path}"
        
        # Show in AI chat as tool activity
        self._ai_chat.show_tool_activity('terminal_' + operation_type, display_info, status)

    def _on_terminal_tab_changed(self, index: int):
        """Update AI agent when active terminal changes."""
        term = self._terminal_tabs.widget(index)
        if isinstance(term, XTermWidget):
            self._ai_agent.set_terminal(term)

    def _close_terminal_tab(self, index: int):
        term = self._terminal_tabs.widget(index)
        if isinstance(term, XTermWidget):
            term._kill_process()
        self._terminal_tabs.removeTab(index)
        if self._terminal_tabs.count() == 0:
            self._terminal_tabs.setVisible(False)

    def _kill_current_terminal(self):
        idx = self._terminal_tabs.currentIndex()
        if idx >= 0:
            self._close_terminal_tab(idx)

    def _current_terminal(self) -> XTermWidget | None:
        w = self._terminal_tabs.currentWidget()
        return w if isinstance(w, XTermWidget) else None

    def _toggle_theme(self):
        new_theme = self._theme_manager.toggle(QApplication_instance())
        self._settings.theme = new_theme
        is_dark = self._theme_manager.is_dark
        self._theme_btn.setText("☀️" if not is_dark else "🌙")
        self._editor_tabs.update_theme(is_dark)
        if isinstance(self._terminal_tabs.tabBar(), CleanTabBar):
            self._terminal_tabs.tabBar().set_dark(is_dark)
        
        self._ai_chat.set_theme(is_dark)
        self._sidebar.set_theme(is_dark)
        if hasattr(self, '_diff_window'):
            self._diff_window.set_theme(is_dark)
        
        # Style the tab widget panels
        tab_bg = "#1e1e1e" if is_dark else "#ffffff"
        tab_border = "#3e3e42" if is_dark else "#dee2e6"
        tab_style = f"""
            QTabWidget::pane {{ border: 1px solid {tab_border}; background: {tab_bg}; }}
            QTabWidget::tab-bar {{ left: 0px; }}
        """
        self._editor_tabs.setStyleSheet(tab_style)
        self._terminal_tabs.setStyleSheet(tab_style)

        # Update all terminal tabs
        for i in range(self._terminal_tabs.count()):
            term = self._terminal_tabs.widget(i)
            if isinstance(term, XTermWidget):
                term.set_theme(is_dark)
        self._apply_welcome_theme(is_dark)


    def _show_terminal_panel(self):
        """Show terminal panel (called from AI chat 'View in terminal' button)."""
        self._terminal_tabs.setVisible(True)
        if self._terminal_tabs.count() == 0:
            self._new_terminal()
        term = self._current_terminal()
        if term:
            term.setFocus()

    def _show_terminal_and_run(self, command: str):
        """Show terminal panel and execute command (called from 'View in terminal' with a command)."""
        self._terminal_tabs.setVisible(True)
        if self._terminal_tabs.count() == 0:
            self._new_terminal()
        # Ensure the terminal panel has a visible height
        self._terminal_tabs.setMinimumHeight(150)
        term = self._current_terminal()
        if term:
            term.setFocus()
            if command and command.strip():
                term.execute_command(command.strip())

    def _toggle_terminal(self):
        """Toggle terminal visibility in AI-first mode"""
        visible = self._terminal_tabs.isVisible()
        self._terminal_tabs.setVisible(not visible)
        if self._terminal_tabs.isVisible():
            if self._terminal_tabs.count() == 0:
                self._new_terminal()
            self._terminal_tabs.setMinimumHeight(150)
            term = self._current_terminal()
            if term:
                term.setFocus()

    def _toggle_left_sidebar(self, show: bool = True):
        """Toggle left sidebar via splitter — show=True to expand, False to collapse."""
        sizes = self._main_splitter.sizes()
        if len(sizes) < 4:
            return
        widget = self._main_splitter.widget(0)
        if show:
            widget.setMinimumWidth(180)
            widget.setMaximumWidth(400)
            sizes[0] = self._left_sidebar_min_width
        else:
            widget.setMinimumWidth(0)
            widget.setMaximumWidth(0)
            sizes[0] = 0
        self._main_splitter.setSizes(sizes)
        self._left_sidebar_hidden = not show

    def _toggle_sidebar(self):
        """Toggle sidebar visibility (hidden by default in AI-first mode)"""
        visible = self._sidebar.isVisible()
        self._sidebar.setVisible(not visible)
        if self._sidebar.isVisible():
            self._sidebar.setFocus()

    def _toggle_file_tree_panel(self, show: bool = True):
        """Toggle file tree panel via splitter (Ctrl+Shift+I)."""
        sizes = self._main_splitter.sizes()
        if len(sizes) < 4:
            return
        widget = self._main_splitter.widget(3)
        if show:
            widget.setMinimumWidth(200)
            widget.setMaximumWidth(500)
            sizes[3] = self._file_tree_min_width
        else:
            widget.setMinimumWidth(0)
            widget.setMaximumWidth(0)
            sizes[3] = 0
        self._main_splitter.setSizes(sizes)
        self._file_tree_hidden = not show

    def _toggle_file_tree(self):
        """Toggle File Tree panel visibility (Ctrl+Shift+I)."""
        if hasattr(self, '_file_tree_panel'):
            visible = self._file_tree_panel.isVisible()
            self._file_tree_panel.setVisible(not visible)
            self._file_tree_hidden = not visible

    def _toggle_ai_chat_panel(self, show: bool = True):
        """Toggle AI chat panel via splitter."""
        sizes = self._main_splitter.sizes()
        if len(sizes) < 4:
            return
        widget = self._main_splitter.widget(1)
        if show:
            widget.setMinimumWidth(300)
            widget.setMaximumWidth(16777215)  # QWIDGETSIZE_MAX
            sizes[1] = self._chat_panel_min_width
        else:
            widget.setMinimumWidth(0)
            widget.setMaximumWidth(0)
            sizes[1] = 0
        self._main_splitter.setSizes(sizes)
        self._chat_panel_hidden = not show

    def _toggle_review_panel(self, show: bool = True):
        """Toggle Review panel via splitter (Alt+Ctrl+B)."""
        sizes = self._main_splitter.sizes()
        if len(sizes) < 4:
            return
        widget = self._main_splitter.widget(2)
        if show:
            widget.setMinimumWidth(250)
            widget.setMaximumWidth(600)
            sizes[2] = self._review_panel_min_width
        else:
            widget.setMinimumWidth(0)
            widget.setMaximumWidth(0)
            sizes[2] = 0
        self._main_splitter.setSizes(sizes)
        self._review_panel_hidden = not show

    def _toggle_review_panel_menu(self):
        """Toggle Review panel visibility (Alt+Ctrl+B) — menu bar action."""
        if hasattr(self, '_review_panel'):
            visible = self._review_panel.isVisible()
            self._review_panel.setVisible(not visible)
            self._review_panel_hidden = not visible

    def _toggle_summary_panel(self, show: bool = True):
        """Toggle summary panel visibility (switch between Summary and Review tabs)."""
        if show:
            if hasattr(self, '_summary_tab_btn'):
                self._summary_tab_btn.setChecked(True)
                self._review_stack.setCurrentIndex(0)
        else:
            if hasattr(self, '_review_tab_btn'):
                self._review_tab_btn.setChecked(True)
                self._review_stack.setCurrentIndex(1)
        self._summary_panel_hidden = not show

    def _toggle_git_panel(self, show: bool = True):
        """Toggle git panel visibility — implemented as a panel-level toggle."""
        if show:
            # Expand review panel to show git content
            if hasattr(self, '_review_panel'):
                self._review_panel.setVisible(True)
                self._review_panel_hidden = False
        else:
            if hasattr(self, '_review_panel'):
                self._review_panel.setVisible(False)
                self._review_panel_hidden = True
        self._git_panel_hidden = not show

    def _toggle_fullscreen(self):
        """Toggle full screen mode (F11)."""
        if self.isFullScreen():
            self.showNormal()
        else:
            self.showFullScreen()

    def _minimize_window(self):
        """Minimize window (Ctrl+M)."""
        self.showMinimized()

    def _zoom_window(self):
        """Zoom window (maximize/restore)."""
        if self.isMaximized():
            self.showNormal()
        else:
            self.showMaximized()

    def _close_window(self):
        """Close window (Ctrl+W)."""
        self.close()

    def _previous_chat(self):
        """Navigate to previous chat (Ctrl+Shift+[)."""
        # TODO: Implement chat history navigation
        log.info("Previous chat shortcut pressed")

    def _next_chat(self):
        """Navigate to next chat (Ctrl+Shift+])."""
        # TODO: Implement chat history navigation
        log.info("Next chat shortcut pressed")

    def _navigate_back(self):
        """Navigate back (Ctrl+[)."""
        # TODO: Implement navigation history
        log.info("Navigate back shortcut pressed")

    def _navigate_forward(self):
        """Navigate forward (Ctrl+])."""
        # TODO: Implement navigation history
        log.info("Navigate forward shortcut pressed")

    def _new_window(self):
        """Open new window (Ctrl+Shift+N)."""
        # TODO: Implement multi-window support
        log.info("New window shortcut pressed")

    def _quick_chat(self):
        """Open quick chat (Alt+Ctrl+N)."""
        self._on_new_chat()

    def _toggle_ai_chat(self):
        """Toggle AI chat panel visibility"""
        # In AI-first mode, the chat is always visible
        # This method can be used to focus it instead
        if hasattr(self, '_ai_chat'):
            self._ai_chat.setFocus()

    def _zoom_in(self):
        """Zoom in (Ctrl+=)."""
        editor = self._editor_tabs.current_editor()
        if editor:
            zoom = editor.zoomIn() + 1
            editor.setZoom(zoom)

    def _zoom_out(self):
        """Zoom out (Ctrl+-)."""
        editor = self._editor_tabs.current_editor()
        if editor:
            zoom = max(0, editor.zoomIn() - 1)
            editor.setZoom(zoom)

    def _zoom_reset(self):
        """Reset zoom (Ctrl+0)."""
        editor = self._editor_tabs.current_editor()
        if editor:
            editor.setZoom(0)

    def _focus_ai_chat(self):
        """Focus AI Chat input (Ctrl+Shift+A)."""
        self._ai_chat.focus_input()
        self._ai_chat.raise_()
        self._ai_chat.activateWindow()

    def _command_palette(self):
        """Show Command Palette (Ctrl+Shift+P) - REMOVED in AI-first mode."""
        # Command palette not implemented
        pass

    def _current_editor_action(self, action: str):
        """Focus-aware edit action handler (supports Editor, AI Chat, and Terminal)."""
        focused = QApplication.focusWidget()
        log.debug(f"Action {action} requested. Current focus: {focused}")

        # Map generic action strings to QWebEnginePage.WebAction enums
        web_action_map = {
            "copy": QWebEnginePage.WebAction.Copy,
            "paste": QWebEnginePage.WebAction.Paste,
            "cut": QWebEnginePage.WebAction.Cut,
            "selectAll": QWebEnginePage.WebAction.SelectAll,
            "undo": QWebEnginePage.WebAction.Undo,
            "redo": QWebEnginePage.WebAction.Redo
        }

        # Determine which "logical" component has focus
        logical_focused = None
        widget = focused
        max_depth = 10
        while widget and max_depth > 0:
            if hasattr(self, '_ai_chat') and (widget == self._ai_chat or widget == self._ai_chat._view):
                logical_focused = "ai_chat"
                break
            
            # Check if this widget belongs to a terminal tab
            term = self._current_terminal()
            if term and (widget == term or widget == term._webview):
                logical_focused = "terminal"
                break
            
            widget = widget.parentWidget()
            max_depth -= 1

        # 1. Route to AI Chat
        if logical_focused == "ai_chat":
            if action in web_action_map:
                log.debug(f"Routing {action} to AI Chat WebEngineView")
                self._ai_chat._view.page().triggerAction(web_action_map[action])
                return

        # 2. Route to Terminal
        if logical_focused == "terminal":
            term = self._current_terminal()
            if action == "copy":
                term.copy()
                return
            elif action == "paste":
                term.paste()
                return
            elif action == "selectAll":
                term.select_all()
                return
            elif action == "cut":
                term.cut()
                return
            
            if action in web_action_map:
                term._webview.page().triggerAction(web_action_map[action])
                return

        # 3. Route to Sidebar explicitly (if focused)
        if hasattr(self, '_sidebar') and self._sidebar.is_explorer_focused():
            log.debug(f"Action {action} ignored globally: Sidebar handles it locally")
            return

        # 4. Fallback to Editor (current tab)
        editor = self._editor_tabs.current_editor()
        if editor:
            log.debug(f"Routing {action} to Code Editor")
            if action == "selectAll":
                if hasattr(editor, "selectAll"): editor.selectAll()
                elif hasattr(editor, "select_all"): editor.select_all()
                return
            if hasattr(editor, action):
                getattr(editor, action)()

    # ------------------------------------------------------------------
    # VS Code Style Keyboard Shortcuts
    # ------------------------------------------------------------------
    def _show_find(self):
        """Show Find dialog (Ctrl+F)."""
        editor = self._editor_tabs.current_editor()
        if editor:
            selected = editor.get_selected_text()
            if selected:
                self._find_replace_dialog.set_find_text(selected)
            self._find_replace_dialog.show_find_only()
            self._find_replace_dialog.show()
            self._find_replace_dialog.raise_()
            self._find_replace_dialog.activateWindow()

    def _show_find_replace(self):
        """Show Find & Replace dialog (Ctrl+H)."""
        editor = self._editor_tabs.current_editor()
        if editor:
            selected = editor.get_selected_text()
            if selected:
                self._find_replace_dialog.set_find_text(selected)
            self._find_replace_dialog.show_find_replace()
            self._find_replace_dialog.show()
            self._find_replace_dialog.raise_()
            self._find_replace_dialog.activateWindow()

    def _rename_file(self):
        """Rename file (F2)."""
        # Check if right-side Explore tree is focused
        if hasattr(self, '_file_tree') and self._file_tree.hasFocus():
            index = self._file_tree.currentIndex()
            if index.isValid():
                file_path = self._file_model.filePath(index)
                if file_path:
                    # Prevent renaming the project root
                    if hasattr(self, '_project_manager') and self._project_manager.root:
                        try:
                            if Path(file_path).resolve() == Path(str(self._project_manager.root)).resolve():
                                return
                        except Exception:
                            pass
                    self._rename_path(file_path)
                    return
        
        # Check if left sidebar explorer is focused
        if self._sidebar.is_explorer_focused():
            if self._sidebar.rename_selected_item():
                return
            return

        # Otherwise rename the currently open file in editor
        current_file = self._editor_tabs.current_filepath()
        if not current_file:
            return
        
        from PyQt6.QtWidgets import QInputDialog
        from pathlib import Path
        
        old_name = Path(current_file).name
        new_name, ok = QInputDialog.getText(
            self, 
            "Rename File", 
            f"New name for '{old_name}':",
            text=old_name
        )
        
        if ok and new_name and new_name != old_name:
            try:
                old_path = Path(current_file)
                new_path = old_path.parent / new_name
                
                # Rename file on disk
                old_path.rename(new_path)
                
                # Close current tab and open renamed file
                index = self._editor_tabs.currentIndex()
                self._editor_tabs.removeTab(index)
                self._open_file(str(new_path))
                
                self._status_bar.showMessage(f"Renamed to {new_name}", 3000)
            except Exception as e:
                from PyQt6.QtWidgets import QMessageBox
                QMessageBox.critical(self, "Rename Failed", f"Could not rename file: {e}")

    def _rename_path(self, path: str) -> bool:
        """Rename a file or folder by path (used by Explore panel)."""
        try:
            name, ok = QInputDialog.getText(self, "Rename", "New name:", text=Path(path).name)
            if not ok or not name or name == Path(path).name:
                return False

            new_path_obj = Path(path).parent / name
            
            # Prevent overwriting existing files
            if new_path_obj.exists() and new_path_obj.resolve() != Path(path).resolve():
                QMessageBox.warning(
                    self, 
                    "Rename Failed", 
                    f"A file or folder with the name '{name}' already exists.\n\nPlease choose a different name."
                )
                return False

            new_path = str(new_path_obj)
            Path(path).rename(new_path)
            log.info(f"[RENAME] Renamed: {Path(path).name} -> {name}")
            
            # Refresh file tree
            QTimer.singleShot(100, lambda: self._file_tree.viewport().update())
            return True
        except Exception as e:
            QMessageBox.critical(self, "Rename Failed", f"Could not rename: {e}")
            return False

    def _on_sidebar_file_renamed(self, old_path: str, new_path: str):
        old_norm = os.path.normpath(old_path)
        new_norm = os.path.normpath(new_path)
        updated = False

        for idx, fp in list(self._editor_tabs._files.items()):
            if os.path.normcase(fp) != os.path.normcase(old_norm):
                continue

            self._editor_tabs._files[idx] = new_norm
            name = Path(new_norm).name
            if fp in self._editor_tabs._modified:
                self._editor_tabs._modified.discard(fp)
                self._editor_tabs._modified.add(new_norm)
                self._editor_tabs.setTabText(idx, f"ƒ-? {name}")
                self._editor_tabs.setTabToolTip(idx, f"{new_norm} (Modified)")
            else:
                self._editor_tabs.setTabText(idx, name)
                self._editor_tabs.setTabToolTip(idx, new_norm)

            self._editor_tabs._set_tab_icon(idx, new_norm)
            updated = True

            if idx == self._editor_tabs.currentIndex():
                self._update_status_file(new_norm)
                if hasattr(self, '_ai_agent'):
                    self._ai_agent.set_active_file(new_norm)

        if updated:
            log.info(f"Updated open tabs for rename: {old_norm} -> {new_norm}")
    
    def _on_sidebar_file_deleted(self, path: str):
        """Handle file/folder deletion from sidebar."""
        import os
        from pathlib import Path
        
        norm_path = os.path.normpath(path)
        
        # Close tab if file is open
        for idx, fp in list(self._editor_tabs._files.items()):
            if os.path.normcase(fp) == os.path.normcase(norm_path):
                self._editor_tabs._close_tab(idx)
                break
        
        # Refresh sidebar to reflect deletion
        self._sidebar.refresh()
        
        # Refresh project context for AI agent
        if hasattr(self, '_ai_agent') and self._ai_agent:
            project_root = str(self._project_manager.root) if self._project_manager.root else None
            if project_root:
                self._ai_agent.set_project_root(project_root)
        
        log.info(f"File deleted: {path}")
    
    def _on_file_deleted_for_undo(self, original_path: str):
        """Track file deletion for undo functionality."""
        log.debug(f"Undo tracking: File moved to trash: {original_path}")
    
    def _on_file_restored_for_redo(self, restored_path: str):
        """Track file restoration for redo functionality."""
        log.debug(f"Redo tracking: File restored: {restored_path}")
        # Refresh sidebar to show restored file
        QTimer.singleShot(100, self._sidebar.refresh)
    
    def _undo(self):
        """Handle undo - prioritize editor undo, then file restore."""
        # Try editor undo first
        editor = self._editor_tabs.current_editor()
        if editor and editor.document().isUndoAvailable():
            editor.undo()
            return
        
        # If no editor or no undo available, try file restore
        if hasattr(self, '_file_manager') and self._file_manager.can_undo():
            restored_path = self._file_manager.undo_operation()
            if restored_path:
                log.info(f"Restored file: {restored_path}")
                self.statusBar().showMessage(f"Restored: {Path(restored_path).name}", 3000)
    
    def _redo(self):
        """Handle redo - prioritize editor redo, then file re-delete."""
        # Try editor redo first
        editor = self._editor_tabs.current_editor()
        if editor and editor.document().isRedoAvailable():
            editor.redo()
            return
        
        if hasattr(self, '_file_manager') and self._file_manager.can_redo():
            deleted_path = self._file_manager.redo_operation()
            if deleted_path:
                log.info(f"Re-deleted file: {deleted_path}")
                self.statusBar().showMessage(f"Deleted: {Path(deleted_path).name}", 3000)

    def _find_in_files(self):
        """Find in files (Ctrl+Shift+F)."""
        self._ai_chat.add_system_message("🔍 Find in Files: Type your search query in the AI chat.")
        self._ai_chat.set_input_text("Search in files for: ")

    def _go_to_line(self):
        """Go to line (Ctrl+G)."""
        editor = self._editor_tabs.current_editor()
        if not editor:
            return
        
        from PyQt6.QtWidgets import QInputDialog
        line, ok = QInputDialog.getInt(self, "Go to Line", "Line number:", min=1, max=editor.blockCount())
        if ok:
            cursor = editor.textCursor()
            cursor.movePosition(cursor.MoveOperation.Start)
            for _ in range(line - 1):
                cursor.movePosition(cursor.MoveOperation.Down)
            editor.setTextCursor(cursor)
            editor.setFocus()

    def _toggle_comment(self):
        """Toggle comment on selected lines (Ctrl+/)."""
        editor = self._editor_tabs.current_editor()
        if not editor:
            return
        
        cursor = editor.textCursor()
        start_line = cursor.blockNumber()
        
        # Get selected text range
        if cursor.hasSelection():
            end_cursor = editor.textCursor()
            end_cursor.setPosition(cursor.selectionEnd())
            end_line = end_cursor.blockNumber()
        else:
            end_line = start_line
        
        # Toggle comments for each line
        for line_num in range(start_line, end_line + 1):
            block = editor.document().findBlockByNumber(line_num)
            text = block.text()
            
            if text.strip().startswith("# "):
                # Remove comment
                new_text = text.replace("# ", "", 1)
            elif text.strip().startswith("#"):
                # Remove comment
                new_text = text.replace("#", "", 1)
            else:
                # Add comment
                new_text = "# " + text
            
            cursor.select(cursor.SelectionType.BlockUnderCursor)
            cursor.insertText(new_text)

    def _delete_line(self):
        """Delete current line (Ctrl+Shift+K)."""
        editor = self._editor_tabs.current_editor()
        if not editor:
            return
        
        cursor = editor.textCursor()
        cursor.select(cursor.SelectionType.BlockUnderCursor)
        cursor.removeSelectedText()
        cursor.deleteChar()  # Remove the newline

    def _duplicate_line(self):
        """Duplicate current line (Ctrl+Shift+D)."""
        editor = self._editor_tabs.current_editor()
        if not editor:
            return
        
        cursor = editor.textCursor()
        line = cursor.block().text()
        cursor.movePosition(cursor.MoveOperation.EndOfLine)
        cursor.insertText("\n" + line)

    def _indent_line(self):
        """Indent selected lines (Ctrl+])."""
        editor = self._editor_tabs.current_editor()
        if not editor:
            return
        
        cursor = editor.textCursor()
        start = cursor.selectionStart()
        end = cursor.selectionEnd()
        
        cursor.beginEditBlock()
        cursor.setPosition(start)
        while cursor.position() <= end:
            cursor.movePosition(cursor.MoveOperation.StartOfLine)
            cursor.insertText("    ")  # 4 spaces
            cursor.movePosition(cursor.MoveOperation.Down)
            if cursor.atEnd():
                break
        cursor.endEditBlock()

    def _outdent_line(self):
        """Outdent selected lines (Ctrl+[)."""
        editor = self._editor_tabs.current_editor()
        if not editor:
            return
        
        cursor = editor.textCursor()
        start = cursor.selectionStart()
        end = cursor.selectionEnd()
        
        cursor.beginEditBlock()
        cursor.setPosition(start)
        while cursor.position() <= end:
            cursor.movePosition(cursor.MoveOperation.StartOfLine)
            line_text = cursor.block().text()
            if line_text.startswith("    "):
                cursor.deleteChar()
                cursor.deleteChar()
                cursor.deleteChar()
                cursor.deleteChar()
            elif line_text.startswith("\t"):
                cursor.deleteChar()
            cursor.movePosition(cursor.MoveOperation.Down)
            if cursor.atEnd():
                break
        cursor.endEditBlock()

    def _move_line_up(self):
        """Move current line up (Alt+Up)."""
        editor = self._editor_tabs.current_editor()
        if not editor:
            return
        
        cursor = editor.textCursor()
        line_num = cursor.blockNumber()
        
        if line_num == 0:
            return  # Already at top
        
        # Get current line and previous line
        current_block = cursor.block()
        prev_block = current_block.previous()
        
        current_text = current_block.text()
        prev_text = prev_block.text()
        
        # Swap lines
        cursor.beginEditBlock()
        cursor.movePosition(cursor.MoveOperation.Start)
        for _ in range(line_num - 1):
            cursor.movePosition(cursor.MoveOperation.Down)
        cursor.movePosition(cursor.MoveOperation.StartOfLine)
        cursor.movePosition(cursor.MoveOperation.Down, cursor.MoveMode.KeepAnchor)
        cursor.movePosition(cursor.MoveOperation.EndOfLine, cursor.MoveMode.KeepAnchor)
        cursor.insertText(current_text + "\n" + prev_text)
        cursor.endEditBlock()

    def _move_line_down(self):
        """Move current line down (Alt+Down)."""
        editor = self._editor_tabs.current_editor()
        if not editor:
            return
        
        cursor = editor.textCursor()
        line_num = cursor.blockNumber()
        total_lines = editor.blockCount()
        
        if line_num >= total_lines - 1:
            return  # Already at bottom
        
        # Get current line and next line
        current_block = cursor.block()
        next_block = current_block.next()
        
        current_text = current_block.text()
        next_text = next_block.text()
        
        # Swap lines
        cursor.beginEditBlock()
        cursor.movePosition(cursor.MoveOperation.StartOfLine)
        cursor.movePosition(cursor.MoveOperation.Down, cursor.MoveMode.KeepAnchor)
        cursor.movePosition(cursor.MoveOperation.EndOfLine, cursor.MoveMode.KeepAnchor)
        cursor.insertText(next_text + "\n" + current_text)
        cursor.endEditBlock()

    def _quick_open(self):
        """Quick open file (Ctrl+P) - Opens file dialog."""
        filepath, _ = QFileDialog.getOpenFileName(
            self, "Open File", 
            str(self._project_manager.root) if self._project_manager.root else "",
            "All Files (*.*)"
        )
        if filepath:
            self._open_file(filepath)

    def _go_to_symbol(self):
        """Go to symbol in file (Ctrl+Shift+O)."""
        editor = self._editor_tabs.current_editor()
        if not editor:
            return
        
        text = editor.get_all_text()
        # Find function/class definitions
        import re
        symbols = []
        for match in re.finditer(r'^(def|class|function|const|let|var)\s+(\w+)', text, re.MULTILINE):
            line_num = text[:match.start()].count('\n') + 1
            symbols.append(f"{match.group(1)} {match.group(2)} (line {line_num})")
        
        if symbols:
            self._ai_chat.add_system_message("📍 Symbols in file:\n" + "\n".join(symbols[:20]))
        else:
            self._ai_chat.add_system_message("No symbols found in current file.")

    def _close_current_tab(self):
        """Close current tab (Ctrl+W)."""
        self._editor_tabs.close_current_tab()

    def _close_all_tabs(self):
        """Close all tabs (Ctrl+Shift+W)."""
        self._editor_tabs.close_all_tabs()

    def _next_tab(self):
        """Go to next tab (Ctrl+Tab)."""
        current = self._editor_tabs.currentIndex()
        count = self._editor_tabs.count()
        if count > 0:
            self._editor_tabs.setCurrentIndex((current + 1) % count)

    def _prev_tab(self):
        """Go to previous tab (Ctrl+Shift+Tab)."""
        current = self._editor_tabs.currentIndex()
        count = self._editor_tabs.count()
        if count > 0:
            self._editor_tabs.setCurrentIndex((current - 1) % count)

    def _show_shortcuts_help(self):
        """Show keyboard shortcuts help dialog."""
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QLabel, QScrollArea
        
        dialog = QDialog(self)
        dialog.setWindowTitle("Keyboard Shortcuts Reference")
        dialog.setMinimumSize(700, 500)
        
        layout = QVBoxLayout(dialog)
        
        shortcuts_html = """
        <html>
        <head>
            <style>
                body { font-family: 'Segoe UI', sans-serif; background: #1e1e1e; color: #f5f5f5; }
                h2 { color: #3b82f6; margin-top: 20px; }
                table { width: 100%; border-collapse: collapse; margin: 10px 0; }
                th, td { padding: 10px; text-align: left; border-bottom: 1px solid #3a3a3a; }
                th { background: #2d2d2d; color: #3b82f6; font-weight: 600; }
                tr:hover { background: #2a2a2a; }
                kbd { 
                    background: #2d2d2d; 
                    border: 1px solid #3a3a3a; 
                    border-radius: 4px; 
                    padding: 2px 6px; 
                    font-family: 'Consolas', monospace;
                    color: #3b82f6;
                }
            </style>
        </head>
        <body>
            <h2>📝 Editing</h2>
            <table>
                <tr><th>Shortcut</th><th>Action</th></tr>
                <tr><td><kbd>Tab</kbd></td><td>Indent (inserts 4 spaces)</td></tr>
                <tr><td><kbd>Shift</kbd>+<kbd>Tab</kbd></td><td>Outdent selected lines</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>Z</kbd></td><td>Undo (current file only)</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>Y</kbd></td><td>Redo (current file only)</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>A</kbd></td><td>Select All</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>C</kbd></td><td>Copy</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>X</kbd></td><td>Cut</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>V</kbd></td><td>Paste</td></tr>
            </table>
            
            <h2>🔍 Find & Replace</h2>
            <table>
                <tr><th>Shortcut</th><th>Action</th></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>F</kbd></td><td>Find</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>H</kbd></td><td>Find and Replace</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>Shift</kbd>+<kbd>F</kbd></td><td>Find in Files</td></tr>
                <tr><td><kbd>F3</kbd></td><td>Find Next</td></tr>
                <tr><td><kbd>Shift</kbd>+<kbd>F3</kbd></td><td>Find Previous</td></tr>
            </table>
            
            <h2>📑 File & Tab Navigation</h2>
            <table>
                <tr><th>Shortcut</th><th>Action</th></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>Tab</kbd></td><td>Next Tab</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>Shift</kbd>+<kbd>Tab</kbd></td><td>Previous Tab</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>W</kbd></td><td>Close Current Tab</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>Shift</kbd>+<kbd>W</kbd></td><td>Close All Tabs</td></tr>
                <tr><td><kbd>F2</kbd></td><td>Rename File</td></tr>
            </table>
            
            <h2>🚀 Quick Open</h2>
            <table>
                <tr><th>Shortcut</th><th>Action</th></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>P</kbd></td><td>Quick Open File</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>Shift</kbd>+<kbd>O</kbd></td><td>Go to Symbol</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>G</kbd></td><td>Go to Line</td></tr>
            </table>
            
            <h2>🎨 View & Tools</h2>
            <table>
                <tr><th>Shortcut</th><th>Action</th></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>B</kbd></td><td>Toggle Sidebar</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>`</kbd></td><td>Toggle Terminal</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>=</kbd></td><td>Zoom In</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>-</kbd></td><td>Zoom Out</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>0</kbd></td><td>Reset Zoom</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>Shift</kbd>+<kbd>P</kbd></td><td>Command Palette</td></tr>
            </table>
        </body>
        </html>
        """
        
        label = QLabel(shortcuts_html)
        label.setWordWrap(True)
        layout.addWidget(label)
        
        dialog.exec()

    # ------------------------------------------------------------------
    # Find/Replace Handlers
    # ------------------------------------------------------------------
    def _on_find_requested(self, text: str, options: dict):
        """Handle find request from dialog."""
        editor = self._editor_tabs.current_editor()
        if not editor:
            return
        
        cursor = editor.textCursor()
        
        # Search options
        case_sensitive = options.get('case_sensitive', False)
        whole_word = options.get('whole_word', False)
        use_regex = options.get('use_regex', False)
        wrap_around = options.get('wrap_around', True)
        search_forward = options.get('forward', True)
        
        # Build search flags
        flags = 0
        if case_sensitive:
            flags |= 0x00010  # QTextDocument.FindFlag.FindCaseSensitively
        
        # Perform search
        from PyQt6.QtGui import QTextDocument
        find_flags = QTextDocument.FindFlag(flags)
        
        if search_forward:
            found = editor.find(text, find_flags)
        else:
            found = editor.find(text, find_flags | QTextDocument.FindFlag.FindBackward)
        
        if not found and wrap_around:
            # Wrap around
            cursor.movePosition(cursor.MoveOperation.Start if search_forward else cursor.MoveOperation.End)
            editor.setTextCursor(cursor)
            if search_forward:
                found = editor.find(text, find_flags)
            else:
                found = editor.find(text, find_flags | QTextDocument.FindFlag.FindBackward)
    
    def _on_replace_requested(self, find_text: str, replace_text: str, options: dict):
        """Handle replace request from dialog."""
        editor = self._editor_tabs.current_editor()
        if not editor:
            return
        
        cursor = editor.textCursor()
        
        # Check if there's selected text matching find_text
        selected = cursor.selectedText()
        if selected == find_text:
            cursor.insertText(replace_text)
        
        # Find next
        self._on_find_requested(find_text, options)
    
    def _on_replace_all_requested(self, find_text: str, replace_text: str, options: dict):
        """Handle replace all request from dialog."""
        import re
        editor = self._editor_tabs.current_editor()
        if not editor:
            return
        
        document = editor.document()
        cursor = editor.textCursor()
        
        count = 0
        case_sensitive = options.get('case_sensitive', False)
        whole_word = options.get('whole_word', False)
        use_regex = options.get('use_regex', False)
        
        if use_regex:
            # Regex replace all
            try:
                flags = 0 if case_sensitive else re.IGNORECASE
                pattern = re.compile(find_text, flags)
                content = editor.toPlainText()
                new_content, count = pattern.subn(replace_text, content)
                
                if count > 0:
                    cursor.beginEditBlock()
                    cursor.select(cursor.SelectionType.Document)
                    cursor.insertText(new_content)
                    cursor.endEditBlock()
                    
            except re.error as e:
                from PyQt6.QtWidgets import QMessageBox
                QMessageBox.warning(self, "Regex Error", f"Invalid regular expression: {e}")
                return
        else:
            # Simple replace all
            cursor.beginEditBlock()
            
            # Save original position
            original_position = cursor.position()
            
            # Move to start of document
            cursor.setPosition(0)
            
            # Find and replace all occurrences
            while True:
                found = document.find(find_text, cursor, 
                                    QTextDocument.FindFlag.FindCaseSensitively if case_sensitive else QTextDocument.FindFlag(0))
                
                if found.isNull():
                    break
                
                # Check whole word if needed
                if whole_word:
                    # Verify it's a whole word
                    start = found.selectionStart()
                    end = found.selectionEnd()
                    text = editor.toPlainText()
                    
                    # Check character before
                    if start > 0 and (text[start-1].isalnum() or text[start-1] == '_'):
                        cursor.setPosition(end)
                        continue
                    
                    # Check character after
                    if end < len(text) and (text[end].isalnum() or text[end] == '_'):
                        cursor.setPosition(end)
                        continue
                
                # Replace
                found.insertText(replace_text)
                count += 1
            
            cursor.endEditBlock()
        
        from PyQt6.QtWidgets import QMessageBox
        if count > 0:
            QMessageBox.information(self, "Replace All", f"Replaced {count} occurrence(s).")
        else:
            QMessageBox.information(self, "Replace All", f"No occurrences of '{find_text}' found.")
        
        # Restore cursor position
        cursor.setPosition(original_position)
        editor.setTextCursor(cursor)

    # ------------------------------------------------------------------
    # AI Actions
    # ------------------------------------------------------------------
    def _on_ai_run_command(self, command: str):
        """Execute command from AI in active terminal."""
        self._terminal_tabs.setVisible(True)
        term = self._current_terminal()
        if not term:
            term = self._new_terminal()
        term.setFocus()
        term.execute_command(command)
        # Ensure terminal is visible
        self._terminal_tabs.setCurrentIndex(self._terminal_tabs.indexOf(term))

    def _on_directory_contents_for_tree(self, path: str, contents: str):
        """Handle directory contents for project tree card display."""
        self._ai_chat.emit_directory_tree(path, contents)

    def _on_ai_chat_message(self, message: str):
        """Handle user message from AI chat with project context."""
        # TODO: POINTS SYSTEM - Disabled for development
        # Will be enabled when connecting to https://logic-practice.com backend
        # Points tracking will be handled via API authentication with production server
        """
        # Check points balance before processing
        try:
            points_mgr = get_points_manager()
            perf_mode = self._get_current_performance_mode()
            estimated_tokens = len(message) // 4  # Rough estimate: 4 chars per token
            
            # Check if user can afford this request
            if not points_mgr.can_afford(estimated_tokens, perf_mode):
                cost = points_mgr.estimate_cost(estimated_tokens, perf_mode)
                balance = points_mgr.get_balance()
                log.warning(
                    f"[MainWindow] Insufficient points: need {cost:,}, have {balance:,}. "
                    f"Mode: {perf_mode}, estimated tokens: {estimated_tokens}"
                )
                # Send error to UI
                self._ai_chat.on_error(
                    f"Insufficient points. This request needs ~{cost:,} points but you have {balance:,} points.\n\n"
                    f"Please purchase more points or switch to Efficient mode (0.3x cost)."
                )
                return
            
            log.info(
                f"[MainWindow] Points check passed: balance={points_mgr.get_balance():,}, "
                f"mode={perf_mode}, estimated_cost={points_mgr.estimate_cost(estimated_tokens, perf_mode):,}"
            )
        except Exception as e:
            log.warning(f"[MainWindow] Points check failed: {e} - proceeding anyway")
        """
        
        # Build context for ALL messages - let agent_bridge decide
        # what to do with simple vs complex queries
        context = []

        # 1. Project Root Info
        if self._project_manager.root:
            context.append(f"Project path: {self._project_manager.root}")

        # 2. Current File Context
        editor = self._editor_tabs.current_editor()
        if editor:
            fp = self._editor_tabs.current_filepath()
            if fp:
                name = Path(fp).name
                content = editor.get_all_text()
                # Limit context size
                if len(content) > 5000:
                    content = content[:5000] + "... (truncated)"
                context.append(f"Current file ({name}):\n```\n{content}\n```")

        full_context = "\n\n".join(context)
        
        # NOTE: _ai_integration was removed - using _ai_agent directly
        # session_id = getattr(self._ai_chat, '_current_conversation_id', 'default-session')
        # self._ai_integration.set_session(...)  # Removed - not needed
        
        # Start AI processing immediately for responsiveness
        # But wait - we need to see if enhancement layer (Intent/Routing) is fast enough
        # Optimization: start AI chat, and if enhancement layer finds a special tool/route, 
        # we can inject that context later or stop/restart. 
        # For now, let's just fix the DUPLICATE problem by only calling it ONCE.
        
        conv_id = getattr(self._ai_chat, '_current_conversation_id', None) if hasattr(self._ai_chat, '_current_conversation_id') else None
        
        if conv_id and not self._title_generator.get_cached_title(conv_id):
            title = self._ai_agent.generate_chat_title(message, conv_id)
            if title:
                log.info(f"Generated chat title: {title}")
                if hasattr(self._ai_chat, 'update_conversation_title'):
                    self._ai_chat.update_conversation_title(conv_id, title)
        
        if conv_id:
            try:
                sessions = self._session_db.list_sessions(self._project_manager.root, limit=1)
                if not sessions:
                    title = self._title_generator.get_cached_title(conv_id) or "New Chat"
                    self._session_db.create_session(
                        conv_id, 
                        title, 
                        self._project_manager.root or "", 
                        self._ai_agent._settings.get("ai", "model", default="mistral-large-latest")
                    )
                import uuid
                self._session_db.add_message(
                    conv_id,
                    str(uuid.uuid4()),
                    "user",
                    message,
                    len(message) // 4
                )
            except Exception as e:
                log.debug(f"Could not store message in database: {e}")
        
        # FIX: Initialize enhancement_result to None (not yet implemented in async flow)
        enhancement_result = None
        if enhancement_result and enhancement_result.get("intent"):
            if enhancement_result.get("testing_decision") and \
               enhancement_result["testing_decision"].decision == 'write_tests':
                log.info("Using chat_with_testing with testing workflow")
                self._ai_agent.chat_with_testing(
                    message,
                    code_changes=[],
                    code_context=full_context
                )
            else:
                log.info("Using chat_with_enhancement with intent classification data")
                self._ai_agent.chat_with_enhancement(
                    message, 
                    intent=enhancement_result["intent"],
                    route=enhancement_result["route"],
                    tools=enhancement_result["tools"],
                    code_context=full_context
                )
        else:
            self._ai_agent.chat(message, full_context)

    def _on_model_changed(self, model_id: str, perf: str, cost: str):
        """Handle model selection change from AI chat."""
        log.info(f"[MainWindow] DEBUG: model_id='{model_id}'")
        
        # Check if this is a performance mode (not an actual model)
        performance_modes = ["efficient", "auto", "performance", "ultimate"]
        if model_id.lower() in performance_modes:
            # This is a performance mode, not a model ID
            # Save it to settings for the performance mode system to use
            try:
                from src.config.settings import get_settings
                settings = get_settings()
                settings.set("ai", "performance_mode", model_id.lower())
                log.info(f"[MainWindow] Performance mode changed to: {model_id}")
                # Don't update agent with this - it's not a real model
                return
            except Exception as e:
                log.error(f"[MainWindow] Failed to save performance mode: {e}")
                return
        
        # This is an actual model ID - determine provider from model_id
        if model_id.startswith("mistral-") or model_id.startswith("codestral-"):
            # Mistral AI models (primary provider)
            provider = "mistral"
        elif model_id.startswith(("gpt-", "o1", "o3")):
            # OpenAI models - determine which API to use
            # Chat Completions: gpt-4o, gpt-4.1-*
            # Responses API: gpt-5.*, o1, o3, codex
            if any(x in model_id.lower() for x in ["codex", "gpt-5", "o1", "o3"]):
                provider = "openai_responses"
            else:
                provider = "openai"
        elif "/" in model_id:
            # Vendor models via SiliconFlow
            provider = "siliconflow"
        else:
            provider = "mistral"  # Default to Mistral
        
        log.info(f"[MainWindow] Model changed to: {model_id} (provider: {provider})")
        self._ai_agent.update_settings(provider=provider, model_id=model_id)

    def _on_ai_stop_requested(self):
        """Handle stop request from AI (via web bridge)."""
        self._ai_agent.stop()
    
    def _on_toggle_autogen(self):
        """Toggle AutoGen multi-agent mode."""
        # Get current status
        status = self._ai_agent.get_autogen_status()
        current_enabled = status.get('enabled', False)
        
        # Toggle
        new_state = not current_enabled
        self._ai_agent.enable_autogen(new_state)
        
        log.info(f"AutoGen {'enabled' if new_state else 'disabled'} via UI toggle")
    
    def _show_command_palette(self):
        """Show the command palette - REMOVED in AI-first mode."""
        # Command palette not implemented
        pass
        # if hasattr(self, '_command_palette'):
        #     self._command_palette.show_palette()
    
    def _on_command_selected(self, action: str, command_data: dict):
        """Handle command selection from command palette - REMOVED in AI-first mode."""
        # Command palette not implemented
        pass

    def _on_generate_plan(self):
        log.info("MainWindow: Automated plan generation triggered")
        self._ai_agent.chat("__GENERATE_PLAN__")

    def _ai_action(self, action: str):
        editor = self._editor_tabs.current_editor()
        if not editor:
            self._ai_chat.add_system_message("Open a file first to use AI actions.")
            return

        code = editor.get_selected_text() or editor.get_all_text()
        language = editor.language
        analyzer = CodeAnalyzer()

        prompts = {
            "explain":   analyzer.build_explain_prompt(code, language),
            "refactor":  analyzer.build_refactor_prompt(code, language),
            "tests":     analyzer.build_test_prompt(code, language),
            "debug":     analyzer.build_debug_prompt(code, "unknown error", language),
            "docstring": f"Add comprehensive docstrings to this {language} code:\n\n```{language}\n{code}\n```",
        }

        prompt = prompts.get(action, f"Help me with this {language} code:\n\n```{language}\n{code}\n```")
        self._ai_chat.add_system_message(f"🤖 Running: {action.title()} Code…")
        self._ai_chat._add_ai_bubble_streaming()
        self._ai_agent.chat(prompt)

    def _get_selected_code(self) -> str:
        editor = self._editor_tabs.current_editor()
        if editor:
            return editor.get_selected_text() or editor.get_all_text()
        return ""

    def _get_code_context(self) -> str:
        """Alias for _get_selected_code to satisfy WebAIChatWidget callback."""
        return self._get_selected_code()

    # ------------------------------------------------------------------
    # Project & Events
    # ------------------------------------------------------------------

    def _cleanup_old_project(self):
        """Clean up all state from the old project before opening a new one."""
        # Only cleanup if we have a previous project loaded
        if not hasattr(self, '_current_project_path') or not self._current_project_path:
            log.info("🆕 First project load - skipping cleanup")
            return
            
        log.info("🧹 Cleaning up old project state...")
        
        # 1. Close all editor tabs
        self._editor_tabs.close_all_tabs()
        log.info("   ✓ Closed all editor tabs")
        
        # 2. Clear file snapshots (diff data)
        if hasattr(self, '_file_snapshots'):
            self._file_snapshots.clear()
            log.info("   ✓ Cleared file snapshots")
        
        # 3. Clear diff cache
        if hasattr(self, '_diff_data_store'):
            self._diff_data_store.clear()
            log.info("   ✓ Cleared diff data store")
        
        # 4. Clear file tracker
        if hasattr(self, '_file_tracker'):
            if hasattr(self._file_tracker, '_edits'):
                self._file_tracker._edits.clear()
            log.info("   ✓ Cleared file edit history")
        
        # 5. Clear AI agent context
        self._ai_agent.clear_active_file()
        log.info("   ✓ Cleared AI agent active file")
        
        # 6. Clear codebase index
        self._codebase_index = None
        log.info("   ✓ Cleared codebase index")
        
        # 7. Prepare terminals for new project (will set CWD after cleanup)
        log.info("   ✓ Prepared terminals for new project")
        
        # 8. Todo list is session-based, no need to clear
        # Todos are tied to chat sessions, not projects
        log.info("   ✓ Skipped todo cleanup (session-based)")
        
        # 9. Clear search results
        if hasattr(self, '_search_results'):
            self._search_results.clear()
            log.info("   ✓ Cleared search results")
        
        log.info("✅ Old project cleanup complete!")
    
    def _on_project_opened(self, folder_path: str):
        log.info(f"Project opened: {folder_path}")
        
        # Clean up old project state BEFORE loading new one (only if switching)
        self._cleanup_old_project()
        
        # Set project root FIRST (this loads project-specific context)
        self._ai_agent.set_project_root(folder_path)
        self._current_project_path = folder_path  # Track current project
        
        # Update LSP manager workspace root so completions/diagnostics work
        try:
            from src.core.lsp_manager import get_lsp_manager
            get_lsp_manager().set_project_root(folder_path)
        except Exception:
            pass
        
        # Reset codebase index for new project
        self._codebase_index = None
        
        # Note: Chat history is now per-project, so we don't clear it.
        # The AI chat will automatically load the history for this project.
        self._ai_agent.clear_active_file()
        
        self._sidebar.set_project(folder_path)
        
        # Update sidebar project name display
        if hasattr(self, '_sidebar_project_item'):
            from pathlib import Path as _Path
            self._sidebar_project_item.setText(f"  📁 {_Path(folder_path).name}")
        
        # Update all current terminal tabs to the new project directory
        for i in range(self._terminal_tabs.count()):
            term = self._terminal_tabs.widget(i)
            if isinstance(term, XTermWidget):
                term.set_cwd(folder_path)
                
        project_name = Path(folder_path).name
        self.setWindowTitle(f"Cortex AI Agent — {project_name}")
        self._ai_chat.add_system_message(f"📂 Opened: {folder_path}")
        
        # Load existing chats for this project from SQLite BEFORE setting project info
        try:
            chats_json = self._ai_chat.load_chats_for_project(folder_path)
        except Exception as e:
            log.warning(f"Failed to load chats for project {folder_path}: {e}")
            chats_json = "[]"
        
        # Populate Qt sidebar with loaded chats immediately
        if chats_json and chats_json != "[]":
            log.info(f"[ChatList] Initial population with {len(json.loads(chats_json))} chats")
            self._refresh_sidebar_chat_list(chats_json)
        
        # Update project indicator in AI chat (this triggers project-specific chat loading)
        self._ai_chat.set_project_info(project_name, folder_path, chats_json)
        
        # Update welcome tab if it exists
        self._update_welcome_project_info()
        
        # Update git status in Summary/Review panel
        QTimer.singleShot(300, self._update_git_summary)
        
        # Auto-detect and activate virtual environment
        self._check_and_activate_venv(folder_path)
        
        # Start background project context scan
        self._start_project_context_scan(folder_path)
    
    def _start_project_context_scan(self, folder_path: str):
        """Start background project scanning for instant AI awareness."""
        from src.ai.project_context import build_project_context
        
        # Show indexing status in chat
        self._ai_chat.show_indexing_status("Indexing project...")
        
        def on_context_ready(ctx):
            """Called when background scan finishes."""
            from PyQt6.QtCore import QMetaObject, Qt
            # Must update UI from main thread
            QMetaObject.invokeMethod(
                self, "_on_project_context_ready",
                Qt.ConnectionType.QueuedConnection,
            )
            # Store context in agent
            self._ai_agent.set_project_context(ctx)
        
        self._context_builder = build_project_context(folder_path, on_context_ready)
    
    @pyqtSlot()
    def _on_project_context_ready(self):
        """Called on main thread when project indexing finishes."""
        from src.ai.project_context import get_project_context
        ctx = get_project_context(self._ai_agent._project_root)
        if ctx:
            self._ai_chat.hide_indexing_status()
            # Show ready indicator  
            self._ai_chat.show_indexing_status(
                f"✓ Indexed {ctx.source_file_count} files ({ctx.build_time_ms:.0f}ms)",
                auto_hide=True
            )

    def _on_project_closed(self):
        """Handle project close - show welcome page and clean up."""
        log.info("Project closed - showing welcome page")

        # Stop Live Server (removed in AI-first mode)
        # if self._live_server and self._live_server.is_running:
        #     self._live_server.stop()
        self._live_server = None

        # Clear all state
        self._current_project_path = None
        self.setWindowTitle("Cortex AI Agent")
        self._ai_chat.clear_project_info()
        
        # Close all editor tabs
        if hasattr(self, '_editor_tabs'):
            self._editor_tabs.close_all_tabs()
        
        # Clear file snapshots and caches
        if hasattr(self, '_file_snapshots'):
            self._file_snapshots.clear()
        if hasattr(self, '_diff_data_store'):
            self._diff_data_store.clear()
        if hasattr(self, '_file_tracker') and hasattr(self._file_tracker, '_edits'):
            self._file_tracker._edits.clear()
        if hasattr(self, '_search_results'):
            self._search_results.clear()
        
        # Clear AI agent context
        self._ai_agent.clear_active_file()
        self._codebase_index = None
        
        # Update welcome page
        self._update_welcome_project_info()
        
        # Show welcome page if not already visible
        self._show_welcome()

    def _update_welcome_project_info(self):
        """Update project info display."""
        # Update AI chat with project info if available
        if hasattr(self, '_ai_chat') and self._ai_chat:
            if self._project_manager.root:
                name = self._project_manager.root.name
                path = str(self._project_manager.root)
                # TODO: Update AI chat project indicator
        
        # Keep old welcome widget updated for backward compatibility
        if hasattr(self, '_welcome_project_info') and self._welcome_project_info is not None:
            if self._project_manager.root:
                name = self._project_manager.root.name
                path = str(self._project_manager.root)
                self._welcome_project_info.setText(f"Project: <b>{name}</b><br/><span style='font-size:11px; color:#858585;'>{path}</span>")
            else:
                self._welcome_project_info.setText("No project opened")


    def _check_and_activate_venv(self, project_path: str):
        """Check for and activate Python virtual environment."""
        import os
        
        venv_names = ["venv", ".venv", "env", ".env", "virtualenv"]
        
        for venv_name in venv_names:
            venv_path = os.path.join(project_path, venv_name)
            if os.path.exists(venv_path):
                # Activate in all terminal tabs
                for i in range(self._terminal_tabs.count()):
                    term = self._terminal_tabs.widget(i)
                    if isinstance(term, XTermWidget):
                        term.activate_virtual_env(venv_path)
                
                # Show notification in chat
                self._ai_chat.add_system_message(f"🐍 Virtual environment detected: {venv_name}")
                
                # Update status bar
                self.statusBar().showMessage(f"Virtual environment: {venv_name}", 5000)
                break

    def _on_tab_changed(self, index: int):
        fp = self._editor_tabs._files.get(index)
        self._update_status_file(fp)
        
        # Update AI agent with active file for context injection
        if fp and hasattr(self, '_ai_agent'):
            try:
                editor = self._editor_tabs.widget(index)
                cursor_pos = None
                if hasattr(editor, 'getCursorPosition'):
                    line, col = editor.getCursorPosition()
                    cursor_pos = (line + 1, col)  # Convert to 1-indexed
                self._ai_agent.set_active_file(fp, cursor_pos)
                log.debug(f"Active file updated for AI: {fp} at {cursor_pos}")
            except Exception as e:
                log.warning(f"Could not update active file for AI: {e}")

    # def _apply_initial_theme(self):
    #     from PyQt6.QtWidgets import QApplication
    #     app = QApplication.instance()
    #     theme = self._settings.theme
    #     self._theme_manager.apply(theme, app)
    #     self._theme_btn.setText("☀️" if theme == "light" else "🌙")

    def _restore_session(self):
        # Restore last project — skip if it no longer exists (blank state)
        log.info("Restoring last session project...")
        restored = self._project_manager.restore_last()

        if not restored:
            # No valid project to restore — show clean blank state, no stale tabs
            log.info("No project to restore — showing blank start state.")
            # Remove any tabs that sneak in during build (stale welcome)
            self._show_welcome()
            return

        # Set project root for file tree (with delay to ensure UI is ready)
        if self._project_manager.root:
            project_path = str(self._project_manager.root)
            log.info(f"[FILE TREE] Restoring project root: {project_path}")
            QTimer.singleShot(200, lambda: self._set_project_root(project_path))
            
            # Initialize Git repository for restored project
            if hasattr(self, '_git_manager'):
                self._git_manager.set_repository(project_path)
                log.info(f"[GIT] Repository restored to: {project_path}")
                QTimer.singleShot(500, self._update_git_summary)

        # Restore open files only if project was successfully restored
        session = self._session_manager.load()
        if session:
            log.info(f"Restoring {len(session.get('open_files', []))} files...")
            for fp in session.get("open_files", []):
                if Path(fp).exists():
                    self._open_file(fp)

        # Focus the active file
        if session:
            active = session.get("active_file")
            if active and Path(active).exists():
                for i in range(self._editor_tabs.count()):
                    if self._editor_tabs._files.get(i) == active:
                        self._editor_tabs.setCurrentIndex(i)
                        break



    @property
    def codebase_index(self):
        """Get the codebase index, creating it if needed."""
        if self._codebase_index is None:
            if self._project_manager.root:
                self._codebase_index = get_codebase_index(str(self._project_manager.root))
                self._codebase_index.index_project()
            else:
                # No project open, return a dummy index?
                raise RuntimeError("No project open for indexing")
        return self._codebase_index

    def _show_about(self):
        QMessageBox.about(self, "About Cortex AI Agent",
                          "<h2>🧠 Cortex AI Agent</h2>"
                          "<p>A modern AI-powered IDE built with Python and PyQt6.</p>"
                          "<p>Features: Multi-file editor · Syntax highlighting · "
                          "AI chat · File explorer · Terminal</p>"
                          "<p><b>Version:</b> 1.0.0</p>")

    def _open_documentation(self):
        """Open Cortex documentation in browser."""
        import webbrowser
        webbrowser.open("https://github.com/cortex-ai/docs")
        log.info("Opening documentation")

    def _show_whats_new(self):
        """Show what's new dialog."""
        log.info("What's new dialog requested")

    def _show_automations(self):
        """Show automations help."""
        log.info("Automations help requested")

    def _show_local_envs(self):
        """Show local environments help."""
        log.info("Local environments help requested")

    def _show_worktrees(self):
        """Show worktrees help."""
        log.info("Worktrees help requested")

    def _show_skills_help(self):
        """Show skills help."""
        log.info("Skills help requested")

    def _show_mcp_help(self):
        """Show Model Context Protocol help."""
        log.info("MCP help requested")

    def _show_troubleshooting(self):
        """Show troubleshooting guide."""
        log.info("Troubleshooting guide requested")

    def _send_feedback(self):
        """Send feedback."""
        import webbrowser
        webbrowser.open("https://github.com/cortex-ai/feedback")
        log.info("Opening feedback page")

    def _start_trace(self):
        """Start trace recording for debugging."""
        log.info("Trace recording started")

    def _show_keyboard_shortcuts(self):
        """Show keyboard shortcuts reference dialog (F1)."""
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QTextEdit
        
        dialog = QDialog(self)
        dialog.setWindowTitle("Keyboard Shortcuts Reference")
        dialog.setMinimumSize(800, 600)
        
        layout = QVBoxLayout(dialog)
        
        # Create a text edit for displaying shortcuts
        shortcuts_text = QTextEdit()
        shortcuts_text.setReadOnly(True)
        shortcuts_text.setFontFamily("Consolas")
        shortcuts_text.setFontPointSize(10)
        
        # Build shortcuts HTML - Simple dark theme
        html_content = """
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body { 
                    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; 
                    background-color: #1e1e1e;
                    color: #ffffff;
                    padding: 30px;
                    margin: 0;
                }
                h1 {
                    color: #4CAF50;
                    text-align: center;
                    font-size: 32px;
                    margin-bottom: 10px;
                }
                .subtitle {
                    text-align: center;
                    color: #9cdcfe;
                    margin-bottom: 30px;
                    font-size: 16px;
                }
                .section {
                    background-color: #252526;
                    border-left: 4px solid #4CAF50;
                    padding: 20px;
                    margin-bottom: 25px;
                    border-radius: 5px;
                }
                .section h2 {
                    color: #4CAF50;
                    margin: 0 0 15px 0;
                    font-size: 20px;
                }
                table {
                    width: 100%;
                    border-collapse: collapse;
                    background-color: #2d2d30;
                }
                th {
                    background-color: #3e3e42;
                    color: #ffffff;
                    padding: 12px;
                    text-align: left;
                    font-weight: 600;
                    border: 1px solid #555;
                }
                td {
                    padding: 10px 12px;
                    border: 1px solid #555;
                    color: #cccccc;
                }
                tr:nth-child(even) {
                    background-color: #333337;
                }
                tr:hover {
                    background-color: #3e3e42;
                }
                .shortcut {
                    font-family: 'Consolas', monospace;
                    background-color: #1e1e1e;
                    padding: 5px 10px;
                    border-radius: 4px;
                    font-weight: bold;
                    color: #9cdcfe;
                    display: inline-block;
                    min-width: 110px;
                    text-align: center;
                }
                .status {
                    color: #4CAF50;
                    font-weight: bold;
                }
                .tip {
                    background-color: #264f78;
                    padding: 15px;
                    border-radius: 5px;
                    margin-top: 25px;
                    border-left: 4px solid #007acc;
                }
                .tip strong {
                    color: #9cdcfe;
                }
                .tip p {
                    margin: 8px 0 0 0;
                    color: #cccccc;
                }
            </style>
        </head>
        <body>
            <h1>⌨️ Keyboard Shortcuts Reference</h1>
            <p class="subtitle">Quick reference for all Cortex IDE shortcuts</p>
            
            <!-- File Operations -->
            <div class="section">
                <h2>📁 File Operations</h2>
                <table>
                    <tr><th>Action</th><th>Shortcut</th><th>Status</th></tr>
                    <tr><td>New File</td><td><span class="shortcut">Ctrl+N</span></td><td class="status">✅ Ready</td></tr>
                    <tr><td>Open Folder</td><td><span class="shortcut">Ctrl+O</span></td><td class="status">✅ Ready</td></tr>
                    <tr><td>Open File</td><td><span class="shortcut">Ctrl+Shift+O</span></td><td class="status">✅ Ready</td></tr>
                    <tr><td>Save</td><td><span class="shortcut">Ctrl+S</span></td><td class="status">✅ Ready</td></tr>
                    <tr><td>Save All</td><td><span class="shortcut">Ctrl+Shift+S</span></td><td class="status">✅ Ready</td></tr>
                    <tr><td>Close Tab</td><td><span class="shortcut">Ctrl+W</span></td><td class="status">✅ Ready</td></tr>
                    <tr><td>Close All Tabs</td><td><span class="shortcut">Ctrl+Shift+W</span></td><td class="status">✅ Ready</td></tr>
                    <tr><td>Rename File</td><td><span class="shortcut">F2</span></td><td class="status">✅ Ready</td></tr>
                </table>
            </div>
            
            <!-- Edit Operations -->
            <div class="section">
                <h2>✏️ Edit Operations</h2>
                <table>
                    <tr><th>Action</th><th>Shortcut</th><th>Status</th></tr>
                    <tr><td>Undo</td><td><span class="shortcut">Ctrl+Z</span></td><td class="status">✅ Ready</td></tr>
                    <tr><td>Redo</td><td><span class="shortcut">Ctrl+Y</span></td><td class="status">✅ Ready</td></tr>
                    <tr><td>Cut</td><td><span class="shortcut">Ctrl+X</span></td><td class="status">✅ Ready</td></tr>
                    <tr><td>Copy</td><td><span class="shortcut">Ctrl+C</span></td><td class="status">✅ Ready</td></tr>
                    <tr><td>Paste</td><td><span class="shortcut">Ctrl+V</span></td><td class="status">✅ Ready</td></tr>
                    <tr><td>Select All</td><td><span class="shortcut">Ctrl+A</span></td><td class="status">✅ Ready</td></tr>
                    <tr><td>Toggle Comment</td><td><span class="shortcut">Ctrl+/</span></td><td class="status">✅ Ready</td></tr>
                    <tr><td>Delete Line</td><td><span class="shortcut">Ctrl+Shift+K</span></td><td class="status">✅ Ready</td></tr>
                    <tr><td>Indent Selection</td><td><span class="shortcut">Tab</span></td><td class="status">✅ Ready</td></tr>
                    <tr><td>Outdent Selection</td><td><span class="shortcut">Shift+Tab</span></td><td class="status">✅ Ready</td></tr>
                </table>
            </div>
            
            <!-- Find & Replace -->
            <div class="section">
                <h2>🔍 Find & Replace</h2>
                <table>
                    <tr><th>Action</th><th>Shortcut</th><th>Status</th></tr>
                    <tr><td>Find</td><td><span class="shortcut">Ctrl+F</span></td><td class="status">✅ Ready</td></tr>
                    <tr><td>Find & Replace</td><td><span class="shortcut">Ctrl+H</span></td><td class="status">✅ Ready</td></tr>
                    <tr><td>Replace All</td><td><span class="shortcut">In Dialog</span></td><td class="status">✅ Ready</td></tr>
                </table>
            </div>
            
            <!-- Navigation -->
            <div class="section">
                <h2>🧭 Navigation</h2>
                <table>
                    <tr><th>Action</th><th>Shortcut</th><th>Status</th></tr>
                    <tr><td>Next Tab</td><td><span class="shortcut">Ctrl+Tab</span></td><td class="status">✅ Ready</td></tr>
                    <tr><td>Previous Tab</td><td><span class="shortcut">Ctrl+Shift+Tab</span></td><td class="status">✅ Ready</td></tr>
                </table>
            </div>
            
            <!-- View -->
            <div class="section">
                <h2>👁️ View</h2>
                <table>
                    <tr><th>Action</th><th>Shortcut</th><th>Status</th></tr>
                    <tr><td>Toggle Theme</td><td><span class="shortcut">Ctrl+Shift+T</span></td><td class="status">✅ Ready</td></tr>
                    <tr><td>Toggle Terminal</td><td><span class="shortcut">Ctrl+`</span></td><td class="status">✅ Ready</td></tr>
                    <tr><td>Toggle Sidebar</td><td><span class="shortcut">Ctrl+B</span></td><td class="status">✅ Ready</td></tr>
                </table>
            </div>
            
            <!-- Tip -->
            <div class="tip">
                <strong>💡 Tip:</strong>
                <p>Press <span class="shortcut" style="min-width: 50px;">F1</span> anytime to open this reference!</p>
            </div>
        </body>
        </html>
        """
        
        shortcuts_text.setHtml(html_content)
        layout.addWidget(shortcuts_text)
        
        dialog.exec()

    def closeEvent(self, event: QCloseEvent):
        """Save session on close, prompt for unsaved files, and kill terminals."""
        # 1. Check for unsaved files
        modified_files = self._editor_tabs._modified
        if modified_files:
            from PyQt6.QtWidgets import QMessageBox
            file_names = [os.path.basename(f) for f in modified_files]
            files_str = ", ".join(file_names[:3]) + ("..." if len(file_names) > 3 else "")
            
            reply = QMessageBox.question(
                self,
                "Unsaved Changes",
                f"You have unsaved changes in: {files_str}\n\nDo you want to save them before closing?",
                QMessageBox.StandardButton.SaveAll | 
                QMessageBox.StandardButton.Discard | 
                QMessageBox.StandardButton.Cancel
            )
            
            if reply == QMessageBox.StandardButton.Cancel:
                event.ignore()
                return
            elif reply == QMessageBox.StandardButton.SaveAll:
                # Save all modified files
                for filepath in list(modified_files):
                    idx = -1
                    for i, fp in self._editor_tabs._files.items():
                        if fp == filepath:
                            idx = i
                            break
                    
                    if idx >= 0:
                        editor = self._editor_tabs.widget(idx)
                        if isinstance(editor, CodeEditor):
                            content = editor.toPlainText()
                            try:
                                with open(filepath, 'w', encoding='utf-8') as f:
                                    f.write(content)
                                self._editor_tabs._mark_saved(filepath)
                            except Exception as e:
                                log.error(f"Failed to auto-save {filepath}: {e}")
        
        # 2. Force AI Chat persistence
        if hasattr(self, '_ai_chat') and self._ai_chat:
            log.info("Persisting AI chat history before close...")
            # Create a localized event loop to wait for the chat persistence to finish
            loop = QEventLoop()
            save_success = [False]
            
            def on_save_done(status):
                log.info(f"AI chat persistence finished with status: {status}")
                save_success[0] = (status == "OK")
                loop.quit()
                
            # Connect the bridge response signal
            self._ai_chat.save_finished.connect(on_save_done)
            
            # Start timer for timeout (3s max)
            QTimer.singleShot(3000, loop.quit)
            
            # Trigger JS to save its logic to SQLite
            self._ai_chat.run_javascript("if(window.saveProjectChats) saveProjectChats(window.chats);")
            
            # Wait for save or timeout
            loop.exec()
            
            if not save_success[0]:
                log.warning("AI chat persistence timed out or failed before close.")
            else:
                log.info("AI chat persistence confirmed.")
        
        # 3. Save IDE UI state
        fps = self._editor_tabs.get_open_files()
        active = self._editor_tabs.current_filepath()
        expanded = self._sidebar.get_expanded_paths()
        self._session_manager.save(fps, active, {"expanded_paths": expanded})
        self._settings.set("window", "maximized", self.isMaximized())
        if not self.isMaximized():
            self._settings.set("window", "width", self.width())
            self._settings.set("window", "height", self.height())
        
        # Save panel widths (Codex 4-panel layout)
        # Left sidebar: 220px, Review: 380px, File tree: 280px are fixed
        # Only chat panel width varies and is not saved (flexible)
            
        # 4. Clean up terminals
        for i in range(self._terminal_tabs.count()):
            term = self._terminal_tabs.widget(i)
            if isinstance(term, XTermWidget):
                term._kill_process()

        # 5. Stop Live Server if running
        # Stop Live Server (removed in AI-first mode)
        if False and self._live_server and self._live_server.is_running:
            self._live_server.stop()
                
        event.accept()

    def dragEnterEvent(self, event):
        """Accept drag of folders or files from Explorer onto the window."""
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        else:
            event.ignore()

    def dragMoveEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()

    def dropEvent(self, event):
        """Handle drop of folders/files — open as project or file."""
        for url in event.mimeData().urls():
            path = url.toLocalFile()
            if os.path.isdir(path):
                self._open_folder_programmatic(path)
                event.acceptProposedAction()
                return
            elif os.path.isfile(path):
                self._open_file(path)
                event.acceptProposedAction()
                return

    def resizeEvent(self, event):
        """Handle window resize — responsive layout + welcome font scaling.

        VS Code behaviour:
        - All three panels stay visible and simply shrink proportionally.
        - Only at very small widths (<700px) does the AI chat auto-hide.
        - Sidebar keeps its width; editor absorbs most resize delta.
        """
        super().resizeEvent(event)
        w = event.size().width()

        # Codex 4-panel layout: fixed widths for sidebar (220), review (380), file tree (280)
        # Chat panel is flexible. On very small screens, we may need to hide panels.
        if w < 1000:
            # Small window: hide file tree panel
            if hasattr(self, '_file_tree_panel'):
                self._file_tree_panel.setVisible(False)
        else:
            if hasattr(self, '_file_tree_panel'):
                self._file_tree_panel.setVisible(True)

        if hasattr(self, '_welcome_widget') and self._welcome_widget is not None:
            self._apply_welcome_theme(self._theme_manager.is_dark)

    # Phase 1, 2, 3 Integration Methods
    def _set_agent_mode(self, mode: str):
        """
        Set AI agent mode (Phase 1 Integration).
        
        Args:
            mode: One of 'build', 'explore', 'debug', 'plan'
        """
        self._ai_agent.set_mode(mode)
        mode_names = {
            'build': '🏗️ Build',
            'explore': '🔍 Explore', 
            'debug': '🐛 Debug',
            'plan': '📋 Plan'
        }
        self._statusbar.showMessage(f"Agent mode: {mode_names.get(mode, mode)}", 3000)
        log.info(f"Agent mode switched to: {mode}")
    
    def _show_skills_browser(self):
        """Show skills browser dialog (Phase 3 Integration)."""
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QListWidget, QLabel, QPushButton, QTextEdit
        
        dialog = QDialog(self)
        dialog.setWindowTitle("Skills Browser")
        dialog.setMinimumSize(600, 400)
        
        layout = QVBoxLayout(dialog)
        
        # Title
        title = QLabel("<h2>🛠️ Available Skills</h2>")
        layout.addWidget(title)
        
        # Skills list
        skills_list = QListWidget()
        skills = self._ai_agent.get_available_skills()
        
        for skill in skills:
            item_text = f"{skill['name']} ({skill['id']})"
            skills_list.addItem(item_text)
        
        layout.addWidget(skills_list)
        
        # Description
        desc_label = QLabel("Select a skill to view capabilities")
        layout.addWidget(desc_label)
        
        # Capability display
        capability_text = QTextEdit()
        capability_text.setReadOnly(True)
        capability_text.setPlaceholderText("Skill capabilities will appear here...")
        layout.addWidget(capability_text)
        
        # Close button
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.close)
        layout.addWidget(close_btn)
        
        dialog.exec()
        log.info("Skills browser opened")
    
    def _show_mcp_connections(self):
        """Show MCP connections dialog (Phase 3 Integration)."""
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QLabel, QPushButton, QLineEdit, QListWidget
        
        dialog = QDialog(self)
        dialog.setWindowTitle("MCP Connections")
        dialog.setMinimumSize(500, 300)
        
        layout = QVBoxLayout(dialog)
        
        # Title
        title = QLabel("<h2>🔗 MCP Server Connections</h2>")
        layout.addWidget(title)
        
        # Connected servers
        servers_label = QLabel("Connected Servers:")
        layout.addWidget(servers_label)
        
        servers_list = QListWidget()
        servers = self._ai_agent._mcp_manager.list_servers()
        if servers:
            for server in servers:
                servers_list.addItem(server)
        else:
            servers_list.addItem("No servers connected")
        
        layout.addWidget(servers_list)
        
        # Connect new server section
        layout.addWidget(QLabel("Connect New Server:"))
        
        name_input = QLineEdit()
        name_input.setPlaceholderText("Server name (e.g., github)")
        layout.addWidget(name_input)
        
        url_input = QLineEdit()
        url_input.setPlaceholderText("Server URL (e.g., https://mcp.github.com)")
        layout.addWidget(url_input)
        
        def connect_server():
            name = name_input.text().strip()
            url = url_input.text().strip()
            if name and url:
                success = self._ai_agent.connect_mcp_server(name, url)
                if success:
                    servers_list.addItem(name)
                    name_input.clear()
                    url_input.clear()
        
        connect_btn = QPushButton("Connect")
        connect_btn.clicked.connect(connect_server)
        layout.addWidget(connect_btn)
        
        # Close button
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.close)
        layout.addWidget(close_btn)
        
        dialog.exec()
        log.info("MCP connections dialog opened")

    # Phase 4 Integration Methods
    def _show_todo_manager(self):
        """Show TODO manager dialog.
        
        NOTE: The legacy _todo_manager has been removed. This dialog now shows
        the current todos from the bridge/JS UI instead.
        """
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QLabel, QPushButton, QListWidget
        
        dialog = QDialog(self)
        dialog.setWindowTitle("Tasks & TODOs")
        dialog.setMinimumSize(500, 350)
        
        layout = QVBoxLayout(dialog)
        
        title = QLabel("<h2>Task Manager</h2>")
        layout.addWidget(title)
        
        info = QLabel("Todos are now managed by the AI agent via TodoWrite.\nUse the chat sidebar to view and track task progress.")
        info.setWordWrap(True)
        layout.addWidget(info)
        
        layout.addWidget(QLabel("Current session todos are shown in the sidebar."))
        
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.close)
        layout.addWidget(close_btn)
        
        dialog.exec()
        log.info("TODO manager dialog opened")
    
    def _add_todo_task(self):
        """Add a new TODO task. (Legacy - todo manager removed)"""
        self._statusbar.showMessage("Todo manager has been removed. Ask the AI to create tasks via TodoWrite.", 5000)
        log.info("[TODO] _add_todo_task called - todo manager removed")
    
    def _complete_todo_task(self):
        """Complete a TODO task. (Legacy - todo manager removed)"""
        self._statusbar.showMessage("Todo manager has been removed. Toggle tasks in the sidebar instead.", 5000)
        log.info("[TODO] _complete_todo_task called - todo manager removed")
    
    def _show_permission_settings(self):
        """Show permission settings dialog (Phase 4 Integration)."""
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QLabel, QPushButton
        
        dialog = QDialog(self)
        dialog.setWindowTitle("Permission Settings")
        dialog.setMinimumSize(500, 300)
        
        layout = QVBoxLayout(dialog)
        
        # Title
        title = QLabel("<h2>🔒 Permission System</h2>")
        layout.addWidget(title)
        
        # Info
        info = QLabel("Permission system is active and monitoring tool usage.")
        layout.addWidget(info)
        
        # Cache info
        cache_size = len(self._permission_evaluator._permission_cache)
        cache_label = QLabel(f"Cached decisions: {cache_size}")
        layout.addWidget(cache_label)
        
        # Clear cache button
        def clear_cache():
            self._permission_evaluator.clear_cache()
            cache_label.setText("Cached decisions: 0")
        
        clear_btn = QPushButton("Clear Permission Cache")
        clear_btn.clicked.connect(clear_cache)
        layout.addWidget(clear_btn)
        
        # Close button
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.close)
        layout.addWidget(close_btn)
        
        dialog.exec()
        log.info("Permission settings dialog opened")
    
    def _show_memory_manager(self):
        """Open the Memory Manager dialog (AI → Memory Manager...)."""
        from src.ui.dialogs.memory_manager import MemoryManagerDialog
        try:
            from src.config.settings import get_settings
            settings = get_settings()
        except Exception:
            settings = None
        # Use current project path if available, fallback to cwd
        project_root = getattr(self, '_current_project_path', None) or os.getcwd()
        dlg = MemoryManagerDialog(project_root, settings=settings, parent=self)
        dlg.exec()

    def _show_github_integration(self):
        """Show GitHub integration dialog (Phase 4 Integration)."""
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QLabel, QPushButton, QLineEdit, QInputDialog
        
        dialog = QDialog(self)
        dialog.setWindowTitle("GitHub Integration")
        dialog.setMinimumSize(500, 300)
        
        layout = QVBoxLayout(dialog)
        
        # Title
        title = QLabel("<h2>🐙 GitHub Automation</h2>")
        layout.addWidget(title)
        
        # Repository info
        if self._github_agent.repo_owner and self._github_agent.repo_name:
            repo_text = f"Repository: {self._github_agent.repo_owner}/{self._github_agent.repo_name}"
        else:
            repo_text = "No repository configured"
        repo_label = QLabel(repo_text)
        layout.addWidget(repo_label)
        
        # Set repository button
        def set_repo():
            owner, ok1 = QInputDialog.getText(self, "GitHub Repository", "Repository owner:")
            if ok1 and owner:
                repo, ok2 = QInputDialog.getText(self, "GitHub Repository", "Repository name:")
                if ok2 and repo:
                    self._github_agent.set_repository(owner, repo)
                    repo_label.setText(f"Repository: {owner}/{repo}")
        
        set_repo_btn = QPushButton("Set Repository")
        set_repo_btn.clicked.connect(set_repo)
        layout.addWidget(set_repo_btn)
        
        # Analyze PR button
        def analyze_pr():
            pr_number, ok = QInputDialog.getInt(self, "Analyze PR", "PR number:")
            if ok:
                result = self._github_agent.analyze_pr(pr_number)
                if result:
                    msg = f"PR Analysis: {result.get('summary', {}).get('files_changed', 0)} files changed"
                    self._statusbar.showMessage(msg, 5000)
        
        analyze_btn = QPushButton("Analyze PR...")
        analyze_btn.clicked.connect(analyze_pr)
        layout.addWidget(analyze_btn)
        
        # Close button
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.close)
        layout.addWidget(close_btn)
        
        dialog.exec()
        log.info("GitHub integration dialog opened")

    # Phase 4: Real-time UI Update Handlers
    # NOTE: The following handlers are LEGACY stubs. The _todo_manager has been removed.
    # Todo state is now managed entirely by the bridge (CortexAgentBridge) via TodoWrite.
    
    def _on_todo_task_added(self, task_id: str):
        """Handle new todo task - LEGACY (todo manager removed)."""
        log.info(f"[TODO] _on_todo_task_added called (legacy): {task_id}")

    def _on_toggle_todo(self, task_id: str, completed: bool):
        """Handle todo toggle from UI.
        
        NOTE: The legacy _todo_manager has been removed. Todo state is now managed
        entirely by the bridge (CortexAgentBridge) and the JS UI. This handler
        logs the toggle but does not persist it - the UI already shows the toggled
        state visually and the bridge will emit updated todos on its next turn.
        """
        status = "completed" if completed else "reopened"
        log.info(f"[TODO] UI toggle: {task_id} -> {status}")
        # The todo state is managed by the bridge's TodoWrite tool and the JS UI.
        # No backend persistence needed here - the UI handles the visual state.

    def _on_todo_task_completed(self, task_id: str):
        """Handle completed todo - LEGACY (todo manager removed)."""
        log.info(f"[TODO] _on_todo_task_completed called (legacy): {task_id}")

    def _on_todo_task_updated(self, task_id: str):
        """Handle updated todo - LEGACY (todo manager removed)."""
        log.info(f"[TODO] _on_todo_task_updated called (legacy): {task_id}")

    def _on_ai_task_complete(self, response: str):
        """Show Windows toast notification when AI task completes."""
        try:
            if response and len(response) > 10:
                summary = response[:150].replace('\n', ' ').strip()
                show_task_complete_notification(summary)
            
            # TODO: POINTS SYSTEM - Disabled for development
            # Will be enabled when connecting to https://logic-practice.com backend
            """
            # Consume points based on actual response length
            self._consume_points_for_response(response)
            """
        except Exception:
            pass
    
    def _consume_points_for_response(self, response: str):
        """Consume points based on actual AI response tokens."""
        try:
            points_mgr = get_points_manager()
            perf_mode = self._get_current_performance_mode()
            
            # Estimate actual tokens used (response text)
            actual_tokens = len(response) // 4
            
            # Consume points
            result = points_mgr.consume_points(actual_tokens, perf_mode)
            
            log.info(
                f"[MainWindow] Points consumed: {actual_tokens:,} tokens × "
                f"{result['multiplier']}x = {result['points_consumed']:,} points. "
                f"Remaining: {result['remaining_balance']:,}"
            )
            
            # Update UI with remaining balance
            if hasattr(self, '_ai_chat') and hasattr(self._ai_chat, 'update_points_balance'):
                self._ai_chat.update_points_balance(result)
                
        except InsufficientPointsError as e:
            log.error(f"[MainWindow] Points consumption error: {e}")
            # Send warning to UI
            if hasattr(self, '_ai_chat') and hasattr(self._ai_chat, 'on_error'):
                self._ai_chat.on_error(
                    f"Points consumed exceeded balance. Please purchase more points.\n"
                    f"Required: {e.required:,}, Had: {e.balance:,}"
                )
        except Exception as e:
            log.warning(f"[MainWindow] Failed to consume points: {e}")
    
    def _get_current_performance_mode(self) -> str:
        """Get current performance mode from settings."""
        try:
            settings = get_settings()
            return settings.get("ai", "performance_mode", default="auto")
        except Exception:
            return "auto"

    def _on_title_generated(self, conversation_id: str, title: str):
        """Update the chat title in the sidebar when a new title is generated.
        
        Args:
            conversation_id: The ID of the conversation.
            title: The new title for the chat.
        """
        # Update the sidebar chat list with the new title
        self._refresh_sidebar_chat_list()
        """Handle auto-generated title - update chat tab and UI."""
        # Update the chat tab title via JavaScript bridge
        if hasattr(self, '_ai_chat') and hasattr(self._ai_chat, '_view'):
            safe_title = title.replace('"', '\\"').replace("'", "\\'")
            js_code = f"if(window.updateChatTitle) window.updateChatTitle('{conversation_id}', '{safe_title}');"
            self._ai_chat._view.page().runJavaScript(js_code)
            log.info(f"Chat title updated in UI: {title}")
        
        # Also update window title if this is current chat
        current_id = getattr(self._ai_chat, '_current_conversation_id', None)
        if current_id == conversation_id and title:
            self.setWindowTitle(f"Cortex - {title}")


def QApplication_instance():
    from PyQt6.QtWidgets import QApplication
    return QApplication.instance()
